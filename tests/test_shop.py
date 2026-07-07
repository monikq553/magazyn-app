import io
import zipfile
import unittest
from datetime import date, datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from werkzeug.datastructures import MultiDict

import app as warehouse_app


def sample_order():
    return (
        42,
        "SK/TEST/42",
        date(2026, 7, 6),
        "Klient Testowy",
        "ul. Testowa 1, Poznań",
        "123456789",
        "klient@example.com",
        25.0,
        "Przelew",
        "Oczekuje na płatność",
        "Towar zarezerwowany",
        "",
        "",
        "Dostawa testowa",
        "1234567890",
    )


def sample_items():
    return [
        (1, 1, "Deska dębowa", 2.0, 100.0, 123.0, 23.0),
        (2, 2, "Olej do drewna", 3.0, 20.0, 24.6, 23.0),
    ]


class ShopModuleTests(unittest.TestCase):
    def test_order_creation_with_multiple_products_commits_without_document_generation(self):
        cursor = MagicMock()
        cursor.fetchone.side_effect = [
            None,
            (42,),
            (1, "Deska dębowa", 10.0, "Drewno", 100.0, 23.0),
            (0.0,),
            (0.0,),
            (2, "Olej do drewna", 20.0, "Farby", 20.0, 23.0),
            (0.0,),
            (0.0,),
        ]
        connection = MagicMock()
        connection.cursor.return_value = cursor
        form = MultiDict(
            [
                ("order_number", "SK/TEST/42"),
                ("date", "2026-07-06"),
                ("customer_name", "Klient Testowy"),
                ("salesperson_email", "sales@example.com"),
                ("delivery_address", "ul. Testowa 1"),
                ("email", "klient@example.com"),
                ("shipping_cost", "25"),
                ("payment_method", "Przelew"),
                ("payment_status", "Oczekuje na płatność"),
                ("product_id", "1"),
                ("qty", "2"),
                ("product_id", "2"),
                ("qty", "3"),
            ]
        )

        with warehouse_app.app.test_request_context(
            "/sklep/orders", method="POST", data=form
        ):
            warehouse_app.session["user"] = "admin@example.com"
            warehouse_app.session["role"] = "admin"
            with (
                patch.object(warehouse_app, "db", return_value=connection),
                patch.object(warehouse_app, "update_shop_stage") as update_stage,
                patch.object(warehouse_app, "assign_order_salesperson") as assign_salesperson,
            ):
                response = warehouse_app.shop_create_order()

        self.assertEqual(response.status_code, 302)
        self.assertIn("/sklep/orders/42?created=1", response.location)
        connection.commit.assert_called_once()
        item_inserts = [
            call for call in cursor.execute.call_args_list
            if "INSERT INTO shop_order_items" in call.args[0]
        ]
        self.assertEqual(len(item_inserts), 2)
        self.assertFalse(
            any(
                "INSERT INTO shop_sales_documents" in call.args[0]
                for call in cursor.execute.call_args_list
            )
        )
        self.assertEqual(update_stage.call_count, 2)
        assign_salesperson.assert_called_once_with(
            cursor, 42, "sales@example.com", "admin@example.com"
        )

    def test_incomplete_order_returns_specific_form_message(self):
        form = MultiDict(
            [
                ("customer_name", "Klient"),
                ("delivery_address", "Adres"),
                ("product_id", "1"),
                ("qty", ""),
            ]
        )
        with warehouse_app.app.test_request_context(
            "/sklep/orders", method="POST", data=form
        ):
            warehouse_app.session["user"] = "admin@example.com"
            warehouse_app.session["role"] = "admin"
            with patch.object(
                warehouse_app,
                "render_shop_dashboard",
                side_effect=lambda message, status, data: (message, status),
            ):
                response = warehouse_app.shop_create_order()

        self.assertEqual(response[1], 400)
        self.assertIn("Pozycja 1: podaj ilość", response[0])

    def test_pdf_and_docx_contain_logo_customer_products_and_totals(self):
        payload = warehouse_app.create_shop_document_payload(
            sample_order(), sample_items()
        )
        docx_bytes = warehouse_app.simple_docx_bytes(payload)
        pdf_bytes = warehouse_app.shop_pdf_bytes(payload)

        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
            media = [
                name for name in archive.namelist()
                if name.startswith("word/media/")
            ]
            self.assertTrue(media)
            self.assertGreater(len(archive.read(media[0])), 100)

        document = Document(io.BytesIO(docx_bytes))
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        self.assertIn("Klient Testowy", text)
        self.assertIn("Deska dębowa", text)
        self.assertIn("RAZEM BRUTTO", text)
        self.assertIn("344.80 zł", text)

        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        self.assertIn(b"/Subtype /Image", pdf_bytes)
        self.assertGreater(len(pdf_bytes), 5000)

    def test_generated_document_is_upserted_and_saved_at_order(self):
        cursor = MagicMock()
        cursor.fetchone.side_effect = [
            None,
            sample_order(),
            (77, "DS/SK/TEST/42"),
            (88, "active"),
        ]
        cursor.fetchall.return_value = sample_items()

        with patch.object(warehouse_app, "update_shop_stage") as update_stage:
            document, payload = warehouse_app.generate_shop_sales_document(
                cursor, 42, "admin@example.com"
            )

        self.assertEqual(document[0], 77)
        self.assertEqual(payload["buyer"], "Klient Testowy")
        upserts = [
            call for call in cursor.execute.call_args_list
            if "INSERT INTO shop_sales_documents" in call.args[0]
        ]
        self.assertEqual(len(upserts), 1)
        self.assertIn("ON CONFLICT (order_id) DO UPDATE", upserts[0].args[0])
        params = upserts[0].args[1]
        self.assertGreater(len(params[3].adapted), 5000)
        self.assertGreater(len(params[4].adapted), 5000)
        update_stage.assert_called_once_with(
            cursor,
            42,
            "sales_document_generated",
            True,
            "admin@example.com",
        )

    def test_generate_document_action_commits_and_redirects_to_order(self):
        cursor = MagicMock()
        connection = MagicMock()
        connection.cursor.return_value = cursor
        payload = warehouse_app.create_shop_document_payload(
            sample_order(), sample_items()
        )
        with warehouse_app.app.test_request_context(
            "/sklep/orders/42/document/generate", method="POST"
        ):
            warehouse_app.session["user"] = "admin@example.com"
            warehouse_app.session["role"] = "admin"
            with (
                patch.object(warehouse_app, "db", return_value=connection),
                patch.object(
                    warehouse_app,
                    "generate_shop_sales_document",
                    return_value=((77, "DS/SK/TEST/42"), payload),
                ),
            ):
                response = warehouse_app.shop_generate_document(42)

        self.assertEqual(response.status_code, 302)
        self.assertIn("/sklep/orders/42?document=generated", response.location)
        connection.commit.assert_called_once()
        self.assertTrue(
            any(
                "INSERT INTO shop_notifications" in call.args[0]
                for call in cursor.execute.call_args_list
            )
        )

    def test_pdf_and_docx_download_endpoints(self):
        for fmt, content, mimetype in (
            ("pdf", b"%PDF-test", "application/pdf"),
            (
                "docx",
                b"PK-test",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        ):
            with self.subTest(fmt=fmt):
                cursor = MagicMock()
                cursor.fetchone.return_value = ("DS/SK/TEST/42", content)
                connection = MagicMock()
                connection.cursor.return_value = cursor
                with warehouse_app.app.test_request_context(
                    f"/sklep/documents/77/{fmt}"
                ):
                    warehouse_app.session["user"] = "admin@example.com"
                    warehouse_app.session["role"] = "admin"
                    with patch.object(
                        warehouse_app, "db", return_value=connection
                    ):
                        response = warehouse_app.shop_download_document(77, fmt)
                self.assertEqual(response.get_data(), content)
                self.assertEqual(response.mimetype, mimetype)
                self.assertIn(f".{fmt}", response.headers["Content-Disposition"])

    def test_migration_adds_columns_missing_from_legacy_shop_database(self):
        source = Path(warehouse_app.__file__).read_text(encoding="utf-8")
        migration_section = source[source.index("CREATE TABLE IF NOT EXISTS shop_orders("):]
        for column in ("payment_status", "tracking_number", "notes"):
            self.assertIn(f"ADD COLUMN IF NOT EXISTS {column}", migration_section)
        self.assertIn(
            "CREATE UNIQUE INDEX IF NOT EXISTS uq_shop_sales_documents_order",
            migration_section,
        )

    def test_stage_change_is_saved_with_actor_previous_and_new_value(self):
        cursor = MagicMock()
        cursor.fetchone.return_value = (
            "Nowe zamówienie",
            {"order_accepted": False},
            "",
        )
        cursor.rowcount = 1

        with warehouse_app.app.test_request_context("/sklep/orders/42/stage"):
            warehouse_app.session["user"] = "shop@example.com"
            warehouse_app.session["role"] = "shop"
            previous, current = warehouse_app.update_shop_stage(
                cursor,
                42,
                "order_accepted",
                True,
                "shop@example.com",
            )

        self.assertFalse(previous)
        self.assertTrue(current)
        stage_insert = next(
            call
            for call in cursor.execute.call_args_list
            if "INSERT INTO shop_order_stage_history" in call.args[0]
        )
        self.assertEqual(
            stage_insert.args[1],
            (
                42,
                "order_accepted",
                False,
                True,
                "Nowe zamówienie",
                "Przyjęte",
                "shop@example.com",
            ),
        )
        order_update = next(
            call
            for call in cursor.execute.call_args_list
            if "UPDATE shop_orders" in call.args[0] and "stages=" in call.args[0]
        )
        self.assertEqual(order_update.args[1][1], "Przyjęte")

    def test_stage_permissions_are_split_between_roles(self):
        cases = (
            ("warehouse", "packed", True),
            ("warehouse", "invoice_issued", False),
            ("accounting", "invoice_issued", True),
            ("accounting", "shipped", False),
            ("shop", "order_accepted", True),
            ("shop", "paid", False),
            ("admin", "paid", True),
            ("admin", "shipped", True),
        )
        for role, stage, expected in cases:
            with self.subTest(role=role, stage=stage):
                with warehouse_app.app.test_request_context("/"):
                    warehouse_app.session["role"] = role
                    self.assertEqual(
                        warehouse_app.shop_stage_can_edit(stage),
                        expected,
                    )

    def test_shop_filters_cover_invoice_receipt_payment_shipping_and_packing(self):
        cursor = MagicMock()
        cursor.fetchall.side_effect = [[], [], []]
        cursor.fetchone.return_value = (0, 0, 0, 0, 0, 0)
        filters = {
            "invoice": "yes",
            "receipt": "no",
            "paid": "yes",
            "shipped": "no",
            "packed": "yes",
            "salesperson": "Anna",
            "client": "Klient",
            "date": "2026-07-06",
            "status": "Przyjęte",
        }

        orders, notifications, reports, products = warehouse_app.shop_dashboard_data(
            cursor, filters
        )

        self.assertEqual((orders, notifications, products), ([], [], []))
        self.assertEqual(reports, (0, 0, 0, 0, 0, 0))
        order_query = cursor.execute.call_args_list[0].args[0]
        self.assertIn("a.invoice_issued", order_query)
        self.assertIn("a.receipt_issued", order_query)
        self.assertIn("o.stages->>'shipped'", order_query)
        self.assertIn("o.stages->>'packed'", order_query)

    def test_reservation_permissions_and_home_tile_are_available(self):
        self.assertIn("reservations", warehouse_app.ROLE_PERMISSIONS["admin"])
        self.assertIn("reservations", warehouse_app.ROLE_PERMISSIONS["warehouse"])
        self.assertIn("reservations", warehouse_app.ROLE_PERMISSIONS["shop"])
        self.assertIn("reservations", warehouse_app.ROLE_PERMISSIONS["sales"])
        home = Path(warehouse_app.app.root_path, "templates", "home.html").read_text(
            encoding="utf-8"
        )
        self.assertIn("/rezerwacje", home)
        self.assertIn("Rezerwacje", home)

    def test_reservation_pdf_is_checklist_for_warehouse(self):
        reservation = (
            7, "REZ/TEST/7", "2026-07-07", "Klient", "Anna Sprzedaż",
            "anna@example.com", "Drewno", "zatwierdzona", "Pilne",
        )
        items = [
            (
                1, 7, 11, 21, "Deska", "PACZKA-1", "20x120", 2.5, "m2",
                "Drewno", "A-01", 100, 123, False, None, "Bez uszkodzeń",
            )
        ]

        pdf = warehouse_app.reservation_pdf_bytes(reservation, items)

        self.assertTrue(pdf.startswith(b"%PDF"))
        self.assertGreater(len(pdf), 1000)

    def test_reserved_product_qty_combines_shop_and_reservation_modules(self):
        class Cursor:
            def __init__(self):
                self.calls = 0

            def execute(self, *_args):
                self.calls += 1

            def fetchone(self):
                return (3.0,) if self.calls == 1 else (2.0,)

        reserved = warehouse_app.reserved_product_qty(Cursor(), 11)

        self.assertEqual(reserved, 5.0)

    def test_approved_reservation_rejects_quantity_above_available_stock(self):
        cursor = MagicMock()
        cursor.fetchone.side_effect = [
            ("Anna Sprzedaż",),
            (7,),
            (11, "Deska", 1.0, "m2", "Drewno", 100.0, 23.0),
            (0.0,),
            (0.0,),
        ]
        connection = MagicMock()
        connection.cursor.return_value = cursor
        form = MultiDict(
            [
                ("action", "approve"),
                ("customer_name", "Klient"),
                ("salesperson_email", "anna@example.com"),
                ("date", "2026-07-07"),
                ("warehouse", "Drewno"),
                ("product_id", "11"),
                ("qty", "2"),
            ]
        )

        with warehouse_app.app.test_request_context(
            "/rezerwacje", method="POST", data=form
        ):
            warehouse_app.session["user"] = "admin@example.com"
            warehouse_app.session["role"] = "admin"
            with patch.object(warehouse_app, "db", return_value=connection):
                response = warehouse_app.create_reservation()

        self.assertEqual(response[1], 400)
        self.assertIn("dostępne jest tylko 1", response[0])
        connection.rollback.assert_called_once()

    def test_reservation_issue_reduces_product_and_package_stock(self):
        cursor = MagicMock()
        cursor.fetchone.return_value = ("skompletowana",)
        cursor.fetchall.return_value = [(1, 11, 21, 2.0, "Drewno")]
        cursor.rowcount = 1
        connection = MagicMock()
        connection.cursor.return_value = cursor

        with warehouse_app.app.test_request_context(
            "/rezerwacje/7/issue", method="POST"
        ):
            warehouse_app.session["user"] = "warehouse@example.com"
            warehouse_app.session["role"] = "warehouse"
            with patch.object(warehouse_app, "db", return_value=connection):
                response = warehouse_app.issue_reservation(7)

        self.assertEqual(response.status_code, 302)
        executed = "\n".join(call.args[0] for call in cursor.execute.call_args_list)
        self.assertIn("UPDATE products SET qty=qty-%s", executed)
        self.assertIn("UPDATE packages SET qty=qty-%s", executed)
        self.assertIn("status='wydana'", executed)
        connection.commit.assert_called_once()

    def test_order_detail_renders_all_stages_and_immediate_save_endpoint(self):
        order = sample_order() + (
            False,
            "admin@example.com",
            datetime(2026, 7, 6, 12, 0),
            datetime(2026, 7, 6, 12, 0),
            {},
        )
        stage_rows = [
            {
                "key": key,
                "label": label,
                "checked": False,
                "can_edit": True,
            }
            for key, label, _permission in warehouse_app.SHOP_ORDER_STAGES
        ]
        with warehouse_app.app.test_request_context("/sklep/orders/42"):
            warehouse_app.session["user"] = "admin@example.com"
            warehouse_app.session["role"] = "admin"
            html = warehouse_app.render_template(
                "shop_order.html",
                order=order,
                items=[],
                document=None,
                history=[],
                stage_history=[],
                stage_labels={
                    key: value["label"]
                    for key, value in warehouse_app.SHOP_STAGE_BY_KEY.items()
                },
                stage_rows=stage_rows,
                statuses=warehouse_app.SHOP_STATUS_FLOW,
                accounting=None,
                payment_methods=warehouse_app.ACCOUNTING_PAYMENT_METHODS,
                can_ship=False,
            )

        self.assertEqual(html.count('class="order-stage-checkbox"'), 16)
        self.assertIn("/sklep/orders/42/stage", html)
        self.assertIn("X-CSRF-Token", html)
        self.assertIn("Status przed", html)
        self.assertIn("Status po", html)


if __name__ == "__main__":
    unittest.main()
