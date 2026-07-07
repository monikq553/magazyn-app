import os
import logging
import math
import secrets
import threading
from flask import (
    Flask,
    Response,
    g,
    has_request_context,
    jsonify,
    make_response,
    redirect,
    render_template,
    request,
    session,
)
from functools import wraps
import psycopg2
from psycopg2.pool import ThreadedConnectionPool
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.exceptions import HTTPException
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
import json
import io
import pandas as pd
from urllib import request as urlrequest
from urllib import error as urlerror
from urllib.parse import parse_qsl, quote_plus, urlencode, urlsplit, urlunsplit
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import (
    Image as ReportLabImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from xml.sax.saxutils import escape
import firebase_admin
from firebase_admin import credentials, auth, storage
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_caching import Cache
from backup_service import (
    decrypt_backup,
    encrypt_backup,
    export_database,
    parse_backup,
    restore_database,
)
from general_import import (
    ENTITY_FIELDS,
    ENTITY_GROUPS,
    ENTITY_LABELS,
    FIELD_LABELS,
    duplicate_identity,
    normalize_row as normalize_import_row,
    parse_workbook,
    validate_row as validate_import_row,
)
from issue_import import (
    ISSUE_IMPORT_FIELDS,
    ISSUE_IMPORT_LABELS,
    issue_history_pdf,
    issue_history_xlsx,
    issue_mapping,
    issue_sheet_selected,
    normalize_issue_row,
    validate_issue_row,
)

load_dotenv()

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

app = Flask(__name__)
IS_PRODUCTION = os.environ.get("RENDER", "").lower() == "true"
configured_secret_key = os.environ.get("SECRET_KEY")
if IS_PRODUCTION and not configured_secret_key:
    logger.critical(
        "SECRET_KEY is missing. A temporary key is used; sessions will be invalidated "
        "after every restart. Configure SECRET_KEY immediately."
    )
app.secret_key = configured_secret_key or secrets.token_hex(32)
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=IS_PRODUCTION,
    SESSION_COOKIE_NAME="magazyn_csrf_session",
    MAX_CONTENT_LENGTH=10 * 1024 * 1024,
)
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1)


@app.route("/manifest.json")
def web_manifest():
    response = make_response(app.send_static_file("manifest.json"))
    response.headers["Content-Type"] = "application/manifest+json"
    return response


@app.route("/service-worker.js")
def service_worker():
    response = make_response(app.send_static_file("service-worker.js"))
    response.headers["Content-Type"] = "application/javascript; charset=utf-8"
    response.headers["Cache-Control"] = "no-cache"
    response.headers["Service-Worker-Allowed"] = "/"
    return response


INVESTMENT_WAREHOUSE = "Inwestycja Suwaj"
MAIN_WAREHOUSE = "Drewno"
WAREHOUSES = (
    "Drewno",
    "Farby",
    "Śruby i wkręty",
    "Pellet",
    "Golden Oak",
    "Inne",
    INVESTMENT_WAREHOUSE,
)
UNITS = {"m2", "m3", "szt", "l", "opak", "kg"}
FIREBASE_CONFIG = {
    "apiKey": os.environ.get("FIREBASE_API_KEY", ""),
    "authDomain": os.environ.get("FIREBASE_AUTH_DOMAIN", ""),
    "projectId": os.environ.get("FIREBASE_PROJECT_ID", ""),
    "storageBucket": os.environ.get("FIREBASE_STORAGE_BUCKET", ""),
    "messagingSenderId": os.environ.get("FIREBASE_MESSAGING_SENDER_ID", ""),
    "appId": os.environ.get("FIREBASE_APP_ID", ""),
    "measurementId": os.environ.get("FIREBASE_MEASUREMENT_ID", ""),
}
ADMIN_EMAILS = {e.strip().lower() for e in os.environ.get("ADMIN_EMAILS", "").split(",") if e.strip()}
ALLOWED_ORIGINS = {
    origin.strip().rstrip("/")
    for origin in os.environ.get(
        "ALLOWED_ORIGINS",
        "https://pmagazyn.pl,https://app.pmagazyn.pl",
    ).split(",")
    if origin.strip()
}
ROLES = {"admin", "warehouse", "shop", "accounting", "sales"}
ROLE_LABELS = {
    "admin": "Administrator",
    "warehouse": "Magazynier",
    "shop": "Obsługa sklepu internetowego",
    "accounting": "Księgowość",
    "sales": "Handlowiec",
}
LOGIN_LOCK_MODE = os.environ.get("LOGIN_LOCK_MODE", "temporary").strip().lower()
if LOGIN_LOCK_MODE not in {"temporary", "admin"}:
    LOGIN_LOCK_MODE = "temporary"
try:
    LOGIN_LOCK_MINUTES = max(1, min(int(os.environ.get("LOGIN_LOCK_MINUTES", "15")), 1440))
except ValueError:
    LOGIN_LOCK_MINUTES = 15
ROLE_PERMISSIONS = {
    "admin": {"dashboard", "inventory", "inventory_manage", "receive", "issue", "history", "reports", "users", "backups", "shop", "accounting", "reservations"},
    "warehouse": {"dashboard", "inventory", "inventory_manage", "receive", "issue", "history", "shop", "reservations"},
    "shop": {"dashboard", "inventory", "shop", "history", "reservations"},
    "accounting": {"dashboard", "history", "reports", "accounting"},
    "sales": {"dashboard", "inventory", "shop", "reservations"},
}
ENDPOINT_PERMISSIONS = {
    "dashboard_page_view": "dashboard",
    "magazyny": "inventory",
    "magazyn": "inventory",
    "packages_for_product": "inventory",
    "package_lookup": "inventory",
    "add_product": "inventory_manage",
    "costs": "inventory_manage",
    "add_cost": "inventory_manage",
    "przyjecie": "receive",
    "receive_doc": "receive",
    "import_excel": "receive",
    "wydanie": "issue",
    "issue_doc": "issue",
    "inwestycja_suwaj_page_view": "inventory_manage",
    "inwestycja_suwaj_magazyn": "inventory_manage",
    "inwestycja_suwaj_przyjecie": "receive",
    "inwestycja_suwaj_receive_doc": "receive",
    "inwestycja_suwaj_wydanie": "issue",
    "inwestycja_suwaj_issue_doc": "issue",
    "historia": "history",
    "doc_detail": "history",
    "view_issue_doc_photo": "history",
    "add_issue_doc_photos": "issue",
    "edit_doc": "issue",
    "report": "reports",
    "report_pdf": "reports",
    "shop_orders_page": "shop",
    "add_shop_order": "shop",
    "reservations_page": "reservations",
    "create_reservation": "reservations",
    "reservation_detail": "reservations",
    "reservation_pdf": "reservations",
    "regenerate_reservation_pdf": "reservations",
    "update_reservation_item": "reservations",
    "complete_reservation": "reservations",
    "cancel_reservation": "reservations",
    "issue_reservation": "reservations",
}
FIREBASE_ADMIN_READY = False
FIREBASE_ADMIN_ERROR = ""
DB_POOL = None
DB_INIT_LOCK = threading.Lock()
DB_INITIALIZED = False
DAILY_BACKUP_LOCK = threading.Lock()
cache = Cache(config={"CACHE_TYPE": "SimpleCache", "CACHE_DEFAULT_TIMEOUT": 60})
cache.init_app(app)
limiter = Limiter(
    key_func=get_remote_address,
    default_limits=["200 per minute"],
    storage_uri=os.environ.get("RATELIMIT_STORAGE_URI", "memory://"),
)
limiter.init_app(app)


class PooledConn:
    def __init__(self, raw_conn):
        self._raw = raw_conn
        self._closed = False

    def __getattr__(self, item):
        return getattr(self._raw, item)

    def close(self):
        global DB_POOL
        if self._closed:
            return
        self._closed = True
        try:
            self._raw.rollback()
        except Exception:
            logger.warning("Database rollback before returning connection failed.", exc_info=True)
        try:
            if DB_POOL:
                DB_POOL.putconn(self._raw)
            else:
                self._raw.close()
        except Exception:
            logger.warning("Returning database connection to the pool failed.", exc_info=True)
            try:
                self._raw.close()
            except Exception:
                logger.warning("Closing failed database connection failed.", exc_info=True)


def init_db_pool():
    global DB_POOL
    if DB_POOL:
        return
    dsn = os.environ.get("DATABASE_URL")
    if not dsn:
        raise RuntimeError("Brak wymaganej zmiennej środowiskowej DATABASE_URL.")
    try:
        pool_size = max(2, int(os.environ.get("DB_POOL_SIZE", "5")))
    except (TypeError, ValueError):
        pool_size = 5
        logger.warning("Invalid DB_POOL_SIZE; using 5.")
    candidates = [dsn]
    try:
        parsed = urlsplit(dsn)
        hostname = parsed.hostname or ""
        if os.environ.get("RENDER") and hostname.endswith("-postgres.render.com"):
            internal_hostname = hostname.split(".", 1)[0]
            userinfo = parsed.netloc.rsplit("@", 1)[0] if "@" in parsed.netloc else ""
            port = f":{parsed.port}" if parsed.port else ""
            internal_netloc = f"{userinfo}@{internal_hostname}{port}" if userinfo else f"{internal_hostname}{port}"
            internal_query = urlencode(
                [(key, value) for key, value in parse_qsl(parsed.query) if key.lower() != "sslmode"]
            )
            internal_dsn = urlunsplit(
                (parsed.scheme, internal_netloc, parsed.path, internal_query, parsed.fragment)
            )
            candidates.insert(0, internal_dsn)
    except ValueError:
        logger.warning("DATABASE_URL is not a standard URL; using it without normalization.")

    last_error = None
    for candidate in candidates:
        try:
            DB_POOL = ThreadedConnectionPool(
                minconn=1,
                maxconn=pool_size,
                dsn=candidate,
                connect_timeout=10,
                application_name="magazyn-app",
            )
            break
        except psycopg2.OperationalError as exc:
            last_error = exc
            logger.warning("Database connection candidate failed; trying fallback if available.")
    if DB_POOL is None:
        raise last_error or RuntimeError("Nie udało się utworzyć puli połączeń z bazą.")
    logger.info("DB pool initialized.")


def init_firebase_admin():
    global FIREBASE_ADMIN_READY, FIREBASE_ADMIN_ERROR
    if firebase_admin._apps:
        FIREBASE_ADMIN_READY = True
        return
    raw = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON")
    if not raw:
        FIREBASE_ADMIN_READY = False
        FIREBASE_ADMIN_ERROR = "Brak FIREBASE_SERVICE_ACCOUNT_JSON w zmiennych środowiskowych."
        logger.error("Firebase Admin init failed: FIREBASE_SERVICE_ACCOUNT_JSON is missing.")
        return
    try:
        # Read strictly from env and parse JSON payload (Render-compatible).
        normalized_raw = raw.strip()
        if (normalized_raw.startswith("'") and normalized_raw.endswith("'")) or (
            normalized_raw.startswith('"') and normalized_raw.endswith('"')
        ):
            normalized_raw = normalized_raw[1:-1]

        try:
            service_account = json.loads(normalized_raw)
        except json.JSONDecodeError:
            # Some dashboards escape JSON one level too deep; unescape and try again.
            service_account = json.loads(bytes(normalized_raw, "utf-8").decode("unicode_escape"))

        # If private_key is escaped (\\n), normalize to real newlines.
        private_key = service_account.get("private_key")
        if isinstance(private_key, str):
            service_account["private_key"] = private_key.replace("\\n", "\n")

        cred = credentials.Certificate(service_account)
        firebase_admin.initialize_app(cred)
        FIREBASE_ADMIN_READY = True
        FIREBASE_ADMIN_ERROR = ""
        logger.info("Firebase Admin initialized successfully.")
    except Exception:
        FIREBASE_ADMIN_READY = False
        FIREBASE_ADMIN_ERROR = "Nieprawidłowy FIREBASE_SERVICE_ACCOUNT_JSON."
        logger.exception("Firebase Admin init failed: invalid FIREBASE_SERVICE_ACCOUNT_JSON.")


def get_missing_firebase_web_envs():
    key_map = {
        "FIREBASE_API_KEY": "apiKey",
        "FIREBASE_AUTH_DOMAIN": "authDomain",
        "FIREBASE_PROJECT_ID": "projectId",
        "FIREBASE_APP_ID": "appId",
    }
    return [env_key for env_key, cfg_key in key_map.items() if not FIREBASE_CONFIG.get(cfg_key)]


if get_missing_firebase_web_envs():
    logger.error(
        "Firebase Web configuration is incomplete: %s",
        ", ".join(get_missing_firebase_web_envs()),
    )
if not os.environ.get("DATABASE_URL"):
    logger.error("DATABASE_URL is missing; database-backed pages will be unavailable.")
if not ADMIN_EMAILS:
    logger.warning(
        "ADMIN_EMAILS is empty; automatic creation of the first administrator is disabled."
    )
if not FIREBASE_CONFIG.get("storageBucket"):
    logger.warning("FIREBASE_STORAGE_BUCKET is missing; database backups are unavailable.")
if not os.environ.get("BACKUP_ENCRYPTION_KEY"):
    logger.warning("BACKUP_ENCRYPTION_KEY is missing; database backups are unavailable.")
init_firebase_admin()


# 🔥 DB
def db():
    init_db_pool()
    raw = DB_POOL.getconn()
    try:
        raw.autocommit = False
        with raw.cursor() as cur:
            cur.execute("SET statement_timeout = 15000")
        connection = PooledConn(raw)
        if has_request_context():
            connections = getattr(g, "_database_connections", None)
            if connections is None:
                connections = []
                g._database_connections = connections
            connections.append(connection)
        return connection
    except Exception:
        DB_POOL.putconn(raw, close=True)
        raise


@app.teardown_request
def close_request_database_connections(_error=None):
    for connection in g.pop("_database_connections", []):
        connection.close()


# 🔥 INIT DB
def run_db_migrations():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT pg_advisory_xact_lock(67431029)")

    # USERS
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users(
        id SERIAL PRIMARY KEY,
        firebase_uid TEXT UNIQUE NOT NULL,
        first_name TEXT NOT NULL DEFAULT '',
        last_name TEXT NOT NULL DEFAULT '',
        email TEXT UNIQUE NOT NULL,
        role TEXT NOT NULL DEFAULT 'warehouse',
        status TEXT NOT NULL DEFAULT 'active',
        phone TEXT,
        must_change_password BOOLEAN NOT NULL DEFAULT FALSE,
        last_login_at TIMESTAMPTZ,
        password_changed_at TIMESTAMPTZ,
        created_by TEXT,
        password_reset_by TEXT,
        failed_login_attempts INTEGER NOT NULL DEFAULT 0,
        last_failed_login_at TIMESTAMPTZ,
        locked_until TIMESTAMPTZ,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS schema_migrations(
        name TEXT PRIMARY KEY,
        applied_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute(
        "SELECT 1 FROM schema_migrations WHERE name='firebase_users_v2'"
    )
    if not cur.fetchone():
        cur.execute("""
            SELECT EXISTS (
                SELECT 1 FROM information_schema.columns
                WHERE table_name='users' AND column_name='password'
            )
        """)
        if cur.fetchone()[0]:
            cur.execute("DROP TABLE users")
            cur.execute("""
                CREATE TABLE users(
                    id SERIAL PRIMARY KEY,
                    firebase_uid TEXT UNIQUE NOT NULL,
                    first_name TEXT NOT NULL DEFAULT '',
                    last_name TEXT NOT NULL DEFAULT '',
                    email TEXT UNIQUE NOT NULL,
                    role TEXT NOT NULL DEFAULT 'warehouse',
                    status TEXT NOT NULL DEFAULT 'active',
                    phone TEXT,
                    must_change_password BOOLEAN NOT NULL DEFAULT FALSE,
                    last_login_at TIMESTAMPTZ,
                    password_changed_at TIMESTAMPTZ,
                    created_by TEXT,
                    password_reset_by TEXT,
                    failed_login_attempts INTEGER NOT NULL DEFAULT 0,
                    last_failed_login_at TIMESTAMPTZ,
                    locked_until TIMESTAMPTZ,
                    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
                    updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
                )
            """)
        else:
            # Tabela może już zawierać poprawne konta Firebase, np. po odtworzeniu
            # starszej kopii bez wpisu w schema_migrations. Nie usuwaj ich ponownie.
            logger.info("Existing Firebase-compatible users table preserved.")
        cur.execute(
            "INSERT INTO schema_migrations(name) VALUES ('firebase_users_v2')"
        )

    # PRODUCTS
    cur.execute("""
    CREATE TABLE IF NOT EXISTS products(
        id SERIAL PRIMARY KEY,
        name TEXT,
        qty REAL,
        unit TEXT,
        warehouse TEXT,
        price_netto REAL DEFAULT 0,
        vat REAL DEFAULT 0,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)

    # PACKAGES
    cur.execute("""
    CREATE TABLE IF NOT EXISTS packages(
        id SERIAL PRIMARY KEY,
        product_id INTEGER,
        number TEXT,
        qty REAL,
        warehouse TEXT
    );
    """)

    # DOCS
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_docs(
        id SERIAL PRIMARY KEY,
        date TEXT,
        kontrahent TEXT,
        warehouse TEXT,
        image TEXT,
        doc_number TEXT
    );
    """)

    # COSTS
    cur.execute("""
    CREATE TABLE IF NOT EXISTS costs(
        id SERIAL PRIMARY KEY,
        name TEXT,
        amount REAL,
        date TEXT,
        warehouse_source BOOLEAN DEFAULT FALSE,
        description TEXT
    );
    """)

    # ITEMS (bez kombinowania w środku)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_items(
        id SERIAL PRIMARY KEY,
        doc_id INTEGER,
        product_id INTEGER,
        qty REAL,
        warehouse TEXT,
        package_id INTEGER,
        price_netto REAL DEFAULT 0,
        price_brutto REAL DEFAULT 0
    );
    """)

    # 🔥 KLUCZOWE — aktualizacja starej bazy
    cur.execute("""
    ALTER TABLE issue_items
    ADD COLUMN IF NOT EXISTS package_id INTEGER;
    """)
    cur.execute("""
    ALTER TABLE issue_items
    ADD COLUMN IF NOT EXISTS warehouse TEXT;
    """)
    cur.execute("""
    ALTER TABLE issue_items
    ADD COLUMN IF NOT EXISTS price_netto REAL DEFAULT 0;
    """)
    cur.execute("""
    ALTER TABLE issue_items
    ADD COLUMN IF NOT EXISTS price_brutto REAL DEFAULT 0;
    """)
    cur.execute("""
    ALTER TABLE packages
    ADD COLUMN IF NOT EXISTS warehouse TEXT;
    """)
    cur.execute("""
    ALTER TABLE packages
    ADD COLUMN IF NOT EXISTS number TEXT;
    """)

    # migracja kompatybilności: stare bazy mogły mieć kolumnę package_number zamiast number
    cur.execute("""
    SELECT EXISTS (
        SELECT 1
        FROM information_schema.columns
        WHERE table_name = 'packages' AND column_name = 'package_number'
    )
    """)
    has_package_number = cur.fetchone()[0]
    if has_package_number:
        cur.execute("""
        UPDATE packages
        SET number = COALESCE(number, package_number::TEXT)
        WHERE number IS NULL
        """)

    cur.execute("ALTER TABLE packages ADD COLUMN IF NOT EXISTS initial_qty REAL")
    cur.execute("ALTER TABLE packages ADD COLUMN IF NOT EXISTS status TEXT DEFAULT 'active'")
    cur.execute("ALTER TABLE packages ADD COLUMN IF NOT EXISTS archived_at TIMESTAMPTZ")
    cur.execute("ALTER TABLE issue_docs ADD COLUMN IF NOT EXISTS movement_type TEXT")
    cur.execute("ALTER TABLE issue_docs ADD COLUMN IF NOT EXISTS created_at TIMESTAMPTZ DEFAULT NOW()")
    cur.execute("ALTER TABLE issue_docs ADD COLUMN IF NOT EXISTS voided_at TIMESTAMPTZ")
    cur.execute("ALTER TABLE issue_docs ADD COLUMN IF NOT EXISTS voided_by TEXT")
    cur.execute("ALTER TABLE issue_items ADD COLUMN IF NOT EXISTS package_number TEXT")
    cur.execute("ALTER TABLE issue_items ADD COLUMN IF NOT EXISTS dimension TEXT")
    cur.execute("ALTER TABLE issue_items ADD COLUMN IF NOT EXISTS species TEXT")
    cur.execute("ALTER TABLE issue_items ADD COLUMN IF NOT EXISTS notes TEXT")
    cur.execute("ALTER TABLE costs ADD COLUMN IF NOT EXISTS source_warehouse TEXT")
    cur.execute("UPDATE products SET qty=0 WHERE qty IS NULL OR qty < 0 OR qty::text='NaN'")
    cur.execute("UPDATE packages SET qty=0 WHERE qty IS NULL OR qty < 0 OR qty::text='NaN'")
    cur.execute("UPDATE packages SET initial_qty=qty WHERE initial_qty IS NULL")
    cur.execute("""
        UPDATE packages
        SET status=CASE WHEN qty <= 0 THEN 'issued' ELSE 'active' END,
            archived_at=CASE WHEN qty <= 0 THEN COALESCE(archived_at, NOW()) ELSE archived_at END
        WHERE status IS NULL OR status NOT IN ('active', 'issued', 'cancelled')
           OR (qty <= 0 AND status='active')
    """)
    cur.execute("""
        UPDATE issue_docs
        SET movement_type=CASE
            WHEN upper(COALESCE(doc_number, '')) LIKE 'PZ%' THEN 'PZ'
            WHEN upper(COALESCE(doc_number, '')) LIKE 'RW%' THEN 'RW'
            WHEN upper(COALESCE(doc_number, '')) LIKE 'WZ%' THEN 'WZ'
            ELSE movement_type
        END
        WHERE movement_type IS NULL
    """)
    cur.execute("""
        UPDATE issue_items i SET package_number=p.number
        FROM packages p
        WHERE i.package_id=p.id AND i.package_number IS NULL
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_products_warehouse_name ON products(warehouse, lower(name))")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_packages_product_active ON packages(product_id, warehouse, status)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_packages_number ON packages(warehouse, lower(number))")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_issue_items_doc ON issue_items(doc_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_issue_docs_date ON issue_docs(date)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_issue_docs_number ON issue_docs(lower(doc_number))")
    cur.execute("""
        DO $$
        BEGIN
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='products_qty_nonnegative') THEN
                ALTER TABLE products
                ADD CONSTRAINT products_qty_nonnegative CHECK (qty >= 0);
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='packages_qty_nonnegative') THEN
                ALTER TABLE packages
                ADD CONSTRAINT packages_qty_nonnegative CHECK (qty >= 0);
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='issue_items_qty_positive') THEN
                ALTER TABLE issue_items
                ADD CONSTRAINT issue_items_qty_positive CHECK (qty > 0) NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='products_qty_finite_nonnegative') THEN
                ALTER TABLE products
                ADD CONSTRAINT products_qty_finite_nonnegative
                CHECK (qty >= 0 AND qty::text <> 'NaN') NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='packages_qty_finite_nonnegative') THEN
                ALTER TABLE packages
                ADD CONSTRAINT packages_qty_finite_nonnegative
                CHECK (qty >= 0 AND qty::text <> 'NaN') NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='packages_status_valid') THEN
                ALTER TABLE packages
                ADD CONSTRAINT packages_status_valid
                CHECK (status IN ('active', 'issued', 'cancelled')) NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='packages_product_fk') THEN
                ALTER TABLE packages
                ADD CONSTRAINT packages_product_fk
                FOREIGN KEY (product_id) REFERENCES products(id) ON DELETE RESTRICT NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='issue_items_doc_fk') THEN
                ALTER TABLE issue_items
                ADD CONSTRAINT issue_items_doc_fk
                FOREIGN KEY (doc_id) REFERENCES issue_docs(id) ON DELETE RESTRICT NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='issue_items_product_fk') THEN
                ALTER TABLE issue_items
                ADD CONSTRAINT issue_items_product_fk
                FOREIGN KEY (product_id) REFERENCES products(id) ON DELETE RESTRICT NOT VALID;
            END IF;
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='issue_items_package_fk') THEN
                ALTER TABLE issue_items
                ADD CONSTRAINT issue_items_package_fk
                FOREIGN KEY (package_id) REFERENCES packages(id) ON DELETE SET NULL NOT VALID;
            END IF;
            UPDATE users SET role='warehouse'
            WHERE role IS NULL OR role NOT IN ('admin', 'warehouse', 'shop', 'accounting', 'sales');
            UPDATE users SET status='active'
            WHERE status IS NULL OR status NOT IN ('active', 'blocked', 'inactive');
            IF EXISTS (SELECT 1 FROM pg_constraint WHERE conname='users_role_valid') THEN
                ALTER TABLE users DROP CONSTRAINT users_role_valid;
            END IF;
            ALTER TABLE users
            ADD CONSTRAINT users_role_valid CHECK (role IN ('admin', 'warehouse', 'shop', 'accounting', 'sales'));
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='users_status_valid') THEN
                ALTER TABLE users
                ADD CONSTRAINT users_status_valid CHECK (status IN ('active', 'blocked', 'inactive'));
            END IF;
        END $$;
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS action_logs(
        id SERIAL PRIMARY KEY,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        actor_email TEXT,
        actor_uid TEXT,
        action TEXT NOT NULL,
        entity_type TEXT,
        entity_id TEXT,
        details JSONB NOT NULL DEFAULT '{}'::jsonb
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_action_logs_created_at ON action_logs(created_at DESC)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_action_logs_actor ON action_logs(lower(actor_email))")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS backup_runs(
        id SERIAL PRIMARY KEY,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        created_by TEXT NOT NULL,
        status TEXT NOT NULL,
        object_name TEXT,
        size_bytes BIGINT,
        checksum TEXT,
        error TEXT
    );
    """)
    cur.execute("ALTER TABLE users DROP CONSTRAINT IF EXISTS users_role_valid")
    cur.execute("""
        ALTER TABLE users
        ADD CONSTRAINT users_role_valid
        CHECK (role IN ('admin', 'warehouse', 'shop', 'accounting', 'sales'))
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_orders(
        id SERIAL PRIMARY KEY,
        order_number TEXT UNIQUE NOT NULL,
        order_date DATE NOT NULL,
        customer_name TEXT NOT NULL,
        delivery_address TEXT NOT NULL,
        phone TEXT,
        email TEXT,
        shipping_cost REAL NOT NULL DEFAULT 0,
        payment_method TEXT,
        payment_status TEXT NOT NULL DEFAULT 'Oczekuje na płatność',
        status TEXT NOT NULL DEFAULT 'Nowe zamówienie',
        sales_document_number TEXT,
        tracking_number TEXT,
        notes TEXT,
        nip TEXT,
        stages JSONB NOT NULL DEFAULT '{}'::jsonb,
        document_confirmed BOOLEAN NOT NULL DEFAULT FALSE,
        created_by TEXT,
        salesperson_email TEXT,
        salesperson_name TEXT,
        salesperson_assigned_by TEXT,
        salesperson_assigned_at TIMESTAMPTZ,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_order_items(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        product_id INTEGER NOT NULL REFERENCES products(id) ON DELETE RESTRICT,
        product_name TEXT NOT NULL,
        qty REAL NOT NULL CHECK (qty > 0),
        price_netto REAL NOT NULL DEFAULT 0,
        price_brutto REAL NOT NULL DEFAULT 0,
        vat REAL NOT NULL DEFAULT 0,
        warehouse TEXT NOT NULL,
        reserved_qty REAL NOT NULL DEFAULT 0,
        issued_qty REAL NOT NULL DEFAULT 0
    );
    """)
    # Reconcile databases created by the older, incompatible /shop/orders module.
    # The legacy columns are retained as nullable so existing data can be migrated
    # without destructive table replacement.
    cur.execute("""
        ALTER TABLE shop_orders
            ADD COLUMN IF NOT EXISTS number TEXT,
            ADD COLUMN IF NOT EXISTS customer TEXT,
            ADD COLUMN IF NOT EXISTS warehouse_status TEXT,
            ADD COLUMN IF NOT EXISTS accounting_status TEXT,
            ADD COLUMN IF NOT EXISTS shipping_status TEXT,
            ADD COLUMN IF NOT EXISTS invoice_number TEXT,
            ADD COLUMN IF NOT EXISTS stage TEXT,
            ADD COLUMN IF NOT EXISTS order_number TEXT,
            ADD COLUMN IF NOT EXISTS order_date DATE,
            ADD COLUMN IF NOT EXISTS customer_name TEXT,
            ADD COLUMN IF NOT EXISTS delivery_address TEXT,
            ADD COLUMN IF NOT EXISTS phone TEXT,
            ADD COLUMN IF NOT EXISTS email TEXT,
            ADD COLUMN IF NOT EXISTS shipping_cost REAL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS payment_method TEXT,
            ADD COLUMN IF NOT EXISTS payment_status TEXT DEFAULT 'Oczekuje na płatność',
            ADD COLUMN IF NOT EXISTS status TEXT,
            ADD COLUMN IF NOT EXISTS sales_document_number TEXT,
            ADD COLUMN IF NOT EXISTS tracking_number TEXT,
            ADD COLUMN IF NOT EXISTS notes TEXT,
            ADD COLUMN IF NOT EXISTS nip TEXT,
            ADD COLUMN IF NOT EXISTS stages JSONB DEFAULT '{}'::jsonb,
            ADD COLUMN IF NOT EXISTS document_confirmed BOOLEAN DEFAULT FALSE,
            ADD COLUMN IF NOT EXISTS created_by TEXT,
            ADD COLUMN IF NOT EXISTS salesperson_email TEXT,
            ADD COLUMN IF NOT EXISTS salesperson_name TEXT,
            ADD COLUMN IF NOT EXISTS salesperson_assigned_by TEXT,
            ADD COLUMN IF NOT EXISTS salesperson_assigned_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS created_at TIMESTAMPTZ DEFAULT NOW(),
            ADD COLUMN IF NOT EXISTS updated_at TIMESTAMPTZ DEFAULT NOW()
    """)
    cur.execute("ALTER TABLE shop_orders ALTER COLUMN number DROP NOT NULL")
    cur.execute("ALTER TABLE shop_orders ALTER COLUMN customer DROP NOT NULL")
    cur.execute("""
        UPDATE shop_orders
        SET order_number=COALESCE(NULLIF(order_number, ''), NULLIF(number, ''), 'LEGACY/' || id),
            order_date=COALESCE(order_date, created_at::date, CURRENT_DATE),
            customer_name=COALESCE(NULLIF(customer_name, ''), NULLIF(customer, ''), 'Nieznany klient'),
            delivery_address=COALESCE(delivery_address, ''),
            shipping_cost=COALESCE(shipping_cost, 0),
            payment_status=COALESCE(NULLIF(payment_status, ''), 'Oczekuje na płatność'),
            status=COALESCE(NULLIF(status, ''),
                CASE stage
                    WHEN 'sent' THEN 'Wysłane'
                    WHEN 'completed' THEN 'Zakończone'
                    WHEN 'cancelled' THEN 'Anulowane'
                    WHEN 'packed' THEN 'Spakowane'
                    WHEN 'document_created' THEN 'Dokument wystawiony'
                    WHEN 'reserved' THEN 'Towar zarezerwowany'
                    ELSE 'Nowe zamówienie'
                END),
            stages=COALESCE(stages, '{}'::jsonb),
            document_confirmed=COALESCE(document_confirmed, FALSE),
            created_at=COALESCE(created_at, NOW()),
            updated_at=COALESCE(updated_at, NOW())
    """)
    cur.execute("""
        ALTER TABLE shop_orders
            ALTER COLUMN order_number SET NOT NULL,
            ALTER COLUMN order_date SET NOT NULL,
            ALTER COLUMN customer_name SET NOT NULL,
            ALTER COLUMN delivery_address SET NOT NULL,
            ALTER COLUMN shipping_cost SET NOT NULL,
            ALTER COLUMN payment_status SET NOT NULL,
            ALTER COLUMN status SET NOT NULL,
            ALTER COLUMN stages SET NOT NULL,
            ALTER COLUMN document_confirmed SET NOT NULL,
            ALTER COLUMN created_at SET NOT NULL,
            ALTER COLUMN updated_at SET NOT NULL
    """)
    cur.execute("""
        ALTER TABLE shop_order_items
            ADD COLUMN IF NOT EXISTS product_name TEXT,
            ADD COLUMN IF NOT EXISTS price_netto REAL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS price_brutto REAL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS vat REAL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS warehouse TEXT,
            ADD COLUMN IF NOT EXISTS reserved_qty REAL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS issued_qty REAL DEFAULT 0
    """)
    cur.execute("""
        UPDATE shop_order_items i
        SET product_name=COALESCE(NULLIF(i.product_name, ''), p.name, 'Nieznany produkt'),
            warehouse=COALESCE(NULLIF(i.warehouse, ''), p.warehouse, 'Inne'),
            price_netto=COALESCE(i.price_netto, p.price_netto, 0),
            price_brutto=COALESCE(i.price_brutto,
                COALESCE(p.price_netto, 0) * (1 + COALESCE(p.vat, 0) / 100), 0),
            vat=COALESCE(i.vat, p.vat, 0),
            reserved_qty=COALESCE(i.reserved_qty, 0),
            issued_qty=COALESCE(i.issued_qty, 0)
        FROM products p
        WHERE p.id=i.product_id
    """)
    cur.execute("""
        ALTER TABLE shop_order_items
            ALTER COLUMN product_name SET NOT NULL,
            ALTER COLUMN price_netto SET NOT NULL,
            ALTER COLUMN price_brutto SET NOT NULL,
            ALTER COLUMN vat SET NOT NULL,
            ALTER COLUMN warehouse SET NOT NULL,
            ALTER COLUMN reserved_qty SET NOT NULL,
            ALTER COLUMN issued_qty SET NOT NULL
    """)
    cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS uq_shop_orders_order_number ON shop_orders(order_number)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_orders_stage ON shop_orders(stage)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_orders_salesperson ON shop_orders(lower(salesperson_email))")
    cur.execute("""
        UPDATE shop_orders o
        SET salesperson_email=u.email,
            salesperson_name=trim(concat_ws(' ',u.first_name,u.last_name)),
            salesperson_assigned_by=COALESCE(o.salesperson_assigned_by,'migracja'),
            salesperson_assigned_at=COALESCE(o.salesperson_assigned_at,NOW())
        FROM users u
        WHERE o.salesperson_email IS NULL
          AND u.role='sales'
          AND lower(u.email)=lower(COALESCE(o.created_by,''))
    """)
    cur.execute("""
        ALTER TABLE users
            ADD COLUMN IF NOT EXISTS phone TEXT,
            ADD COLUMN IF NOT EXISTS must_change_password BOOLEAN NOT NULL DEFAULT FALSE,
            ADD COLUMN IF NOT EXISTS last_login_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS password_changed_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS created_by TEXT,
            ADD COLUMN IF NOT EXISTS password_reset_by TEXT,
            ADD COLUMN IF NOT EXISTS failed_login_attempts INTEGER NOT NULL DEFAULT 0,
            ADD COLUMN IF NOT EXISTS last_failed_login_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS locked_until TIMESTAMPTZ
    """)
    cur.execute("ALTER TABLE users DROP CONSTRAINT IF EXISTS users_status_valid")
    cur.execute("""
        ALTER TABLE users
        ADD CONSTRAINT users_status_valid
        CHECK (status IN ('active','blocked','inactive'))
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS user_security_events(
        id SERIAL PRIMARY KEY,
        user_id INTEGER REFERENCES users(id) ON DELETE SET NULL,
        user_email TEXT NOT NULL,
        event_type TEXT NOT NULL,
        actor_email TEXT,
        details JSONB NOT NULL DEFAULT '{}'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_user_security_events_user ON user_security_events(lower(user_email),created_at DESC)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_order_items_order ON shop_order_items(order_id)")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_order_history(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        user_email TEXT NOT NULL,
        action TEXT NOT NULL,
        details TEXT,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_notifications(
        id SERIAL PRIMARY KEY,
        order_id INTEGER REFERENCES shop_orders(id) ON DELETE CASCADE,
        type TEXT NOT NULL,
        message TEXT NOT NULL,
        recipient_email TEXT,
        resolved BOOLEAN NOT NULL DEFAULT FALSE,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_sales_documents(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL UNIQUE REFERENCES shop_orders(id) ON DELETE CASCADE,
        document_number TEXT NOT NULL,
        editable_data JSONB NOT NULL,
        docx BYTEA,
        pdf BYTEA,
        confirmed BOOLEAN NOT NULL DEFAULT FALSE,
        created_by TEXT,
        confirmed_by TEXT,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        confirmed_at TIMESTAMPTZ
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_accounting(
        order_id INTEGER PRIMARY KEY REFERENCES shop_orders(id) ON DELETE CASCADE,
        proforma_issued BOOLEAN NOT NULL DEFAULT FALSE,
        reserved_by_proforma BOOLEAN NOT NULL DEFAULT FALSE,
        waiting_for_payment BOOLEAN NOT NULL DEFAULT FALSE,
        partial_payment BOOLEAN NOT NULL DEFAULT FALSE,
        paid BOOLEAN NOT NULL DEFAULT FALSE,
        payment_method TEXT,
        invoice_issued BOOLEAN NOT NULL DEFAULT FALSE,
        receipt_issued BOOLEAN NOT NULL DEFAULT FALSE,
        invoice_sent BOOLEAN NOT NULL DEFAULT FALSE,
        document_to_warehouse BOOLEAN NOT NULL DEFAULT FALSE,
        ready_to_ship BOOLEAN NOT NULL DEFAULT FALSE,
        settled BOOLEAN NOT NULL DEFAULT FALSE,
        proforma_number TEXT,
        invoice_number TEXT,
        receipt_number TEXT,
        document_issue_date DATE,
        payment_received_date DATE,
        amount_paid REAL NOT NULL DEFAULT 0,
        amount_due REAL NOT NULL DEFAULT 0,
        accounting_notes TEXT,
        salesperson TEXT,
        updated_by TEXT,
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        sales_document_generated BOOLEAN NOT NULL DEFAULT FALSE,
        document_sent BOOLEAN NOT NULL DEFAULT FALSE,
        payment_due_date DATE
    );
    """)
    cur.execute("""
        ALTER TABLE shop_accounting
            ADD COLUMN IF NOT EXISTS sales_document_generated BOOLEAN NOT NULL DEFAULT FALSE,
            ADD COLUMN IF NOT EXISTS document_sent BOOLEAN NOT NULL DEFAULT FALSE,
            ADD COLUMN IF NOT EXISTS payment_due_date DATE
    """)
    cur.execute("""
        ALTER TABLE shop_sales_documents
            ADD COLUMN IF NOT EXISTS voided_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS voided_by TEXT,
            ADD COLUMN IF NOT EXISTS void_reason TEXT,
            ADD COLUMN IF NOT EXISTS restored_at TIMESTAMPTZ,
            ADD COLUMN IF NOT EXISTS restored_by TEXT
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS accounting_documents(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        document_type TEXT NOT NULL,
        document_number TEXT,
        status TEXT NOT NULL DEFAULT 'draft',
        data JSONB NOT NULL DEFAULT '{}'::jsonb,
        void_reason TEXT,
        voided_by TEXT,
        voided_at TIMESTAMPTZ,
        restored_by TEXT,
        restored_at TIMESTAMPTZ,
        created_by TEXT,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        updated_by TEXT,
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        CONSTRAINT accounting_document_type_valid
            CHECK (document_type IN ('proforma','invoice','receipt','sales')),
        CONSTRAINT accounting_document_status_valid
            CHECK (status IN ('draft','active','cancelled')),
        UNIQUE(order_id,document_type)
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS accounting_history(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        entity_type TEXT NOT NULL,
        entity_id INTEGER,
        field_name TEXT NOT NULL,
        previous_value TEXT,
        new_value TEXT,
        changed_by TEXT NOT NULL,
        changed_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_accounting_documents_order ON accounting_documents(order_id,status,document_type)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_accounting_history_order ON accounting_history(order_id,changed_at DESC)")
    cur.execute("""
        UPDATE shop_accounting a
        SET sales_document_generated=EXISTS(
                SELECT 1 FROM shop_sales_documents d
                WHERE d.order_id=a.order_id AND d.voided_at IS NULL
            ),
            document_sent=COALESCE((
                SELECT (o.stages->>'document_sent')::boolean
                FROM shop_orders o WHERE o.id=a.order_id
            ),FALSE)
    """)
    cur.execute("""
        INSERT INTO accounting_documents(
            order_id,document_type,document_number,status,created_by
        )
        SELECT a.order_id,document_type,document_number,
               CASE WHEN issued THEN 'active' ELSE 'draft' END,'migracja'
        FROM shop_accounting a
        CROSS JOIN LATERAL (VALUES
            ('proforma',a.proforma_number,a.proforma_issued),
            ('invoice',a.invoice_number,a.invoice_issued),
            ('receipt',a.receipt_number,a.receipt_issued)
        ) AS d(document_type,document_number,issued)
        WHERE COALESCE(document_number,'')<>'' OR issued
        ON CONFLICT (order_id,document_type) DO NOTHING
    """)
    cur.execute("""
        INSERT INTO accounting_documents(
            order_id,document_type,document_number,status,created_by,created_at
        )
        SELECT d.order_id,'sales',d.document_number,
               CASE WHEN d.voided_at IS NULL THEN 'active' ELSE 'cancelled' END,
               d.created_by,d.created_at
        FROM shop_sales_documents d
        ON CONFLICT (order_id,document_type) DO NOTHING
    """)
    cur.execute("""
        UPDATE shop_orders o
        SET salesperson_email=u.email,
            salesperson_name=trim(concat_ws(' ',u.first_name,u.last_name)),
            salesperson_assigned_by=COALESCE(o.salesperson_assigned_by,'migracja'),
            salesperson_assigned_at=COALESCE(o.salesperson_assigned_at,NOW())
        FROM shop_accounting a,users u
        WHERE a.order_id=o.id
          AND o.salesperson_email IS NULL
          AND u.role='sales'
          AND (
              lower(u.email)=lower(COALESCE(a.salesperson,''))
              OR lower(trim(concat_ws(' ',u.first_name,u.last_name)))=
                 lower(COALESCE(a.salesperson,''))
          )
    """)
    cur.execute("ALTER TABLE products ADD COLUMN IF NOT EXISTS created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()")
    cur.execute("ALTER TABLE shop_notifications ADD COLUMN IF NOT EXISTS recipient_email TEXT")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_notifications_recipient ON shop_notifications(lower(recipient_email),resolved,created_at DESC)")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS reservations(
        id SERIAL PRIMARY KEY,
        reservation_number TEXT UNIQUE NOT NULL,
        reservation_date DATE NOT NULL DEFAULT CURRENT_DATE,
        customer_name TEXT NOT NULL,
        salesperson_email TEXT,
        salesperson_name TEXT,
        warehouse TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'robocza',
        notes TEXT,
        cancel_reason TEXT,
        cancelled_by TEXT,
        cancelled_at TIMESTAMPTZ,
        created_by TEXT NOT NULL,
        approved_by TEXT,
        approved_at TIMESTAMPTZ,
        pdf BYTEA,
        pdf_generated_by TEXT,
        pdf_generated_at TIMESTAMPTZ,
        completed_by TEXT,
        completed_at TIMESTAMPTZ,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS reservation_items(
        id SERIAL PRIMARY KEY,
        reservation_id INTEGER NOT NULL REFERENCES reservations(id) ON DELETE CASCADE,
        product_id INTEGER NOT NULL REFERENCES products(id) ON DELETE RESTRICT,
        package_id INTEGER REFERENCES packages(id) ON DELETE SET NULL,
        product_name TEXT NOT NULL,
        package_number TEXT,
        dimension TEXT,
        qty REAL NOT NULL CHECK (qty > 0),
        unit TEXT NOT NULL,
        warehouse TEXT NOT NULL,
        location TEXT,
        price_netto REAL NOT NULL DEFAULT 0,
        price_brutto REAL NOT NULL DEFAULT 0,
        prepared BOOLEAN NOT NULL DEFAULT FALSE,
        prepared_by TEXT,
        prepared_at TIMESTAMPTZ,
        warehouse_note TEXT
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS reservation_history(
        id SERIAL PRIMARY KEY,
        reservation_id INTEGER NOT NULL REFERENCES reservations(id) ON DELETE CASCADE,
        actor_email TEXT NOT NULL,
        action TEXT NOT NULL,
        details TEXT,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_reservations_status ON reservations(status, reservation_date DESC)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_reservations_customer ON reservations(lower(customer_name))")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_reservation_items_product ON reservation_items(product_id, package_id, warehouse)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_reservation_history_reservation ON reservation_history(reservation_id, created_at DESC)")
    cur.execute("""
        DO $$
        BEGIN
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname='reservations_status_valid') THEN
                ALTER TABLE reservations
                ADD CONSTRAINT reservations_status_valid
                CHECK (status IN (
                    'robocza','zatwierdzona','przekazana do magazynu',
                    'w trakcie kompletowania','skompletowana','wydana','anulowana'
                )) NOT VALID;
            END IF;
        END $$;
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_order_salesperson_history(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        previous_salesperson TEXT,
        new_salesperson TEXT,
        changed_by TEXT NOT NULL,
        changed_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_salesperson_history_order ON shop_order_salesperson_history(order_id,changed_at DESC)")
    cur.execute("""
        UPDATE shop_orders o
        SET stages=jsonb_build_object(
            'order_accepted', o.status<>'Nowe zamówienie',
            'stock_checked', FALSE,
            'stock_reserved', o.status IN (
                'Towar zarezerwowany','Dokument wystawiony','W trakcie pakowania',
                'Spakowane','Wysłane','Dostarczone','Zakończone'
            ),
            'payment_checked', o.payment_status IS NOT NULL,
            'paid', o.payment_status='Opłacone',
            'proforma_issued', COALESCE((
                SELECT a.proforma_issued FROM shop_accounting a WHERE a.order_id=o.id
            ),FALSE),
            'invoice_issued', COALESCE((
                SELECT a.invoice_issued FROM shop_accounting a WHERE a.order_id=o.id
            ),FALSE),
            'receipt_issued', COALESCE((
                SELECT a.receipt_issued FROM shop_accounting a WHERE a.order_id=o.id
            ),FALSE),
            'sales_document_generated', EXISTS (
                SELECT 1 FROM shop_sales_documents sd WHERE sd.order_id=o.id
            ),
            'document_sent', FALSE,
            'sent_to_packing', o.status IN (
                'W trakcie pakowania','Spakowane','Wysłane','Dostarczone','Zakończone'
            ),
            'packed', o.status IN ('Spakowane','Wysłane','Dostarczone','Zakończone'),
            'shipped', o.status IN ('Wysłane','Dostarczone','Zakończone'),
            'tracking_entered', COALESCE(NULLIF(o.tracking_number,''),'')<>'',
            'completed', o.status='Zakończone',
            'cancelled', o.status='Anulowane'
        )
        WHERE COALESCE(o.stages,'{}'::jsonb)='{}'::jsonb
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_accounting_filters ON shop_accounting(payment_method, proforma_number, invoice_number, receipt_number, salesperson)")

    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_orders_search ON shop_orders(lower(order_number), lower(customer_name), lower(status), order_date)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_shop_items_product ON shop_order_items(product_id)")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_doc_photos(
        id SERIAL PRIMARY KEY,
        doc_id INTEGER NOT NULL REFERENCES issue_docs(id) ON DELETE CASCADE,
        filename TEXT NOT NULL,
        content_type TEXT NOT NULL,
        data BYTEA NOT NULL,
        note TEXT,
        added_by TEXT NOT NULL,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_doc_history(
        id SERIAL PRIMARY KEY,
        doc_id INTEGER NOT NULL REFERENCES issue_docs(id) ON DELETE CASCADE,
        user_email TEXT NOT NULL,
        action TEXT NOT NULL,
        details TEXT,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_issue_doc_photos_doc ON issue_doc_photos(doc_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_issue_doc_history_doc ON issue_doc_history(doc_id)")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS system_settings(
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL,
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS shop_order_stage_history(
        id SERIAL PRIMARY KEY,
        order_id INTEGER NOT NULL REFERENCES shop_orders(id) ON DELETE CASCADE,
        stage_key TEXT NOT NULL,
        previous_value BOOLEAN NOT NULL,
        new_value BOOLEAN NOT NULL,
        previous_status TEXT,
        new_status TEXT,
        changed_by TEXT NOT NULL,
        changed_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
        ALTER TABLE shop_order_stage_history
            ADD COLUMN IF NOT EXISTS previous_status TEXT,
            ADD COLUMN IF NOT EXISTS new_status TEXT
    """)
    cur.execute(
        "CREATE INDEX IF NOT EXISTS idx_shop_stage_history_order "
        "ON shop_order_stage_history(order_id, changed_at DESC)"
    )
    cur.execute("""
        ALTER TABLE shop_sales_documents
            ADD COLUMN IF NOT EXISTS document_number TEXT,
            ADD COLUMN IF NOT EXISTS editable_data JSONB DEFAULT '{}'::jsonb,
            ADD COLUMN IF NOT EXISTS docx BYTEA,
            ADD COLUMN IF NOT EXISTS pdf BYTEA,
            ADD COLUMN IF NOT EXISTS confirmed BOOLEAN DEFAULT FALSE,
            ADD COLUMN IF NOT EXISTS created_by TEXT,
            ADD COLUMN IF NOT EXISTS confirmed_by TEXT,
            ADD COLUMN IF NOT EXISTS created_at TIMESTAMPTZ DEFAULT NOW(),
            ADD COLUMN IF NOT EXISTS confirmed_at TIMESTAMPTZ
    """)
    cur.execute(
        """
        UPDATE shop_sales_documents
        SET document_number=COALESCE(NULLIF(document_number,''), 'DS/' || order_id),
            editable_data=COALESCE(editable_data, '{}'::jsonb),
            confirmed=COALESCE(confirmed, FALSE),
            created_at=COALESCE(created_at, NOW())
        """
    )
    cur.execute(
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_shop_sales_documents_order ON shop_sales_documents(order_id)"
    )
    cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS can_import_warehouse BOOLEAN NOT NULL DEFAULT FALSE")
    cur.execute("ALTER TABLE users ADD COLUMN IF NOT EXISTS can_import_accounting BOOLEAN NOT NULL DEFAULT FALSE")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS contractors(
        id SERIAL PRIMARY KEY,
        name TEXT NOT NULL,
        nip TEXT,
        email TEXT,
        phone TEXT,
        address TEXT,
        active BOOLEAN NOT NULL DEFAULT TRUE,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_contractors_name ON contractors(lower(name))")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_contractors_nip ON contractors(nip)")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS general_imports(
        id SERIAL PRIMARY KEY,
        filename TEXT NOT NULL,
        imported_by TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'draft',
        detected_sheets JSONB NOT NULL DEFAULT '[]'::jsonb,
        summary JSONB NOT NULL DEFAULT '{}'::jsonb,
        errors JSONB NOT NULL DEFAULT '[]'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        completed_at TIMESTAMPTZ
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS general_import_rows(
        id SERIAL PRIMARY KEY,
        import_id INTEGER NOT NULL REFERENCES general_imports(id) ON DELETE CASCADE,
        sheet_name TEXT NOT NULL,
        row_number INTEGER NOT NULL,
        entity_type TEXT NOT NULL,
        source_data JSONB NOT NULL DEFAULT '{}'::jsonb,
        normalized_data JSONB NOT NULL DEFAULT '{}'::jsonb,
        duplicate_data JSONB,
        resolution TEXT NOT NULL DEFAULT 'new',
        included BOOLEAN NOT NULL DEFAULT TRUE,
        validation_errors JSONB NOT NULL DEFAULT '[]'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_general_import_rows_import ON general_import_rows(import_id, id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_general_imports_created ON general_imports(created_at DESC)")
    cur.execute("""
    CREATE TABLE IF NOT EXISTS app_settings(
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL,
        updated_by TEXT,
        updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute(
        """
        INSERT INTO app_settings(key,value)
        VALUES ('issue_import_allow_general_stock','false')
        ON CONFLICT (key) DO NOTHING
        """
    )
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_imports(
        id SERIAL PRIMARY KEY,
        filename TEXT NOT NULL,
        imported_by TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'draft',
        detected_sheets JSONB NOT NULL DEFAULT '[]'::jsonb,
        summary JSONB NOT NULL DEFAULT '{}'::jsonb,
        errors JSONB NOT NULL DEFAULT '[]'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
        completed_at TIMESTAMPTZ,
        undone_at TIMESTAMPTZ,
        undone_by TEXT
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_import_rows(
        id SERIAL PRIMARY KEY,
        import_id INTEGER NOT NULL REFERENCES issue_imports(id) ON DELETE CASCADE,
        sheet_name TEXT NOT NULL,
        row_number INTEGER NOT NULL,
        source_data JSONB NOT NULL DEFAULT '{}'::jsonb,
        normalized_data JSONB NOT NULL DEFAULT '{}'::jsonb,
        duplicate_data JSONB,
        resolution TEXT NOT NULL DEFAULT 'new',
        included BOOLEAN NOT NULL DEFAULT TRUE,
        validation_errors JSONB NOT NULL DEFAULT '[]'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS issue_import_effects(
        id SERIAL PRIMARY KEY,
        import_id INTEGER NOT NULL REFERENCES issue_imports(id) ON DELETE CASCADE,
        row_id INTEGER NOT NULL REFERENCES issue_import_rows(id) ON DELETE CASCADE,
        action TEXT NOT NULL,
        doc_id INTEGER REFERENCES issue_docs(id) ON DELETE RESTRICT,
        item_id INTEGER REFERENCES issue_items(id) ON DELETE SET NULL,
        product_id INTEGER REFERENCES products(id) ON DELETE RESTRICT,
        package_id INTEGER REFERENCES packages(id) ON DELETE SET NULL,
        qty REAL NOT NULL,
        prior_data JSONB NOT NULL DEFAULT '{}'::jsonb,
        created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
    );
    """)
    cur.execute("""
        ALTER TABLE issue_import_effects
            DROP CONSTRAINT IF EXISTS issue_import_effects_item_id_fkey
    """)
    cur.execute("""
        ALTER TABLE issue_import_effects
            ADD CONSTRAINT issue_import_effects_item_id_fkey
            FOREIGN KEY (item_id) REFERENCES issue_items(id) ON DELETE SET NULL
    """)
    cur.execute(
        "CREATE INDEX IF NOT EXISTS idx_issue_import_rows_import "
        "ON issue_import_rows(import_id,id)"
    )
    cur.execute(
        "CREATE INDEX IF NOT EXISTS idx_issue_imports_created "
        "ON issue_imports(created_at DESC)"
    )
    cur.execute(
        "CREATE INDEX IF NOT EXISTS idx_issue_import_effects_import "
        "ON issue_import_effects(import_id,id)"
    )

    conn.commit()
    conn.close()


def ensure_db_initialized():
    global DB_INITIALIZED, DB_POOL
    if DB_INITIALIZED:
        return
    with DB_INIT_LOCK:
        if DB_INITIALIZED:
            return
        try:
            run_db_migrations()
            DB_INITIALIZED = True
        except Exception:
            logger.exception("Database initialization failed.")
            if DB_POOL:
                DB_POOL.closeall()
                DB_POOL = None
            raise


def csrf_token():
    token = session.get("_csrf_token")
    if not token:
        token = secrets.token_urlsafe(32)
        session["_csrf_token"] = token
    return token


app.jinja_env.globals["csrf_token"] = csrf_token


def parse_positive_number(value, field_name="Ilość"):
    try:
        number = float(str(value or "").strip().replace(",", "."))
    except (TypeError, ValueError):
        raise ValueError(f"{field_name}: podaj prawidłową liczbę.")
    if not math.isfinite(number) or number <= 0:
        raise ValueError(f"{field_name} musi być większa od zera.")
    return number


def parse_nonnegative_number(value, field_name):
    if value in (None, ""):
        return 0.0
    try:
        number = float(str(value).strip().replace(",", "."))
    except (TypeError, ValueError):
        raise ValueError(f"{field_name}: podaj prawidłową liczbę.")
    if not math.isfinite(number) or number < 0:
        raise ValueError(f"{field_name} nie może być ujemna.")
    return number


def close_with_rollback(conn):
    conn.rollback()

ALLOWED_ISSUE_PHOTO_TYPES = {"image/jpeg", "image/png", "image/webp", "image/gif"}
MAX_ISSUE_PHOTO_BYTES = 5 * 1024 * 1024


def is_issue_document(movement_type, doc_number):
    movement = str(movement_type or "").upper()
    number = str(doc_number or "").upper()
    return movement in {"WZ", "RW"} or number.startswith("WZ") or number.startswith("RW")


def issue_doc_history(cur, doc_id, action, details=""):
    actor = session.get("user", "system") if has_request_context() else "system"
    cur.execute(
        """
        INSERT INTO issue_doc_history(doc_id, user_email, action, details)
        VALUES (%s,%s,%s,%s)
        """,
        (doc_id, actor, action, details),
    )


def clean_photo_filename(filename):
    filename = (filename or "zdjecie").rsplit("/", 1)[-1].rsplit("\\", 1)[-1].strip()
    filename = "".join(char for char in filename if ord(char) >= 32 and char not in {'"', "'"})
    return filename[:180] or "zdjecie"


def save_issue_photos(cur, doc_id, uploads, note=""):
    saved = 0
    for upload in uploads:
        if not upload or not upload.filename:
            continue
        content_type = (upload.mimetype or "").lower()
        if content_type not in ALLOWED_ISSUE_PHOTO_TYPES:
            raise ValueError("Dozwolone są tylko zdjęcia JPG, PNG, WEBP lub GIF.")
        data = upload.read()
        if not data:
            continue
        if len(data) > MAX_ISSUE_PHOTO_BYTES:
            raise ValueError("Jedno zdjęcie może mieć maksymalnie 5 MB.")
        cur.execute(
            """
            INSERT INTO issue_doc_photos(doc_id, filename, content_type, data, note, added_by)
            VALUES (%s,%s,%s,%s,%s,%s)
            """,
            (
                doc_id,
                clean_photo_filename(upload.filename),
                content_type,
                psycopg2.Binary(data),
                (note or "").strip()[:500],
                session.get("user", "system"),
            ),
        )
        saved += 1
    if saved:
        issue_doc_history(cur, doc_id, "dodano zdjęcia", f"Liczba zdjęć: {saved}")
    return saved

SHOP_STATUS_FLOW = [
    "Nowe zamówienie", "Przyjęte", "Oczekuje na płatność", "Opłacone",
    "Towar zarezerwowany", "Dokument wystawiony", "W trakcie pakowania",
    "Spakowane", "Wysłane", "Dostarczone", "Zakończone", "Anulowane",
]
SHOP_ROLE_LABELS = {
    "admin": "Administrator", "employee": "Pracownik", "warehouse": "Magazynier",
    "shop": "Obsługa sklepu internetowego", "accounting": "Księgowość",
    "sales": "Handlowiec",
}

ACCOUNTING_PAYMENT_METHODS = ["Przelew", "Gotówka", "Karta", "BLIK", "Autopay", "Pobranie", "Inny"]
ACCOUNTING_BOOL_FIELDS = [
    "proforma_issued", "reserved_by_proforma", "waiting_for_payment", "partial_payment",
    "paid", "invoice_issued", "receipt_issued", "invoice_sent", "document_to_warehouse",
    "sales_document_generated", "document_sent", "ready_to_ship", "settled",
]
ACCOUNTING_FIELD_LABELS = {
    "proforma_issued": "Proforma wystawiona",
    "reserved_by_proforma": "Towar zarezerwowany na podstawie proformy",
    "waiting_for_payment": "Oczekiwanie na płatność",
    "partial_payment": "Płatność częściowa",
    "paid": "Zapłacono",
    "invoice_issued": "Faktura wystawiona",
    "receipt_issued": "Paragon wystawiony",
    "sales_document_generated": "Dokument sprzedaży wygenerowany",
    "document_sent": "Dokument wysłany do klienta",
    "invoice_sent": "Faktura wysłana do klienta",
    "document_to_warehouse": "Dokument przekazany do magazynu",
    "ready_to_ship": "Zamówienie gotowe do wysyłki",
    "settled": "Zamówienie rozliczone",
}

SHOP_ORDER_STAGES = (
    ("order_accepted", "Zamówienie przyjęte", "sales"),
    ("stock_checked", "Towar sprawdzony w magazynie", "warehouse"),
    ("stock_reserved", "Towar zarezerwowany", "warehouse"),
    ("payment_checked", "Płatność sprawdzona", "accounting"),
    ("paid", "Zapłacone", "accounting"),
    ("proforma_issued", "Proforma wystawiona", "accounting"),
    ("invoice_issued", "Faktura wystawiona", "accounting"),
    ("receipt_issued", "Paragon wystawiony", "accounting"),
    ("sales_document_generated", "Dokument sprzedaży wygenerowany", "accounting"),
    ("document_sent", "Dokument wysłany do klienta", "sales"),
    ("sent_to_packing", "Przekazane do pakowania", "warehouse"),
    ("packed", "Spakowane", "warehouse"),
    ("shipped", "Wysłane", "warehouse"),
    ("tracking_entered", "Numer przesyłki wpisany", "warehouse"),
    ("completed", "Zakończone", "shop_edit"),
    ("cancelled", "Anulowane", "shop_edit"),
)
SHOP_STAGE_BY_KEY = {
    key: {"label": label, "permission": permission}
    for key, label, permission in SHOP_ORDER_STAGES
}
SHOP_STAGE_STATUS = {
    "order_accepted": "Przyjęte",
    "stock_reserved": "Towar zarezerwowany",
    "payment_checked": "Oczekuje na płatność",
    "paid": "Opłacone",
    "sales_document_generated": "Dokument wystawiony",
    "sent_to_packing": "W trakcie pakowania",
    "packed": "Spakowane",
    "shipped": "Wysłane",
    "completed": "Zakończone",
    "cancelled": "Anulowane",
}


def accounting_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return redirect('/login')
        if session.get('role') not in {'admin', 'accounting'}:
            return "Brak uprawnień", 403
        return f(*args, **kwargs)
    return decorated


def order_can_be_shipped(accounting_row):
    if not accounting_row:
        return False
    payment_method = accounting_row[6]
    paid = bool(accounting_row[5])
    ready_to_ship = bool(accounting_row[11])
    return ready_to_ship and (paid or payment_method == "Pobranie")


def accounting_history_change(
    cur, order_id, entity_type, entity_id, field_name, previous, new
):
    if previous == new:
        return
    cur.execute(
        """
        INSERT INTO accounting_history(
            order_id,entity_type,entity_id,field_name,previous_value,new_value,changed_by
        ) VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (
            order_id,
            entity_type,
            entity_id,
            field_name,
            "" if previous is None else str(previous),
            "" if new is None else str(new),
            session.get("user", "system") if has_request_context() else "system",
        ),
    )


def sync_accounting_document(
    cur, order_id, document_type, document_number, issued, actor, data=None
):
    cur.execute(
        """
        INSERT INTO accounting_documents(
            order_id,document_type,document_number,status,data,created_by,updated_by
        ) VALUES (%s,%s,%s,%s,%s::jsonb,%s,%s)
        ON CONFLICT (order_id,document_type) DO UPDATE SET
            document_number=EXCLUDED.document_number,
            status=CASE
                WHEN accounting_documents.status='cancelled' THEN 'cancelled'
                ELSE EXCLUDED.status
            END,
            data=EXCLUDED.data,
            updated_by=EXCLUDED.updated_by,
            updated_at=NOW()
        RETURNING id,status
        """,
        (
            order_id,
            document_type,
            (document_number or "").strip() or None,
            "active" if issued else "draft",
            json.dumps(data or {}, ensure_ascii=False, default=str),
            actor,
            actor,
        ),
    )
    return cur.fetchone()


def accounting_document_warnings(cur, order_id, document_type):
    warnings = ["zamówieniem"]
    cur.execute(
        "SELECT 1 FROM shop_order_items WHERE order_id=%s LIMIT 1",
        (order_id,),
    )
    if cur.fetchone():
        warnings.append("magazynem i produktami")
    cur.execute(
        """
        SELECT paid,partial_payment,amount_paid,invoice_issued,receipt_issued
        FROM shop_accounting WHERE order_id=%s
        """,
        (order_id,),
    )
    accounting = cur.fetchone()
    if accounting and (accounting[0] or accounting[1] or float(accounting[2] or 0) > 0):
        warnings.append("płatnością")
    if accounting and accounting[3] and document_type != "invoice":
        warnings.append("fakturą")
    if accounting and accounting[4] and document_type != "receipt":
        warnings.append("paragonem")
    return warnings


def ensure_shop_accounting_row(cur, order_id):
    cur.execute(
        """
        INSERT INTO shop_accounting(order_id, amount_due)
        SELECT o.id, COALESCE(SUM(i.qty*i.price_brutto),0)+o.shipping_cost
        FROM shop_orders o
        LEFT JOIN shop_order_items i ON i.order_id=o.id
        WHERE o.id=%s
        GROUP BY o.id
        ON CONFLICT (order_id) DO NOTHING
        """,
        (order_id,),
    )


def sync_accounting_payment_status(cur, order_id):
    cur.execute("SELECT paid, partial_payment, waiting_for_payment, payment_method FROM shop_accounting WHERE order_id=%s", (order_id,))
    row = cur.fetchone()
    if not row:
        return
    if row[0]:
        status = "Opłacone"
    elif row[1]:
        status = "Płatność częściowa"
    elif row[3] == "Pobranie":
        status = "Pobranie"
    elif row[2]:
        status = "Oczekuje na płatność"
    else:
        status = "Oczekuje na płatność"
    cur.execute("UPDATE shop_orders SET payment_status=%s, payment_method=COALESCE(NULLIF((SELECT payment_method FROM shop_accounting WHERE order_id=%s), ''), payment_method), updated_at=NOW() WHERE id=%s", (status, order_id, order_id))


def current_user_role():
    return session.get("role", "employee") if has_request_context() else "employee"


def can_shop(action):
    role = current_user_role()
    if role == "admin":
        return True
    return role in {
        "shop_edit": {"shop"},
        "sales": {"shop", "sales"},
        "create": {"shop", "sales"},
        "warehouse": {"warehouse"},
        "accounting": {"accounting"},
        "view": {"shop", "warehouse", "accounting", "sales"},
    }.get(action, set())


def require_shop_permission(action):
    if not can_shop(action):
        return "Brak uprawnień do tej funkcji modułu sklepu internetowego.", 403
    return None


RESERVATION_STATUSES = (
    "robocza",
    "zatwierdzona",
    "przekazana do magazynu",
    "w trakcie kompletowania",
    "skompletowana",
    "wydana",
    "anulowana",
)
ACTIVE_RESERVATION_STATUSES = (
    "zatwierdzona",
    "przekazana do magazynu",
    "w trakcie kompletowania",
    "skompletowana",
)


def reservation_user_can_edit():
    return current_user_role() in {"admin", "shop", "sales"}


def reservation_user_can_pick():
    return current_user_role() in {"admin", "warehouse"}


def reservation_history(cur, reservation_id, action, details=""):
    actor = session.get("user", "system") if has_request_context() else "system"
    cur.execute(
        """
        INSERT INTO reservation_history(reservation_id, actor_email, action, details)
        VALUES (%s,%s,%s,%s)
        """,
        (reservation_id, actor, action, details),
    )


def active_reservation_statuses_sql():
    return list(ACTIVE_RESERVATION_STATUSES)


def reserved_product_qty(cur, product_id, exclude_reservation_id=None):
    cur.execute(
        """
        SELECT COALESCE(SUM(i.reserved_qty-i.issued_qty), 0)
        FROM shop_order_items i
        JOIN shop_orders o ON o.id=i.order_id
        WHERE i.product_id=%s
          AND o.status NOT IN ('Anulowane', 'Zakończone')
        """,
        (product_id,),
    )
    reserved = float((cur.fetchone() or (0,))[0] or 0)
    params = [product_id, active_reservation_statuses_sql()]
    exclude_sql = ""
    if exclude_reservation_id:
        exclude_sql = "AND r.id<>%s"
        params.append(exclude_reservation_id)
    cur.execute(
        f"""
        SELECT COALESCE(SUM(ri.qty), 0)
        FROM reservation_items ri
        JOIN reservations r ON r.id=ri.reservation_id
        WHERE ri.product_id=%s
          AND r.status = ANY(%s)
          {exclude_sql}
        """,
        tuple(params),
    )
    reserved += float((cur.fetchone() or (0,))[0] or 0)
    return max(reserved, 0.0)


def reserved_package_qty(cur, package_id, exclude_reservation_id=None):
    params = [package_id, active_reservation_statuses_sql()]
    exclude_sql = ""
    if exclude_reservation_id:
        exclude_sql = "AND r.id<>%s"
        params.append(exclude_reservation_id)
    cur.execute(
        f"""
        SELECT COALESCE(SUM(ri.qty), 0)
        FROM reservation_items ri
        JOIN reservations r ON r.id=ri.reservation_id
        WHERE ri.package_id=%s
          AND r.status = ANY(%s)
          {exclude_sql}
        """,
        tuple(params),
    )
    return max(float((cur.fetchone() or (0,))[0] or 0), 0.0)


def product_available_qty(cur, product_id, product_qty, exclude_reservation_id=None):
    return max(float(product_qty or 0) - reserved_product_qty(cur, product_id, exclude_reservation_id), 0.0)


def package_available_qty(cur, package_id, package_qty, exclude_reservation_id=None):
    return max(float(package_qty or 0) - reserved_package_qty(cur, package_id, exclude_reservation_id), 0.0)


def notify_warehouse_reservation(cur, reservation_id, message):
    cur.execute(
        """
        SELECT email FROM users
        WHERE role='warehouse' AND status='active'
        ORDER BY lower(email)
        """
    )
    for (email,) in cur.fetchall():
        cur.execute(
            """
            INSERT INTO shop_notifications(type,message,recipient_email)
            VALUES ('rezerwacja',%s,%s)
            """,
            (message, (email or "").strip().lower()),
        )
    reservation_history(cur, reservation_id, "powiadomiono magazyn", message)


def shop_history(cur, order_id, action, details=""):
    actor = session.get("user", "system") if has_request_context() else "system"
    cur.execute(
        """
        INSERT INTO shop_order_history(order_id, user_email, action, details)
        VALUES (%s,%s,%s,%s)
        """,
        (order_id, actor, action, details),
    )


def sales_can_access_order(cur, order_id):
    if current_user_role() != "sales":
        return True
    cur.execute(
        """
        SELECT 1 FROM shop_orders
        WHERE id=%s AND lower(COALESCE(salesperson_email,''))=lower(%s)
        """,
        (order_id, session.get("user", "")),
    )
    return cur.fetchone() is not None


def sales_order_access_error(cur, order_id):
    if sales_can_access_order(cur, order_id):
        return None
    return "Handlowiec ma dostęp wyłącznie do własnych zamówień.", 403


def notify_order_salesperson(cur, order_id, notification_type, message):
    cur.execute(
        "SELECT salesperson_email FROM shop_orders WHERE id=%s",
        (order_id,),
    )
    row = cur.fetchone()
    recipient = (row[0] or "").strip().lower() if row else ""
    if recipient:
        cur.execute(
            """
            INSERT INTO shop_notifications(order_id,type,message,recipient_email)
            VALUES (%s,%s,%s,%s)
            """,
            (order_id, notification_type, message, recipient),
        )


def reservation_pdf_bytes(reservation, items):
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    normal = styles["Normal"]
    small = ParagraphStyle("reservation-small", parent=normal, fontSize=8, leading=10)
    story = [
        Paragraph("Rezerwacja - checklista magazyniera", title_style),
        Spacer(1, 8),
        Paragraph(f"Numer rezerwacji: {escape(str(reservation[1]))}", normal),
        Paragraph(f"Data: {escape(str(reservation[2]))}", normal),
        Paragraph(f"Klient: {escape(str(reservation[3]))}", normal),
        Paragraph(f"Handlowiec: {escape(str(reservation[4] or reservation[5] or ''))}", normal),
        Paragraph(f"Magazyn: {escape(str(reservation[6]))}", normal),
        Paragraph(f"Uwagi: {escape(str(reservation[8] or ''))}", normal),
        Spacer(1, 10),
    ]
    data = [[
        "", "Paczka", "Produkt", "Wymiar", "Ilość", "Jm", "Lokalizacja", "Uwagi",
    ]]
    for item in items:
        data.append([
            "[ ]",
            Paragraph(escape(str(item[5] or "")), small),
            Paragraph(escape(str(item[4] or "")), small),
            Paragraph(escape(str(item[6] or "")), small),
            f"{float(item[7] or 0):.3f}".rstrip("0").rstrip("."),
            Paragraph(escape(str(item[8] or "")), small),
            Paragraph(escape(str(item[10] or "")), small),
            Paragraph(escape(str(item[15] or "")), small),
        ])
    table = Table(
        data,
        repeatRows=1,
        colWidths=[10 * mm, 24 * mm, 43 * mm, 22 * mm, 18 * mm, 13 * mm, 26 * mm, 25 * mm],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFC067")),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 1), (0, -1), "CENTER"),
        ("ALIGN", (4, 1), (4, -1), "RIGHT"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([
        table,
        Spacer(1, 16),
        Paragraph("Przygotował: ________________________________", normal),
        Spacer(1, 8),
        Paragraph("Data: ____________________", normal),
        Spacer(1, 8),
        Paragraph("Podpis: ________________________________", normal),
        Spacer(1, 8),
        Paragraph("Uwagi magazyniera: ________________________________________________", normal),
    ])
    document.build(story)
    return output.getvalue()


def load_reservation_for_pdf(cur, reservation_id):
    cur.execute(
        """
        SELECT id,reservation_number,reservation_date,customer_name,
               salesperson_name,salesperson_email,warehouse,status,notes
        FROM reservations WHERE id=%s
        """,
        (reservation_id,),
    )
    reservation = cur.fetchone()
    cur.execute(
        """
        SELECT id,reservation_id,product_id,package_id,product_name,package_number,
               dimension,qty,unit,warehouse,location,price_netto,price_brutto,
               prepared,prepared_by,warehouse_note
        FROM reservation_items WHERE reservation_id=%s ORDER BY id
        """,
        (reservation_id,),
    )
    return reservation, cur.fetchall()


def generate_and_store_reservation_pdf(cur, reservation_id, actor):
    reservation, items = load_reservation_for_pdf(cur, reservation_id)
    if not reservation:
        raise ValueError("Nie znaleziono rezerwacji.")
    content = reservation_pdf_bytes(reservation, items)
    cur.execute(
        """
        UPDATE reservations
        SET pdf=%s,pdf_generated_by=%s,pdf_generated_at=NOW(),updated_at=NOW()
        WHERE id=%s
        """,
        (psycopg2.Binary(content), actor, reservation_id),
    )
    reservation_history(cur, reservation_id, "wygenerowano PDF", "Checklista magazyniera")
    return content


def assign_order_salesperson(cur, order_id, email, changed_by):
    normalized = (email or "").strip().lower()
    cur.execute(
        "SELECT salesperson_email FROM shop_orders WHERE id=%s FOR UPDATE",
        (order_id,),
    )
    current = cur.fetchone()
    if not current:
        raise ValueError("Nie znaleziono zamówienia.")
    previous = (current[0] or "").strip().lower()
    if normalized:
        cur.execute(
            """
            SELECT trim(concat_ws(' ',first_name,last_name)),email
            FROM users
            WHERE lower(email)=lower(%s) AND role='sales' AND status='active'
            """,
            (normalized,),
        )
        salesperson = cur.fetchone()
        if not salesperson:
            raise ValueError("Wybrany handlowiec nie istnieje lub jest nieaktywny.")
        name = salesperson[0] or salesperson[1]
        normalized = salesperson[1].strip().lower()
    else:
        name = None
    if previous == normalized:
        return False
    cur.execute(
        """
        UPDATE shop_orders
        SET salesperson_email=%s,salesperson_name=%s,salesperson_assigned_by=%s,
            salesperson_assigned_at=NOW(),updated_at=NOW()
        WHERE id=%s
        """,
        (normalized or None, name, changed_by, order_id),
    )
    cur.execute(
        """
        INSERT INTO shop_order_salesperson_history(
            order_id,previous_salesperson,new_salesperson,changed_by
        ) VALUES (%s,%s,%s,%s)
        """,
        (order_id, previous or None, normalized or None, changed_by),
    )
    shop_history(
        cur,
        order_id,
        "zmieniono handlowca",
        f"{previous or 'brak'} → {normalized or 'brak'}",
    )
    return True


def shop_stage_can_edit(stage_key):
    stage = SHOP_STAGE_BY_KEY.get(stage_key)
    if not stage:
        return False
    return current_user_role() == "admin" or can_shop(stage["permission"])


def shop_stage_dict(value):
    if isinstance(value, dict):
        return dict(value)
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except (TypeError, ValueError):
            return {}
    return {}


def shop_status_from_stages(stages):
    precedence = (
        "cancelled",
        "completed",
        "shipped",
        "packed",
        "sent_to_packing",
        "sales_document_generated",
        "paid",
        "payment_checked",
        "stock_reserved",
        "order_accepted",
    )
    for stage_key in precedence:
        if stages.get(stage_key):
            return SHOP_STAGE_STATUS[stage_key]
    return "Nowe zamówienie"


def set_shop_inventory_issued(cur, order_id, issued):
    ensure_shop_accounting_row(cur, order_id)
    if issued:
        cur.execute(
            "SELECT * FROM shop_accounting WHERE order_id=%s FOR UPDATE",
            (order_id,),
        )
        if not order_can_be_shipped(cur.fetchone()):
            raise ValueError(
                "Księgowość nie odblokowała wydania: wymagane opłacenie albo "
                "pobranie oraz oznaczenie gotowości do wysyłki."
            )
    cur.execute(
        """
        SELECT id,product_id,warehouse,qty,issued_qty
        FROM shop_order_items WHERE order_id=%s ORDER BY id FOR UPDATE
        """,
        (order_id,),
    )
    for item_id, product_id, warehouse, qty, issued_qty in cur.fetchall():
        current = float(issued_qty or 0)
        target = float(qty or 0) if issued else 0.0
        delta = target - current
        if delta > 0:
            cur.execute(
                """
                UPDATE products SET qty=qty-%s
                WHERE id=%s AND warehouse=%s AND qty>=%s
                """,
                (delta, product_id, warehouse, delta),
            )
            if cur.rowcount != 1:
                raise ValueError(
                    "Brak stanu magazynowego lub próba podwójnego wydania."
                )
        elif delta < 0:
            cur.execute(
                "UPDATE products SET qty=qty+%s WHERE id=%s AND warehouse=%s",
                (-delta, product_id, warehouse),
            )
            if cur.rowcount != 1:
                raise ValueError("Nie można odtworzyć stanu magazynowego.")
        cur.execute(
            "UPDATE shop_order_items SET issued_qty=%s WHERE id=%s",
            (target, item_id),
        )


def update_shop_stage(cur, order_id, stage_key, new_value, actor, tracking_number=""):
    if stage_key not in SHOP_STAGE_BY_KEY:
        raise ValueError("Nieprawidłowy etap zamówienia.")
    cur.execute(
        """
        SELECT status,stages,tracking_number
        FROM shop_orders WHERE id=%s FOR UPDATE
        """,
        (order_id,),
    )
    order = cur.fetchone()
    if not order:
        raise ValueError("Nie znaleziono zamówienia.")
    stages = shop_stage_dict(order[1])
    previous = bool(stages.get(stage_key))
    new_value = bool(new_value)
    if previous == new_value:
        return previous, new_value
    tracking_number = (tracking_number or order[2] or "").strip()
    if stage_key == "tracking_entered" and new_value and not tracking_number:
        raise ValueError("Najpierw wpisz numer przesyłki.")
    if stage_key == "sales_document_generated" and new_value:
        cur.execute(
            "SELECT 1 FROM shop_sales_documents WHERE order_id=%s",
            (order_id,),
        )
        if not cur.fetchone():
            raise ValueError("Najpierw wygeneruj dokument sprzedaży.")
    if stage_key == "shipped":
        set_shop_inventory_issued(cur, order_id, new_value)
    if stage_key == "cancelled" and new_value:
        set_shop_inventory_issued(cur, order_id, False)
        if stages.get("shipped"):
            stages["shipped"] = False
            cur.execute(
                """
                INSERT INTO shop_order_stage_history(
                    order_id,stage_key,previous_value,new_value,
                    previous_status,new_status,changed_by
                ) VALUES (%s,'shipped',TRUE,FALSE,%s,'Anulowane',%s)
                """,
                (order_id, order[0], actor),
            )
    accounting_column = {
        "paid": "paid",
        "proforma_issued": "proforma_issued",
        "invoice_issued": "invoice_issued",
        "receipt_issued": "receipt_issued",
        "sales_document_generated": "sales_document_generated",
        "document_sent": "document_sent",
    }.get(stage_key)
    if accounting_column:
        ensure_shop_accounting_row(cur, order_id)
        cur.execute(
            f"""
            UPDATE shop_accounting
            SET {accounting_column}=%s,updated_by=%s,updated_at=NOW()
            WHERE order_id=%s
            """,
            (new_value, actor, order_id),
        )
        sync_accounting_payment_status(cur, order_id)
    stages[stage_key] = new_value
    status = shop_status_from_stages(stages)
    cur.execute(
        """
        UPDATE shop_orders
        SET stages=%s::jsonb,status=%s,
            tracking_number=COALESCE(NULLIF(%s,''),tracking_number),
            updated_at=NOW()
        WHERE id=%s
        """,
        (
            json.dumps(stages, ensure_ascii=False),
            status,
            tracking_number,
            order_id,
        ),
    )
    cur.execute(
        """
        INSERT INTO shop_order_stage_history(
            order_id,stage_key,previous_value,new_value,
            previous_status,new_status,changed_by
        ) VALUES (%s,%s,%s,%s,%s,%s,%s)
        """,
        (order_id, stage_key, previous, new_value, order[0], status, actor),
    )
    label = SHOP_STAGE_BY_KEY[stage_key]["label"]
    shop_history(
        cur,
        order_id,
        "zmieniono etap",
        f"{label}: {'tak' if previous else 'nie'} → {'tak' if new_value else 'nie'}",
    )
    if new_value and stage_key in {
        "paid", "invoice_issued", "receipt_issued", "shipped", "cancelled"
    }:
        notify_order_salesperson(
            cur,
            order_id,
            stage_key,
            f"{label}: zamówienie ma teraz status „{status}”.",
        )
    elif status != order[0]:
        notify_order_salesperson(
            cur,
            order_id,
            "zmiana statusu",
            f"Status zamówienia zmieniono: {order[0]} → {status}.",
        )
    return previous, new_value


def create_shop_document_payload(order, items):
    subtotal_net = sum((item[4] or 0) * (item[3] or 0) for item in items)
    subtotal_gross = sum((item[5] or 0) * (item[3] or 0) for item in items)
    shipping = order[7] or 0
    return {
        "document_number": order[11] or f"DS/{order[1]}",
        "order_number": order[1],
        "date": order[2],
        "receipt_or_invoice": order[11] or "Do uzupełnienia",
        "seller": "Primadera",
        "buyer": order[3],
        "address": order[4],
        "phone": order[5] or "",
        "email": order[6] or "",
        "payment_method": order[8] or "",
        "payment_status": order[9] or "",
        "nip": order[14] or "",
        "items": [
            {
                "name": item[2], "qty": item[3], "net": item[4], "vat": item[6],
                "gross": item[5],
                "total_net": (item[4] or 0) * (item[3] or 0),
                "total_gross": (item[5] or 0) * (item[3] or 0),
            }
            for item in items
        ],
        "shipping": shipping,
        "total_net": subtotal_net,
        "subtotal_gross": subtotal_gross,
        "total_gross": subtotal_gross + shipping,
        "notes": order[13] or "",
    }


def simple_docx_bytes(payload):
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    logo_path = os.path.join(app.root_path, "static", "primadera-logo.png")
    if not os.path.isfile(logo_path):
        raise FileNotFoundError("Brak logo Primadera.")

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)

    def set_cell_shading(cell, fill):
        properties = cell._tc.get_or_add_tcPr()
        shading = properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            properties.append(shading)
        shading.set(qn("w:fill"), fill)

    def set_cell_text(cell, value, bold=False, color=None, align=None):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(value or ""))
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    logo = document.add_paragraph()
    logo.paragraph_format.space_after = Pt(4)
    logo.add_run().add_picture(logo_path, width=Mm(76))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("DOKUMENT SPRZEDAŻY")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(24, 92, 67)

    metadata = document.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.autofit = False
    labels = [
        ("Numer dokumentu", payload["document_number"]),
        ("Numer zamówienia", payload["order_number"]),
        ("Data", payload["date"]),
        ("Płatność", " / ".join(filter(None, [payload["payment_method"], payload["payment_status"]]))),
    ]
    for row, (label, value) in zip(metadata.rows, labels):
        row.cells[0].width = Mm(42)
        row.cells[1].width = Mm(132)
        set_cell_text(row.cells[0], label, bold=True, color=(24, 92, 67))
        set_cell_text(row.cells[1], value)

    document.add_paragraph()
    buyer_heading = document.add_paragraph()
    buyer_heading.paragraph_format.space_after = Pt(3)
    heading_run = buyer_heading.add_run("NABYWCA")
    heading_run.bold = True
    heading_run.font.name = "Arial"
    heading_run.font.size = Pt(11)
    heading_run.font.color.rgb = RGBColor(24, 92, 67)
    buyer_lines = [
        payload["buyer"],
        payload["address"],
        f"NIP: {payload['nip']}" if payload["nip"] else "",
        " · ".join(filter(None, [payload["phone"], payload["email"]])),
    ]
    buyer = document.add_paragraph("\n".join(line for line in buyer_lines if line))
    buyer.paragraph_format.space_after = Pt(10)

    products = document.add_table(rows=1, cols=6)
    products.alignment = WD_TABLE_ALIGNMENT.CENTER
    products.style = "Table Grid"
    products.autofit = False
    widths = [Mm(62), Mm(17), Mm(25), Mm(17), Mm(25), Mm(28)]
    headers = ["Produkt", "Ilość", "Netto/szt.", "VAT", "Brutto/szt.", "Wartość brutto"]
    for cell, header, width in zip(products.rows[0].cells, headers, widths):
        cell.width = width
        set_cell_shading(cell, "185C43")
        set_cell_text(cell, header, bold=True, color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in payload["items"]:
        row = products.add_row()
        values = [
            item["name"],
            f"{item['qty']:g}",
            f"{item['net']:.2f} zł",
            f"{item['vat']:.0f}%",
            f"{item['gross']:.2f} zł",
            f"{item['total_gross']:.2f} zł",
        ]
        for index, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            cell.width = width
            set_cell_text(
                cell,
                value,
                align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT,
            )

    totals = document.add_table(rows=4, cols=2)
    totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
    totals.autofit = False
    total_rows = [
        ("Razem netto", payload["total_net"]),
        ("Razem brutto produktów", payload["subtotal_gross"]),
        ("Wysyłka", payload["shipping"]),
        ("RAZEM BRUTTO", payload["total_gross"]),
    ]
    for index, (row, (label, value)) in enumerate(zip(totals.rows, total_rows)):
        row.cells[0].width = Mm(55)
        row.cells[1].width = Mm(35)
        highlight = index == len(total_rows) - 1
        if highlight:
            set_cell_shading(row.cells[0], "F58220")
            set_cell_shading(row.cells[1], "F58220")
        set_cell_text(
            row.cells[0], label, bold=True,
            color=(255, 255, 255) if highlight else (24, 92, 67),
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        set_cell_text(
            row.cells[1], f"{value:.2f} zł", bold=True,
            color=(255, 255, 255) if highlight else None,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )

    if payload["notes"]:
        notes = document.add_paragraph()
        notes.paragraph_format.space_before = Pt(8)
        notes.add_run("Uwagi: ").bold = True
        notes.add_run(str(payload["notes"]))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def shop_pdf_fonts():
    regular_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    bold_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        r"C:\Windows\Fonts\arialbd.ttf",
    ]
    regular = next((path for path in regular_candidates if os.path.isfile(path)), None)
    bold = next((path for path in bold_candidates if os.path.isfile(path)), None)
    if regular and bold:
        if "ShopSans" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("ShopSans", regular))
            pdfmetrics.registerFont(TTFont("ShopSans-Bold", bold))
        return "ShopSans", "ShopSans-Bold"
    return "Helvetica", "Helvetica-Bold"


def shop_pdf_bytes(payload):
    from xml.sax.saxutils import escape

    logo_path = os.path.join(app.root_path, "static", "primadera-logo.png")
    if not os.path.isfile(logo_path):
        raise FileNotFoundError("Brak logo Primadera.")
    body_font, bold_font = shop_pdf_fonts()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=f"Dokument sprzedaży {payload['document_number']}",
        author="Primadera",
    )
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "ShopNormal",
        parent=styles["Normal"],
        fontName=body_font,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#222222"),
    )
    small = ParagraphStyle(
        "ShopSmall",
        parent=normal,
        fontSize=8,
        leading=10,
    )
    heading = ParagraphStyle(
        "ShopHeading",
        parent=styles["Heading1"],
        fontName=bold_font,
        fontSize=17,
        leading=21,
        textColor=colors.HexColor("#185C43"),
        alignment=TA_RIGHT,
        spaceAfter=8,
    )
    logo = ReportLabImage(logo_path, width=76 * mm, height=17.1 * mm)
    header = Table(
        [[logo, Paragraph("DOKUMENT SPRZEDAŻY", heading)]],
        colWidths=[85 * mm, 83 * mm],
    )
    header.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    metadata = [
        [Paragraph("<b>Numer dokumentu</b>", normal), Paragraph(escape(str(payload["document_number"])), normal)],
        [Paragraph("<b>Numer zamówienia</b>", normal), Paragraph(escape(str(payload["order_number"])), normal)],
        [Paragraph("<b>Data</b>", normal), Paragraph(escape(str(payload["date"])), normal)],
        [
            Paragraph("<b>Płatność</b>", normal),
            Paragraph(
                escape(" / ".join(filter(None, [payload["payment_method"], payload["payment_status"]]))),
                normal,
            ),
        ],
    ]
    metadata_table = Table(metadata, colWidths=[42 * mm, 126 * mm])
    metadata_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), bold_font),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#185C43")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    buyer_lines = [
        f"<b>NABYWCA</b><br/>{escape(str(payload['buyer']))}",
        escape(str(payload["address"])),
        f"NIP: {escape(str(payload['nip']))}" if payload["nip"] else "",
        escape(" · ".join(filter(None, [payload["phone"], payload["email"]]))),
    ]
    story = [
        header,
        metadata_table,
        Spacer(1, 7 * mm),
        Paragraph("<br/>".join(line for line in buyer_lines if line), normal),
        Spacer(1, 6 * mm),
    ]
    data = [[
        Paragraph('<font color="#FFFFFF"><b>Produkt</b></font>', small),
        Paragraph('<font color="#FFFFFF"><b>Ilość</b></font>', small),
        Paragraph('<font color="#FFFFFF"><b>Netto/szt.</b></font>', small),
        Paragraph('<font color="#FFFFFF"><b>VAT</b></font>', small),
        Paragraph('<font color="#FFFFFF"><b>Brutto/szt.</b></font>', small),
        Paragraph('<font color="#FFFFFF"><b>Wartość brutto</b></font>', small),
    ]]
    for item in payload["items"]:
        data.append([
            Paragraph(escape(str(item["name"])), small),
            Paragraph(f"{item['qty']:g}", small),
            Paragraph(f"{item['net']:.2f} zł", small),
            Paragraph(f"{item['vat']:.0f}%", small),
            Paragraph(f"{item['gross']:.2f} zł", small),
            Paragraph(f"{item['total_gross']:.2f} zł", small),
        ])
    table = Table(
        data,
        repeatRows=1,
        colWidths=[57 * mm, 16 * mm, 24 * mm, 16 * mm, 25 * mm, 30 * mm],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#185C43")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C8C0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    totals_data = [
        [Paragraph("Razem netto", normal), Paragraph(f"{payload['total_net']:.2f} zł", normal)],
        [Paragraph("Razem brutto produktów", normal), Paragraph(f"{payload['subtotal_gross']:.2f} zł", normal)],
        [Paragraph("Wysyłka", normal), Paragraph(f"{payload['shipping']:.2f} zł", normal)],
        [
            Paragraph('<font color="#FFFFFF"><b>RAZEM BRUTTO</b></font>', normal),
            Paragraph(
                f'<font color="#FFFFFF"><b>{payload["total_gross"]:.2f} zł</b></font>',
                normal,
            ),
        ],
    ]
    totals = Table(totals_data, colWidths=[55 * mm, 35 * mm], hAlign="RIGHT")
    totals.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F58220")),
        ("TEXTCOLOR", (0, -1), (-1, -1), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [
        table,
        Spacer(1, 5 * mm),
        totals,
    ]
    if payload["notes"]:
        story += [
            Spacer(1, 5 * mm),
            Paragraph(f"<b>Uwagi:</b> {escape(str(payload['notes']))}", normal),
        ]
    doc.build(story)
    return buffer.getvalue()


def load_shop_document_data(cur, order_id):
    cur.execute(
        """
        SELECT id,order_number,order_date,customer_name,delivery_address,phone,email,
               shipping_cost,payment_method,payment_status,status,sales_document_number,
               tracking_number,notes,nip
        FROM shop_orders WHERE id=%s
        """,
        (order_id,),
    )
    order = cur.fetchone()
    if not order:
        raise ValueError("Nie znaleziono zamówienia.")
    cur.execute(
        """
        SELECT id,product_id,product_name,qty,price_netto,price_brutto,vat
        FROM shop_order_items WHERE order_id=%s ORDER BY id
        """,
        (order_id,),
    )
    items = cur.fetchall()
    if not items:
        raise ValueError("Zamówienie nie zawiera produktów.")
    return order, items


def generate_shop_sales_document(cur, order_id, created_by):
    cur.execute(
        """
        SELECT id,voided_at FROM shop_sales_documents WHERE order_id=%s
        """,
        (order_id,),
    )
    existing_document = cur.fetchone()
    if existing_document and existing_document[1] is not None:
        raise ValueError(
            "Dokument sprzedaży jest anulowany. Administrator musi go najpierw przywrócić."
        )
    order, items = load_shop_document_data(cur, order_id)
    payload = create_shop_document_payload(order, items)
    docx = simple_docx_bytes(payload)
    pdf = shop_pdf_bytes(payload)
    cur.execute(
        """
        INSERT INTO shop_sales_documents(
            order_id,document_number,editable_data,docx,pdf,created_by
        ) VALUES (%s,%s,%s::jsonb,%s,%s,%s)
        ON CONFLICT (order_id) DO UPDATE SET
            document_number=EXCLUDED.document_number,
            editable_data=EXCLUDED.editable_data,
            docx=EXCLUDED.docx,
            pdf=EXCLUDED.pdf,
            created_by=EXCLUDED.created_by,
            created_at=NOW()
        RETURNING id,document_number
        """,
        (
            order_id,
            payload["document_number"],
            json.dumps(payload, ensure_ascii=False, default=str),
            psycopg2.Binary(docx),
            psycopg2.Binary(pdf),
            created_by,
        ),
    )
    document = cur.fetchone()
    sync_accounting_document(
        cur,
        order_id,
        "sales",
        payload["document_number"],
        True,
        created_by or "system",
        payload,
    )
    ensure_shop_accounting_row(cur, order_id)
    cur.execute(
        """
        UPDATE shop_accounting
        SET sales_document_generated=TRUE,updated_by=%s,updated_at=NOW()
        WHERE order_id=%s
        """,
        (created_by, order_id),
    )
    update_shop_stage(
        cur,
        order_id,
        "sales_document_generated",
        True,
        created_by or "system",
    )
    return document, payload


# 🔒 LOGIN REQUIRED
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return redirect('/login')
        if session.get('role') != 'admin':
            return "Brak uprawnień", 403
        return f(*args, **kwargs)
    return decorated


def current_user_can(permission):
    role = session.get("role")
    return permission in ROLE_PERMISSIONS.get(role, set())


def permission_required(permission):
    def outer(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if 'user' not in session:
                return redirect('/login')
            if not current_user_can(permission):
                return "Brak uprawnień", 403
            return f(*args, **kwargs)
        return decorated
    return outer


def general_import_groups():
    if session.get("role") == "admin":
        return {"warehouse", "accounting"}
    groups = set()
    if session.get("role") == "warehouse" and session.get("can_import_warehouse"):
        groups.add("warehouse")
    if session.get("role") == "accounting" and session.get("can_import_accounting"):
        groups.add("accounting")
    return groups


def import_access_error():
    if not general_import_groups():
        return "Brak uprawnień do importu ogólnego.", 403
    return None


def import_run_allowed(run):
    return bool(
        run
        and (
            session.get("role") == "admin"
            or str(run[2] or "").casefold() == str(session.get("user") or "").casefold()
        )
    )


def issue_import_access_error():
    if session.get("role") == "admin":
        return None
    if session.get("role") == "warehouse" and session.get("can_import_warehouse"):
        return None
    return "Brak uprawnień do importu wydań.", 403


def issue_import_run_allowed(run):
    return bool(
        run
        and (
            session.get("role") == "admin"
            or str(run[2] or "").casefold()
            == str(session.get("user") or "").casefold()
        )
    )


def issue_import_allow_general_stock(cur):
    cur.execute(
        "SELECT value FROM app_settings WHERE key='issue_import_allow_general_stock'"
    )
    row = cur.fetchone()
    return bool(row and str(row[0]).strip().lower() == "true")


def issue_import_run_query(cur, import_id, for_update=False):
    cur.execute(
        """
        SELECT id,filename,imported_by,status,detected_sheets,summary,errors,
               created_at,completed_at,undone_at,undone_by
        FROM issue_imports WHERE id=%s
        """
        + (" FOR UPDATE" if for_update else ""),
        (import_id,),
    )
    return cur.fetchone()


def issue_import_rows_query(cur, import_id, for_update=False):
    cur.execute(
        """
        SELECT id,sheet_name,row_number,source_data,normalized_data,
               duplicate_data,resolution,included,validation_errors
        FROM issue_import_rows
        WHERE import_id=%s ORDER BY sheet_name,row_number,id
        """
        + (" FOR UPDATE" if for_update else ""),
        (import_id,),
    )
    return cur.fetchall()


def detect_issue_import_duplicate(cur, data):
    number = str(data.get("doc_number") or "").strip()
    if not number:
        return None
    cur.execute(
        """
        SELECT id,doc_number,date,kontrahent
        FROM issue_docs
        WHERE lower(doc_number)=lower(%s)
          AND COALESCE(movement_type,'WZ') IN ('WZ','RW')
          AND voided_at IS NULL
        ORDER BY id LIMIT 1
        """,
        (number,),
    )
    row = cur.fetchone()
    if not row:
        return None
    return {
        "id": row[0],
        "label": " · ".join(str(value or "") for value in row[1:]),
    }


def issue_import_row_context(
    cur,
    data,
    allow_general_stock,
    duplicate=None,
    resolution="new",
):
    errors = validate_issue_row(data, UNITS, WAREHOUSES)
    context = {"product": None, "package": None, "existing_item": None}
    if errors:
        return errors, context
    product_name = str(data.get("product_name") or "").strip()
    warehouse = str(data.get("warehouse") or "").strip()
    cur.execute(
        """
        SELECT id,qty,unit,price_netto,vat
        FROM products
        WHERE lower(name)=lower(%s) AND warehouse=%s
        ORDER BY id LIMIT 1 FOR UPDATE
        """,
        (product_name, warehouse),
    )
    product = cur.fetchone()
    if not product:
        errors.append(
            f"Produkt {product_name} nie istnieje w magazynie {warehouse}."
        )
        return errors, context
    context["product"] = product
    unit = str(data.get("unit") or "").strip()
    if unit != str(product[2] or "").strip():
        errors.append(
            f"Jednostka produktu to {product[2]}, a w pliku podano {unit}."
        )
    package = None
    raw_package_id = str(data.get("package_id") or "").strip()
    package_number = str(data.get("package_number") or "").strip()
    if raw_package_id:
        try:
            package_id = int(raw_package_id)
        except ValueError:
            errors.append("Wybrana paczka jest nieprawidłowa.")
        else:
            cur.execute(
                """
                SELECT id,number,qty FROM packages
                WHERE id=%s AND product_id=%s AND warehouse=%s AND status='active'
                FOR UPDATE
                """,
                (package_id, product[0], warehouse),
            )
            package = cur.fetchone()
    elif package_number:
        cur.execute(
            """
            SELECT id,number,qty FROM packages
            WHERE product_id=%s AND warehouse=%s AND lower(number)=lower(%s)
              AND status='active'
            ORDER BY id LIMIT 1 FOR UPDATE
            """,
            (product[0], warehouse, package_number),
        )
        package = cur.fetchone()
    if (raw_package_id or package_number) and not package:
        errors.append(
            "Nie znaleziono wskazanej aktywnej paczki. Wybierz właściwą paczkę."
        )
    if not raw_package_id and not package_number and not allow_general_stock:
        errors.append(
            "Numer paczki jest wymagany. Administrator może zezwolić na "
            "wydanie z ogólnego stanu."
        )
    if package:
        context["package"] = package
        data["package_id"] = package[0]
        data["package_number"] = package[1]
    old_qty = 0.0
    if duplicate and resolution == "update":
        package_clause = "package_id=%s" if package else "package_id IS NULL"
        item_params = (
            (duplicate["id"], product[0], package[0])
            if package
            else (duplicate["id"], product[0])
        )
        cur.execute(
            f"""
            SELECT id,qty,package_id,package_number,dimension,species,notes
            FROM issue_items
            WHERE doc_id=%s AND product_id=%s
              AND {package_clause}
            ORDER BY id LIMIT 1 FOR UPDATE
            """,
            item_params,
        )
        existing_item = cur.fetchone()
        if existing_item:
            context["existing_item"] = existing_item
            old_qty = float(existing_item[1] or 0)
    qty = float(data.get("qty") or 0)
    if float(product[1] or 0) + old_qty + 1e-9 < qty:
        errors.append(
            f"Dostępny stan produktu to {float(product[1] or 0) + old_qty:g}."
        )
    if package and float(package[2] or 0) + old_qty + 1e-9 < qty:
        errors.append(
            f"W paczce {package[1]} dostępne jest "
            f"{float(package[2] or 0) + old_qty:g}."
        )
    return errors, context


def detect_import_duplicate(cur, entity_type, data):
    if entity_type == "product":
        cur.execute(
            """
            SELECT id,name,qty,unit,warehouse FROM products
            WHERE lower(name)=lower(%s) AND warehouse=%s
            ORDER BY id LIMIT 1
            """,
            (str(data.get("name") or "").strip(), str(data.get("warehouse") or "").strip()),
        )
    elif entity_type == "package":
        cur.execute(
            """
            SELECT id,number,qty,warehouse FROM packages
            WHERE lower(number)=lower(%s) AND warehouse=%s
            ORDER BY id LIMIT 1
            """,
            (
                str(data.get("package_number") or "").strip(),
                str(data.get("warehouse") or "").strip(),
            ),
        )
    elif entity_type in {"issue", "receipt", "document"}:
        number = str(data.get("doc_number") or "").strip()
        if not number:
            return None
        cur.execute(
            "SELECT id,doc_number,date,kontrahent FROM issue_docs WHERE lower(doc_number)=lower(%s) LIMIT 1",
            (number,),
        )
    elif entity_type == "contractor":
        name = str(data.get("contractor") or "").strip()
        nip = str(data.get("nip") or "").strip()
        cur.execute(
            """
            SELECT id,name,nip FROM contractors
            WHERE lower(name)=lower(%s)
               OR (%s<>'' AND regexp_replace(COALESCE(nip,''),'[^0-9]','','g')=
                   regexp_replace(%s,'[^0-9]','','g'))
            ORDER BY id LIMIT 1
            """,
            (name, nip, nip),
        )
    elif entity_type in {"shop_order", "accounting"}:
        cur.execute(
            "SELECT id,order_number,customer_name,status FROM shop_orders WHERE lower(order_number)=lower(%s) LIMIT 1",
            (str(data.get("order_number") or "").strip(),),
        )
    else:
        return None
    row = cur.fetchone()
    if not row:
        return None
    return {"id": row[0], "label": " · ".join(str(value or "") for value in row[1:])}


def unique_import_identifier(cur, table, column, requested, import_id, row_id):
    requested = str(requested or "").strip()
    candidate = requested or f"IMPORT/{import_id}/{row_id}"
    cur.execute(
        f"SELECT 1 FROM {table} WHERE lower({column})=lower(%s) LIMIT 1",
        (candidate,),
    )
    if not cur.fetchone():
        return candidate
    return f"{candidate}-IMPORT-{import_id}-{row_id}"


def log_action(action, entity_type=None, entity_id=None, details=None, conn=None):
    target_conn = conn or db()
    try:
        cur = target_conn.cursor()
        cur.execute(
            """
            INSERT INTO action_logs(actor_email, actor_uid, action, entity_type, entity_id, details)
            VALUES (%s,%s,%s,%s,%s,%s::jsonb)
            """,
            (
                session.get("user"),
                session.get("uid"),
                action,
                entity_type,
                str(entity_id) if entity_id is not None else None,
                json.dumps(details or {}, ensure_ascii=False),
            ),
        )
        if conn is None:
            target_conn.commit()
    except Exception:
        logger.warning("Action log skipped.", exc_info=True)
    finally:
        if conn is None:
            target_conn.close()


def password_validation_error(password):
    password = password or ""
    if len(password) < 12:
        return "Hasło musi mieć co najmniej 12 znaków."
    if not any(character.islower() for character in password):
        return "Hasło musi zawierać małą literę."
    if not any(character.isupper() for character in password):
        return "Hasło musi zawierać dużą literę."
    if not any(character.isdigit() for character in password):
        return "Hasło musi zawierać cyfrę."
    if not any(not character.isalnum() for character in password):
        return "Hasło musi zawierać znak specjalny."
    return None


def security_event(cur, user_id, email, event_type, actor=None, details=None):
    cur.execute(
        """
        INSERT INTO user_security_events(
            user_id,user_email,event_type,actor_email,details
        ) VALUES (%s,%s,%s,%s,%s::jsonb)
        """,
        (
            user_id,
            (email or "").strip().lower(),
            event_type,
            actor or (session.get("user") if has_request_context() else None),
            json.dumps(details or {}, ensure_ascii=False),
        ),
    )


def firebase_password_sign_in(email, password):
    api_key = FIREBASE_CONFIG.get("apiKey")
    if not api_key:
        raise RuntimeError("Firebase Web API nie jest skonfigurowane.")
    endpoint = (
        "https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword"
        f"?key={api_key}"
    )
    payload = json.dumps(
        {
            "email": (email or "").strip().lower(),
            "password": password or "",
            "returnSecureToken": True,
        }
    ).encode("utf-8")
    request_object = urlrequest.Request(
        endpoint,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urlrequest.urlopen(request_object, timeout=12) as response:
            return json.loads(response.read().decode("utf-8"))
    except urlerror.HTTPError as exc:
        # Firebase returns credential details in the response. Do not expose them.
        raise ValueError("Nieprawidłowy adres e-mail lub hasło.") from exc


@app.context_processor
def inject_permissions():
    return {"can": current_user_can, "role_labels": ROLE_LABELS}


# 🔐 LOGIN
@app.route('/login')
def login():
    if 'user' in session:
        return redirect('/')
    try:
        return render_template(
            "login.html",
            firebase_config=FIREBASE_CONFIG,
            firebase_admin_ready=FIREBASE_ADMIN_READY,
            firebase_admin_error=FIREBASE_ADMIN_ERROR,
            missing_firebase_web_envs=get_missing_firebase_web_envs(),
        )
    except Exception:
        logger.exception("Login page rendering failed.")
        return (
            "<!doctype html><html lang='pl'><meta charset='utf-8'>"
            "<title>Logowanie niedostępne</title><body>"
            "<h1>Logowanie jest chwilowo niedostępne</h1>"
            "<p>Administrator został poinformowany. Spróbuj ponownie później.</p>"
            "</body></html>",
            503,
        )


@app.route('/register')
def register():
    return redirect('/login')


@app.route('/auth/password-login', methods=['POST'])
@limiter.limit("10 per minute")
def password_login():
    payload = request.get_json(silent=True) or {}
    email = (payload.get("email") or "").strip().lower()
    password = payload.get("password") or ""
    if not email or not password:
        return jsonify({"ok": False, "error": "Podaj adres e-mail i hasło."}), 400
    conn = None
    try:
        ensure_db_initialized()
        conn = db()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id,firebase_uid,status,failed_login_attempts,
                   (locked_until IS NOT NULL AND locked_until>NOW())
            FROM users WHERE lower(email)=lower(%s) FOR UPDATE
            """,
            (email,),
        )
        account = cur.fetchone()
        if not account:
            conn.rollback()
            return jsonify({
                "ok": False,
                "error": "Nieprawidłowy adres e-mail lub hasło.",
            }), 401
        user_id, uid, status, failed_attempts, is_temporarily_locked = account
        if status == "inactive":
            conn.rollback()
            return jsonify({"ok": False, "error": "Konto jest nieaktywne."}), 403
        if status == "blocked":
            conn.rollback()
            return jsonify({
                "ok": False,
                "error": "Konto jest zablokowane. Skontaktuj się z administratorem.",
            }), 423
        if is_temporarily_locked:
            conn.rollback()
            return jsonify({
                "ok": False,
                "error": (
                    f"Konto jest tymczasowo zablokowane. Spróbuj ponownie za "
                    f"{LOGIN_LOCK_MINUTES} minut."
                ),
            }), 423
        try:
            result = firebase_password_sign_in(email, password)
        except ValueError:
            attempts = int(failed_attempts or 0) + 1
            if attempts >= 5 and LOGIN_LOCK_MODE == "admin":
                cur.execute(
                    """
                    UPDATE users SET failed_login_attempts=%s,status='blocked',
                        locked_until=NULL,last_failed_login_at=NOW(),updated_at=NOW()
                    WHERE id=%s
                    """,
                    (attempts, user_id),
                )
                security_event(
                    cur, user_id, email, "account_locked", None,
                    {"mode": "admin", "failed_attempts": attempts},
                )
                try:
                    auth.update_user(uid, disabled=True)
                except Exception:
                    logger.warning("Firebase account lock synchronization failed.", exc_info=True)
            elif attempts >= 5:
                cur.execute(
                    """
                    UPDATE users SET failed_login_attempts=%s,
                        locked_until=NOW()+(%s*INTERVAL '1 minute'),
                        last_failed_login_at=NOW(),updated_at=NOW()
                    WHERE id=%s
                    """,
                    (attempts, LOGIN_LOCK_MINUTES, user_id),
                )
                security_event(
                    cur, user_id, email, "account_temporarily_locked", None,
                    {"minutes": LOGIN_LOCK_MINUTES, "failed_attempts": attempts},
                )
            else:
                cur.execute(
                    """
                    UPDATE users SET failed_login_attempts=%s,
                        last_failed_login_at=NOW(),updated_at=NOW()
                    WHERE id=%s
                    """,
                    (attempts, user_id),
                )
                security_event(
                    cur, user_id, email, "login_failed", None,
                    {"failed_attempts": attempts},
                )
            conn.commit()
            message = (
                "Konto zostało zablokowane po 5 nieudanych próbach."
                if attempts >= 5
                else "Nieprawidłowy adres e-mail lub hasło."
            )
            return jsonify({"ok": False, "error": message}), (
                423 if attempts >= 5 else 401
            )
        cur.execute(
            """
            UPDATE users SET failed_login_attempts=0,locked_until=NULL,
                updated_at=NOW() WHERE id=%s
            """,
            (user_id,),
        )
        conn.commit()
        return jsonify({"ok": True, "idToken": result.get("idToken")})
    except Exception:
        if conn:
            conn.rollback()
        logger.exception("Password login failed.")
        return jsonify({
            "ok": False,
            "error": "Logowanie jest chwilowo niedostępne. Spróbuj ponownie.",
        }), 503
    finally:
        if conn:
            conn.close()


@app.route('/auth/forgot-password', methods=['POST'])
@limiter.limit("5 per hour")
def forgot_password():
    email = ((request.get_json(silent=True) or {}).get("email") or "").strip().lower()
    if email:
        conn = None
        try:
            ensure_db_initialized()
            conn = db()
            cur = conn.cursor()
            cur.execute(
                "SELECT id,status FROM users WHERE lower(email)=lower(%s)",
                (email,),
            )
            account = cur.fetchone()
            if account and account[1] == "active":
                send_firebase_password_reset(email)
                security_event(cur, account[0], email, "password_reset_requested")
                conn.commit()
        except Exception:
            if conn:
                conn.rollback()
            logger.warning("Password reset request could not be completed.", exc_info=True)
        finally:
            if conn:
                conn.close()
    # Always return the same response to prevent account enumeration.
    return jsonify({
        "ok": True,
        "message": (
            "Jeśli konto istnieje i jest aktywne, wysłaliśmy jednorazowy link "
            "do zmiany hasła."
        ),
    })


@app.route('/auth/session', methods=['POST'])
@limiter.limit("10 per minute")
def create_session():
    payload = request.get_json(silent=True) or {}
    id_token = payload.get("idToken")
    if not id_token:
        return jsonify({"ok": False, "error": "Brak tokenu"}), 400

    if not FIREBASE_ADMIN_READY:
        return jsonify({
            "ok": False,
            "error": FIREBASE_ADMIN_ERROR or "Firebase Admin nie jest skonfigurowany.",
        }), 503
    try:
        decoded = auth.verify_id_token(id_token, check_revoked=True)
        email = (decoded.get("email") or "").strip().lower()
        uid = decoded.get("uid")
        firebase_user = auth.get_user(uid)
        if firebase_user.disabled:
            return jsonify({"ok": False, "error": "Konto jest zablokowane."}), 403
    except Exception:
        logger.info("Firebase ID token rejected.", exc_info=True)
        return jsonify({"ok": False, "error": "Nieprawidłowy lub wygasły token"}), 401
    if not email or not uid:
        return jsonify({"ok": False, "error": "Brak danych użytkownika"}), 400

    conn = None
    try:
        ensure_db_initialized()
        conn = db()
        cur = conn.cursor()
        cur.execute(
            """
            SELECT role,status,can_import_warehouse,can_import_accounting,
                   must_change_password,id,
                   (locked_until IS NOT NULL AND locked_until>NOW())
            FROM users
            WHERE firebase_uid=%s OR lower(email)=lower(%s)
            FOR UPDATE
            """,
            (uid, email),
        )
        row = cur.fetchone()
        if not row:
            if email not in ADMIN_EMAILS:
                if not ADMIN_EMAILS:
                    logger.error(
                        "Login rejected because ADMIN_EMAILS is empty and no application user exists."
                    )
                return jsonify({
                    "ok": False,
                    "error": "Brak konta. Skontaktuj się z administratorem.",
                }), 403
            display_name = (decoded.get("name") or firebase_user.display_name or "").strip()
            first_name, _, last_name = display_name.partition(" ")
            role = "admin"
            can_import_warehouse = True
            can_import_accounting = True
            must_change_password = False
            cur.execute(
                """
                INSERT INTO users(
                    firebase_uid,first_name,last_name,email,role,status,last_login_at
                )
                VALUES (%s,%s,%s,%s,'admin','active',NOW())
                RETURNING id
                """,
                (uid, first_name, last_name, email),
            )
            user_id = cur.fetchone()[0]
            security_event(cur, user_id, email, "login_successful", email)
        else:
            role, status = row[:2]
            can_import_warehouse = bool(row[2]) if len(row) > 2 else False
            can_import_accounting = bool(row[3]) if len(row) > 3 else False
            must_change_password = bool(row[4]) if len(row) > 4 else False
            user_id = row[5] if len(row) > 5 else None
            is_temporarily_locked = bool(row[6]) if len(row) > 6 else False
            if role == "employee":
                role = "warehouse"
            if status != "active":
                return jsonify({"ok": False, "error": "Konto jest zablokowane."}), 403
            if is_temporarily_locked:
                return jsonify({
                    "ok": False,
                    "error": "Konto jest tymczasowo zablokowane.",
                }), 423
            cur.execute(
                """
                UPDATE users SET firebase_uid=%s,email=%s,role=%s,
                    failed_login_attempts=0,locked_until=NULL,last_login_at=NOW(),
                    updated_at=NOW()
                WHERE firebase_uid=%s OR lower(email)=lower(%s)
                """,
                (uid, email, role, uid, email),
            )
            security_event(cur, user_id, email, "login_successful", email)
        conn.commit()
    except Exception:
        if conn:
            conn.rollback()
        logger.exception("Application session database synchronization failed.")
        return jsonify({
            "ok": False,
            "error": "Baza danych jest chwilowo niedostępna. Spróbuj ponownie później.",
        }), 503
    finally:
        if conn:
            conn.close()
    expires = timedelta(days=5)
    try:
        firebase_cookie = auth.create_session_cookie(id_token, expires_in=expires)
    except Exception:
        logger.exception("Firebase session cookie creation failed.")
        return jsonify({"ok": False, "error": "Nie udało się utworzyć bezpiecznej sesji."}), 500
    session['user'] = email
    session['role'] = role
    session['uid'] = uid
    session['can_import_warehouse'] = role == "admin" or can_import_warehouse
    session['can_import_accounting'] = role == "admin" or can_import_accounting
    session['must_change_password'] = bool(must_change_password)
    response = make_response(jsonify({
        "ok": True,
        "role": role,
        "mustChangePassword": bool(must_change_password),
        "redirect": "/change-password?required=1" if must_change_password else "/",
    }))
    response.set_cookie(
        "firebase_session",
        firebase_cookie,
        max_age=int(expires.total_seconds()),
        httponly=True,
        secure=os.environ.get("RENDER", "").lower() == "true",
        samesite="Lax",
    )
    return response


def render_change_password(error=None, status=200):
    return render_template(
        "change_password.html",
        error=error,
        required=bool(session.get("must_change_password")),
    ), status


@app.route('/change-password')
@login_required
def change_password_page():
    return render_change_password()


@app.route('/auth/change-password', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def change_password():
    current_password = request.form.get("current_password") or ""
    new_password = request.form.get("new_password") or ""
    confirmation = request.form.get("confirm_password") or ""
    if not current_password:
        return render_change_password("Podaj aktualne hasło.", 400)
    if new_password != confirmation:
        return render_change_password("Nowe hasła nie są identyczne.", 400)
    validation_error = password_validation_error(new_password)
    if validation_error:
        return render_change_password(validation_error, 400)
    if secrets.compare_digest(current_password, new_password):
        return render_change_password("Nowe hasło musi różnić się od aktualnego.", 400)
    try:
        firebase_password_sign_in(session.get("user"), current_password)
    except ValueError:
        return render_change_password("Aktualne hasło jest nieprawidłowe.", 400)
    except Exception:
        logger.exception("Current password verification failed.")
        return render_change_password("Nie udało się zweryfikować hasła.", 503)

    conn = db()
    cur = conn.cursor()
    try:
        auth.update_user(session.get("uid"), password=new_password)
        sign_in_result = firebase_password_sign_in(session.get("user"), new_password)
        new_id_token = sign_in_result.get("idToken")
        if not new_id_token:
            raise RuntimeError("Firebase did not return an ID token.")
        expires = timedelta(days=5)
        firebase_cookie = auth.create_session_cookie(new_id_token, expires_in=expires)
        cur.execute(
            """
            UPDATE users SET must_change_password=FALSE,password_changed_at=NOW(),
                failed_login_attempts=0,locked_until=NULL,updated_at=NOW()
            WHERE firebase_uid=%s RETURNING id,email
            """,
            (session.get("uid"),),
        )
        user_row = cur.fetchone()
        if not user_row:
            raise ValueError("Nie znaleziono konta użytkownika.")
        security_event(
            cur,
            user_row[0],
            user_row[1],
            "password_changed",
            session.get("user"),
        )
        conn.commit()
        session["must_change_password"] = False
        response = make_response(redirect("/?password=changed"))
        response.set_cookie(
            "firebase_session",
            firebase_cookie,
            max_age=int(expires.total_seconds()),
            httponly=True,
            secure=os.environ.get("RENDER", "").lower() == "true",
            samesite="Lax",
        )
        return response
    except ValueError as exc:
        conn.rollback()
        return render_change_password(str(exc), 404)
    except Exception:
        conn.rollback()
        logger.exception("Password change failed.")
        return render_change_password(
            "Nie udało się zmienić hasła. Spróbuj ponownie.", 500
        )
    finally:
        conn.close()


@app.route('/logout', methods=['POST'])
def logout():
    session.clear()
    response = make_response(redirect('/login'))
    response.delete_cookie("firebase_session")
    return response


@app.before_request
def require_login_for_private_app():
    if IS_PRODUCTION and not request.is_secure:
        secure_url = request.url.replace("http://", "https://", 1)
        return redirect(secure_url, code=308)
    if request.method == "OPTIONS":
        return Response(status=204)
    public_routes = {"static", "favicon", "health", "web_manifest", "service_worker"}
    if request.endpoint in public_routes:
        return None

    def csrf_error():
        if request.is_json or request.path.startswith("/api/"):
            return jsonify({"ok": False, "error": "Sesja formularza wygasła. Odśwież stronę."}), 400
        return "Sesja formularza wygasła. Odśwież stronę.", 400

    def validate_csrf():
        sent_token = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token")
        if not sent_token or not secrets.compare_digest(sent_token, session.get("_csrf_token", "")):
            return csrf_error()
        return None

    if request.endpoint in {"create_session", "password_login", "forgot_password"}:
        return validate_csrf()

    if request.endpoint == "logout":
        if request.method in {"POST", "PUT", "PATCH", "DELETE"}:
            return validate_csrf()
        return None

    firebase_cookie = request.cookies.get("firebase_session")
    decoded = None
    if firebase_cookie and FIREBASE_ADMIN_READY:
        try:
            decoded = auth.verify_session_cookie(firebase_cookie, check_revoked=True)
        except Exception:
            logger.info("Firebase session cookie rejected.", exc_info=True)
            decoded = None

    if decoded:
        uid = decoded.get("uid")
        email = (decoded.get("email") or "").strip().lower()
        try:
            ensure_db_initialized()
            conn = db()
            cur = conn.cursor()
            cur.execute(
                """
                SELECT role,status,can_import_warehouse,can_import_accounting,
                       must_change_password,
                       (locked_until IS NOT NULL AND locked_until>NOW())
                FROM users
                WHERE firebase_uid=%s AND lower(email)=lower(%s)
                """,
                (uid, email),
            )
            account = cur.fetchone()
            conn.close()
        except Exception:
            logger.exception("Authenticated account lookup failed.")
            if request.path.startswith("/api/") or request.is_json:
                return jsonify({
                    "ok": False,
                    "error": "Baza danych jest chwilowo niedostępna.",
                }), 503
            return render_template(
                "error.html",
                title="Aplikacja chwilowo niedostępna",
                message="Nie udało się połączyć z bazą danych. Spróbuj ponownie później.",
            ), 503
        if account and account[1] == "active" and not bool(account[5]):
            session["user"] = email
            session["uid"] = uid
            session["role"] = account[0]
            session["can_import_warehouse"] = (
                account[0] == "admin" or (bool(account[2]) if len(account) > 2 else False)
            )
            session["can_import_accounting"] = (
                account[0] == "admin" or (bool(account[3]) if len(account) > 3 else False)
            )
            session["must_change_password"] = bool(account[4])
            if request.endpoint == "login":
                return redirect("/")
        else:
            decoded = None

    if not decoded:
        session.pop("user", None)
        session.pop("uid", None)
        session.pop("role", None)
        if request.endpoint == "login":
            return None
        if request.path.startswith("/api/") or request.is_json:
            return jsonify({"ok": False, "error": "Wymagane logowanie."}), 401
        return redirect("/login")

    if request.method in {"POST", "PUT", "PATCH", "DELETE"}:
        csrf_failure = validate_csrf()
        if csrf_failure:
            return csrf_failure
    if session.get("must_change_password") and request.endpoint not in {
        "change_password_page",
        "change_password",
        "logout",
    }:
        if request.path.startswith("/api/") or request.is_json:
            return jsonify({
                "ok": False,
                "error": "Przed użyciem aplikacji zmień hasło tymczasowe.",
            }), 403
        return redirect("/change-password?required=1")
    required_permission = ENDPOINT_PERMISSIONS.get(request.endpoint)
    if required_permission and not current_user_can(required_permission):
        return "Brak uprawnień", 403
    return None


@app.route('/favicon.ico')
def favicon():
    return Response(status=204)


@app.route('/health')
def health():
    try:
        ensure_db_initialized()
        conn = db()
        cur = conn.cursor()
        cur.execute("SELECT 1")
        cur.fetchone()
        conn.close()
        maybe_start_daily_backup()
        return jsonify({"status": "ok"})
    except Exception:
        logger.exception("Health check failed.")
        return jsonify({"status": "error", "database": "unavailable"}), 503


def maybe_start_daily_backup():
    if not os.environ.get("BACKUP_ENCRYPTION_KEY") or not FIREBASE_ADMIN_READY:
        return
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT 1 FROM backup_runs
        WHERE status='completed' AND created_at::date=(NOW() AT TIME ZONE 'UTC')::date
        LIMIT 1
        """
    )
    already_completed = bool(cur.fetchone())
    conn.close()
    if already_completed or not DAILY_BACKUP_LOCK.acquire(blocking=False):
        return

    def run_backup():
        try:
            with app.app_context():
                perform_database_backup("automatic-fallback")
        except Exception:
            logger.exception("Automatic fallback backup failed.")
        finally:
            DAILY_BACKUP_LOCK.release()

    threading.Thread(
        target=run_backup,
        name="daily-database-backup",
        daemon=True,
    ).start()


@app.after_request
def add_private_cache_headers(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"] = (
        "camera=(), microphone=(), geolocation=(), payment=(), usb=()"
    )
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://www.gstatic.com; "
        "style-src 'self' 'unsafe-inline'; "
        "font-src 'self' data:; "
        "img-src 'self' data:; "
        "connect-src 'self' https://identitytoolkit.googleapis.com "
        "https://securetoken.googleapis.com https://www.googleapis.com; "
        "frame-ancestors 'none'; base-uri 'self'; form-action 'self'"
    )
    if request.is_secure:
        response.headers["Strict-Transport-Security"] = (
            "max-age=31536000; includeSubDomains"
        )
    origin = request.headers.get("Origin", "").rstrip("/")
    if origin in ALLOWED_ORIGINS:
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Vary"] = "Origin"
        response.headers["Access-Control-Allow-Credentials"] = "true"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-CSRF-Token"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response


@app.errorhandler(psycopg2.Error)
def handle_database_error(error):
    logger.exception("Unhandled database error.", exc_info=error)
    if request.path.startswith("/api/") or request.is_json:
        return jsonify({
            "ok": False,
            "error": "Baza danych jest chwilowo niedostępna.",
        }), 503
    return render_template(
        "error.html",
        title="Aplikacja chwilowo niedostępna",
        message="Nie udało się wykonać operacji na bazie danych. Spróbuj ponownie później.",
    ), 503


@app.errorhandler(Exception)
def handle_unexpected_error(error):
    if isinstance(error, HTTPException):
        return error
    logger.exception("Unhandled application error.", exc_info=error)
    if request.path.startswith("/api/") or request.is_json:
        return jsonify({
            "ok": False,
            "error": "Wystąpił nieoczekiwany błąd. Spróbuj ponownie później.",
        }), 500
    return render_template(
        "error.html",
        title="Nie udało się wykonać operacji",
        message="Wystąpił nieoczekiwany błąd. Spróbuj ponownie później.",
    ), 500


@app.route('/')
@login_required
def home():
    if current_user_role() == "sales":
        return redirect("/dashboard")
    return render_template("home.html")


@app.route('/dashboard', endpoint='dashboard_page_view')
@login_required
def dashboard_page():
    conn = db()
    cur = conn.cursor()
    if current_user_role() == "sales":
        email = session.get("user", "")
        cur.execute(
            """
            SELECT COUNT(*),
                   COUNT(*) FILTER (WHERE status IN ('Nowe zamówienie','Przyjęte','Oczekuje na płatność')),
                   COUNT(*) FILTER (WHERE status NOT IN ('Nowe zamówienie','Przyjęte','Oczekuje na płatność','Zakończone','Anulowane')),
                   COUNT(*) FILTER (WHERE status='Zakończone')
            FROM shop_orders
            WHERE lower(COALESCE(salesperson_email,''))=lower(%s)
            """,
            (email,),
        )
        order_counts = cur.fetchone()
        cur.execute(
            """
            SELECT id,order_number,order_date,customer_name,status
            FROM shop_orders
            WHERE lower(COALESCE(salesperson_email,''))=lower(%s)
            ORDER BY updated_at DESC,id DESC LIMIT 10
            """,
            (email,),
        )
        my_orders = cur.fetchall()
        cur.execute(
            """
            SELECT id,name,qty,unit,warehouse
            FROM products WHERE qty<=5 ORDER BY qty,lower(name) LIMIT 10
            """
        )
        low_stock = cur.fetchall()
        cur.execute(
            """
            SELECT id,name,qty,unit,warehouse
            FROM products ORDER BY created_at DESC,id DESC LIMIT 10
            """
        )
        recent_products = cur.fetchall()
        cur.execute(
            """
            SELECT type,message,created_at,order_id
            FROM shop_notifications
            WHERE resolved=FALSE AND lower(COALESCE(recipient_email,''))=lower(%s)
            ORDER BY created_at DESC,id DESC LIMIT 12
            """,
            (email,),
        )
        notifications = cur.fetchall()
        conn.close()
        return render_template(
            "sales_dashboard.html",
            order_counts=order_counts,
            my_orders=my_orders,
            low_stock=low_stock,
            recent_products=recent_products,
            notifications=notifications,
        )
    cur.execute("SELECT COUNT(*), COALESCE(SUM(qty), 0) FROM products")
    total_products, total_qty = cur.fetchone()
    cur.execute("SELECT name, qty FROM products ORDER BY qty DESC LIMIT 10")
    top_products = cur.fetchall()
    conn.close()
    return render_template(
        "dashboard.html",
        total_products=total_products,
        total_qty=round(total_qty or 0, 3),
        top_products=top_products,
        max_top_qty=max((row[1] or 0 for row in top_products), default=0),
    )


@app.route('/magazyny')
@login_required
@cache.cached(timeout=30)
def magazyny():
    return render_template("magazyny.html", warehouses=WAREHOUSES)


@app.route('/magazyn')
@login_required
def magazyn_redirect():
    return redirect('/magazyn/Wszystko')


@app.route('/users')
@admin_required
def users():
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, firebase_uid, first_name, last_name, email, role, status,
               can_import_warehouse, can_import_accounting,phone,
               must_change_password,last_login_at,password_changed_at,created_by,
               password_reset_by,failed_login_attempts,locked_until
        FROM users ORDER BY lower(last_name), lower(first_name), lower(email)
        """
    )
    users_list = cur.fetchall()
    conn.close()
    return render_template("users.html", users=users_list, role_labels=SHOP_ROLE_LABELS)


@app.route('/admin/security-history')
@admin_required
def security_history():
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT created_at,user_email,event_type,actor_email,details
        FROM user_security_events ORDER BY created_at DESC,id DESC LIMIT 500
        """
    )
    events = cur.fetchall()
    conn.close()
    return render_template("security_history.html", events=events)


@app.route('/add_user', methods=['POST'])
@admin_required
def add_user():
    email = (request.form.get("email") or "").strip().lower()
    first_name = (request.form.get("first_name") or "").strip()
    last_name = (request.form.get("last_name") or "").strip()
    phone = (request.form.get("phone") or "").strip()
    temporary_password = request.form.get("temporary_password") or ""
    account_active = request.form.get("active") == "on"
    if not email or "@" not in email or len(email) > 254:
        return "Podaj prawidłowy adres e-mail.", 400
    if not first_name or not last_name or len(first_name) > 100 or len(last_name) > 100:
        return "Imię i nazwisko są wymagane (maksymalnie 100 znaków).", 400
    if len(phone) > 50:
        return "Numer telefonu może mieć maksymalnie 50 znaków.", 400
    validation_error = password_validation_error(temporary_password)
    if validation_error:
        return validation_error, 400
    role = request.form.get("role", "warehouse")
    if role not in ROLES:
        return "Nieprawidłowa rola.", 400
    if not FIREBASE_ADMIN_READY:
        return "Firebase Admin nie jest skonfigurowany.", 503

    created_in_firebase = False
    try:
        try:
            auth.get_user_by_email(email)
            return "Konto o tym adresie e-mail już istnieje.", 409
        except auth.UserNotFoundError:
            firebase_user = auth.create_user(
                email=email,
                display_name=f"{first_name} {last_name}",
                password=temporary_password,
                disabled=not account_active,
            )
            created_in_firebase = True
    except Exception:
        logger.exception("Firebase user creation failed.")
        return "Nie udało się utworzyć użytkownika w Firebase.", 502

    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            INSERT INTO users(
                firebase_uid,first_name,last_name,email,phone,role,status,
                must_change_password,created_by
            )
            VALUES (%s,%s,%s,%s,%s,%s,%s,TRUE,%s)
            ON CONFLICT (email) DO UPDATE SET
                firebase_uid=EXCLUDED.firebase_uid,
                first_name=EXCLUDED.first_name,
                last_name=EXCLUDED.last_name,
                phone=EXCLUDED.phone,
                role=EXCLUDED.role,
                status=EXCLUDED.status,
                must_change_password=TRUE,
                failed_login_attempts=0,
                locked_until=NULL,
                created_by=COALESCE(users.created_by,EXCLUDED.created_by),
                updated_at=NOW()
            RETURNING id
            """,
            (
                firebase_user.uid,
                first_name,
                last_name,
                email,
                phone or None,
                role,
                "active" if account_active else "inactive",
                session.get("user"),
            ),
        )
        user_id = cur.fetchone()[0]
        security_event(
            cur,
            user_id,
            email,
            "account_created",
            session.get("user"),
            {"role": role, "active": account_active, "temporary_password": True},
        )
        conn.commit()
    except Exception:
        conn.rollback()
        if created_in_firebase:
            try:
                auth.delete_user(firebase_user.uid)
            except Exception:
                logger.exception("Firebase compensation delete failed.")
        logger.exception("Application user creation failed.")
        return "Nie udało się zapisać użytkownika.", 500
    finally:
        conn.close()
    return redirect('/users')


@app.route('/delete_user/<int:user_id>', methods=['POST'])
@admin_required
def delete_user(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        "SELECT firebase_uid, email, role, status FROM users WHERE id=%s FOR UPDATE",
        (user_id,),
    )
    user = cur.fetchone()
    if not user:
        conn.close()
        return redirect('/users')
    if user[0] == session.get("uid"):
        conn.close()
        return "Nie możesz usunąć aktualnie zalogowanego konta.", 400
    if user[2] == "admin" and user[3] == "active":
        cur.execute("SELECT COUNT(*) FROM users WHERE role='admin' AND status='active'")
        if cur.fetchone()[0] <= 1:
            conn.close()
            return "Nie można usunąć ostatniego administratora.", 400
    try:
        auth.delete_user(user[0])
        cur.execute("DELETE FROM users WHERE id=%s", (user_id,))
        conn.commit()
    except auth.UserNotFoundError:
        cur.execute("DELETE FROM users WHERE id=%s", (user_id,))
        conn.commit()
    except Exception:
        conn.rollback()
        logger.exception("User deletion failed.")
        return "Nie udało się usunąć użytkownika.", 502
    finally:
        conn.close()
    return redirect('/users')


@app.route('/update_user_role/<int:user_id>', methods=['POST'])
@admin_required
def update_user_role(user_id):
    new_role = request.form.get('role', 'warehouse')
    if new_role not in ROLES:
        return redirect('/users')

    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT role, status FROM users WHERE id=%s FOR UPDATE", (user_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return redirect('/users')

    old_role, status = row
    if old_role == 'admin' and new_role != 'admin' and status == "active":
        cur.execute("SELECT COUNT(*) FROM users WHERE role='admin' AND status='active'")
        admins_count = cur.fetchone()[0]
        if admins_count <= 1:
            conn.close()
            return redirect('/users')

    cur.execute(
        """
        UPDATE users SET role=%s,
            can_import_warehouse=CASE WHEN %s='warehouse' THEN can_import_warehouse ELSE FALSE END,
            can_import_accounting=CASE WHEN %s='accounting' THEN can_import_accounting ELSE FALSE END,
            updated_at=NOW()
        WHERE id=%s
        """,
        (new_role, new_role, new_role, user_id),
    )
    conn.commit()
    conn.close()
    return redirect('/users')


@app.route('/users/<int:user_id>/import-permissions', methods=['POST'])
@admin_required
def update_user_import_permissions(user_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT role FROM users WHERE id=%s FOR UPDATE", (user_id,))
        user = cur.fetchone()
        if not user:
            return "Nie znaleziono użytkownika.", 404
        warehouse_allowed = (
            user[0] == "warehouse"
            and request.form.get("can_import_warehouse") == "on"
        )
        accounting_allowed = (
            user[0] == "accounting"
            and request.form.get("can_import_accounting") == "on"
        )
        cur.execute(
            """
            UPDATE users
            SET can_import_warehouse=%s, can_import_accounting=%s, updated_at=NOW()
            WHERE id=%s
            """,
            (warehouse_allowed, accounting_allowed, user_id),
        )
        conn.commit()
        return redirect("/users")
    except Exception:
        conn.rollback()
        logger.exception("Import permissions update failed.")
        return "Nie udało się zapisać uprawnień importu.", 500
    finally:
        conn.close()


def send_firebase_password_reset(email):
    endpoint = (
        "https://identitytoolkit.googleapis.com/v1/accounts:sendOobCode"
        f"?key={FIREBASE_CONFIG['apiKey']}"
    )
    payload = json.dumps({"requestType": "PASSWORD_RESET", "email": email}).encode("utf-8")
    req = urlrequest.Request(
        endpoint,
        data=payload,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urlrequest.urlopen(req, timeout=10) as response:
        return response.status == 200


@app.route('/users/<int:user_id>/reset-password', methods=['POST'])
@admin_required
@limiter.limit("10 per hour")
def reset_user_password(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT email FROM users WHERE id=%s FOR UPDATE", (user_id,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return "Nie znaleziono użytkownika.", 404
    try:
        send_firebase_password_reset(row[0])
        cur.execute(
            """
            UPDATE users SET password_reset_by=%s,updated_at=NOW() WHERE id=%s
            """,
            (session.get("user"), user_id),
        )
        security_event(
            cur, user_id, row[0], "password_reset_link_sent", session.get("user")
        )
        conn.commit()
    except Exception:
        conn.rollback()
        logger.exception("Password reset email failed.")
        return "Nie udało się wysłać wiadomości resetującej hasło.", 502
    finally:
        conn.close()
    return redirect("/users?reset=sent")


@app.route('/users/<int:user_id>/temporary-password', methods=['POST'])
@admin_required
@limiter.limit("20 per hour")
def set_temporary_password(user_id):
    temporary_password = request.form.get("temporary_password") or ""
    validation_error = password_validation_error(temporary_password)
    if validation_error:
        return validation_error, 400
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT firebase_uid,email FROM users WHERE id=%s FOR UPDATE",
            (user_id,),
        )
        account = cur.fetchone()
        if not account:
            return "Nie znaleziono użytkownika.", 404
        auth.update_user(account[0], password=temporary_password)
        auth.revoke_refresh_tokens(account[0])
        cur.execute(
            """
            UPDATE users SET must_change_password=TRUE,password_reset_by=%s,
                failed_login_attempts=0,locked_until=NULL,updated_at=NOW()
            WHERE id=%s
            """,
            (session.get("user"), user_id),
        )
        security_event(
            cur,
            user_id,
            account[1],
            "temporary_password_set",
            session.get("user"),
        )
        conn.commit()
        return redirect("/users?temporary=saved")
    except Exception:
        conn.rollback()
        logger.exception("Temporary password reset failed.")
        return "Nie udało się ustawić hasła tymczasowego.", 502
    finally:
        conn.close()


@app.route('/users/<int:user_id>/force-password-change', methods=['POST'])
@admin_required
def force_password_change(user_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT firebase_uid,email FROM users WHERE id=%s FOR UPDATE",
            (user_id,),
        )
        account = cur.fetchone()
        if not account:
            return "Nie znaleziono użytkownika.", 404
        auth.revoke_refresh_tokens(account[0])
        cur.execute(
            "UPDATE users SET must_change_password=TRUE,updated_at=NOW() WHERE id=%s",
            (user_id,),
        )
        security_event(
            cur,
            user_id,
            account[1],
            "password_change_forced",
            session.get("user"),
        )
        conn.commit()
        return redirect("/users?force=saved")
    except Exception:
        conn.rollback()
        logger.exception("Forced password change failed.")
        return "Nie udało się wymusić zmiany hasła.", 502
    finally:
        conn.close()


@app.route('/users/<int:user_id>/unlock', methods=['POST'])
@admin_required
def unlock_user(user_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT firebase_uid,email,status FROM users WHERE id=%s FOR UPDATE",
            (user_id,),
        )
        account = cur.fetchone()
        if not account:
            return "Nie znaleziono użytkownika.", 404
        if account[2] == "inactive":
            return "Najpierw aktywuj konto.", 409
        auth.update_user(account[0], disabled=False)
        cur.execute(
            """
            UPDATE users SET status='active',failed_login_attempts=0,
                locked_until=NULL,updated_at=NOW() WHERE id=%s
            """,
            (user_id,),
        )
        security_event(
            cur, user_id, account[1], "account_unlocked", session.get("user")
        )
        conn.commit()
        return redirect("/users?unlock=saved")
    except Exception:
        conn.rollback()
        logger.exception("Account unlock failed.")
        return "Nie udało się odblokować konta.", 502
    finally:
        conn.close()


@app.route('/users/<int:user_id>/toggle-active', methods=['POST'])
@admin_required
def toggle_user_active(user_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT firebase_uid,email,role,status FROM users WHERE id=%s FOR UPDATE",
            (user_id,),
        )
        account = cur.fetchone()
        if not account:
            return "Nie znaleziono użytkownika.", 404
        uid, email, role, status = account
        activating = status == "inactive"
        if not activating and uid == session.get("uid"):
            return "Nie możesz dezaktywować własnego konta.", 400
        if not activating and role == "admin" and status == "active":
            cur.execute("SELECT COUNT(*) FROM users WHERE role='admin' AND status='active'")
            if cur.fetchone()[0] <= 1:
                return "Nie można dezaktywować ostatniego administratora.", 400
        target_status = "active" if activating else "inactive"
        auth.update_user(uid, disabled=not activating)
        cur.execute(
            """
            UPDATE users SET status=%s,failed_login_attempts=0,locked_until=NULL,
                updated_at=NOW() WHERE id=%s
            """,
            (target_status, user_id),
        )
        security_event(
            cur,
            user_id,
            email,
            "account_activated" if activating else "account_deactivated",
            session.get("user"),
        )
        conn.commit()
        return redirect("/users?active=saved")
    except Exception:
        conn.rollback()
        logger.exception("Account activation change failed.")
        return "Nie udało się zmienić aktywności konta.", 502
    finally:
        conn.close()


@app.route('/users/<int:user_id>/toggle-block', methods=['POST'])
@admin_required
def toggle_user_block(user_id):
    conn = db()
    cur = conn.cursor()
    cur.execute(
        "SELECT firebase_uid, role, status FROM users WHERE id=%s FOR UPDATE",
        (user_id,),
    )
    row = cur.fetchone()
    if not row:
        conn.close()
        return "Nie znaleziono użytkownika.", 404
    uid, role, status = row
    if status == "inactive":
        conn.close()
        return "Nieaktywne konto należy najpierw aktywować.", 409
    if uid == session.get("uid"):
        conn.close()
        return "Nie możesz zablokować własnego konta.", 400
    target_status = "active" if status == "blocked" else "blocked"
    if role == "admin" and status == "active" and target_status == "blocked":
        cur.execute("SELECT COUNT(*) FROM users WHERE role='admin' AND status='active'")
        if cur.fetchone()[0] <= 1:
            conn.close()
            return "Nie można zablokować ostatniego administratora.", 400
    try:
        auth.update_user(uid, disabled=(target_status == "blocked"))
        cur.execute(
            """
            UPDATE users SET status=%s,failed_login_attempts=0,locked_until=NULL,
                updated_at=NOW() WHERE id=%s
            """,
            (target_status, user_id),
        )
        cur.execute("SELECT email FROM users WHERE id=%s", (user_id,))
        email_row = cur.fetchone()
        security_event(
            cur,
            user_id,
            email_row[0] if email_row else "",
            "account_unblocked" if target_status == "active" else "account_blocked",
            session.get("user"),
        )
        conn.commit()
    except Exception:
        conn.rollback()
        try:
            auth.update_user(uid, disabled=(status == "blocked"))
        except Exception:
            logger.exception("Firebase block-state compensation failed.")
        logger.exception("User block toggle failed.")
        return "Nie udało się zmienić statusu konta.", 502
    finally:
        conn.close()
    return redirect("/users")


@app.route('/magazyn/<name>')
@login_required
def magazyn(name):
    if name != "Wszystko" and name not in WAREHOUSES:
        return "Nie znaleziono magazynu.", 404
    conn = db()
    cur = conn.cursor()

    search_query = (request.args.get("q") or "").strip()
    if name == "Wszystko" and search_query:
        like = f"%{search_query.lower()}%"
        cur.execute(
            """
            SELECT DISTINCT p.* FROM products p
            LEFT JOIN packages pk ON pk.product_id=p.id
            WHERE lower(p.name) LIKE %s OR lower(COALESCE(pk.number,'')) LIKE %s
            ORDER BY p.warehouse,lower(p.name),p.id
            """,
            (like, like),
        )
    elif name == "Wszystko":
        cur.execute("SELECT * FROM products ORDER BY warehouse, lower(name), id")
    elif search_query:
        like = f"%{search_query.lower()}%"
        cur.execute(
            """
            SELECT DISTINCT p.* FROM products p
            LEFT JOIN packages pk ON pk.product_id=p.id
            WHERE p.warehouse=%s
              AND (lower(p.name) LIKE %s OR lower(COALESCE(pk.number,'')) LIKE %s)
            ORDER BY lower(p.name),p.id
            """,
            (name, like, like),
        )
    else:
        cur.execute("SELECT * FROM products WHERE warehouse=%s ORDER BY lower(name), id", (name,))

    products = cur.fetchall()
    product_ids = [row[0] for row in products]
    package_modes = {}
    reservations = {}
    if product_ids:
        cur.execute(
            """
            SELECT p.id,
                   COALESCE(SUM(CASE WHEN pk.status='active' AND pk.qty>0 THEN pk.qty ELSE 0 END),0) AS numbered_qty
            FROM products p
            LEFT JOIN packages pk ON pk.product_id=p.id AND pk.warehouse=p.warehouse
            WHERE p.id = ANY(%s)
            GROUP BY p.id, p.qty
            """,
            (product_ids,),
        )
        for product_id, numbered_qty in cur.fetchall():
            product = next((row for row in products if row[0] == product_id), None)
            total_qty = product[2] if product else 0
            has_numbered = (numbered_qty or 0) > 0
            has_unnumbered = (total_qty or 0) - (numbered_qty or 0) > 1e-9
            package_modes[product_id] = (
                "mixed" if has_numbered and has_unnumbered else
                "numbered" if has_numbered else
                "unnumbered" if has_unnumbered else
                "empty"
            )
        cur.execute(
            """
            SELECT product_id, SUM(qty) FROM (
                SELECT i.product_id, COALESCE(SUM(i.reserved_qty-i.issued_qty),0) AS qty
                FROM shop_order_items i
                JOIN shop_orders o ON o.id=i.order_id
                WHERE i.product_id=ANY(%s)
                  AND o.status NOT IN ('Zakończone','Anulowane')
                GROUP BY i.product_id
                UNION ALL
                SELECT ri.product_id, COALESCE(SUM(ri.qty),0) AS qty
                FROM reservation_items ri
                JOIN reservations r ON r.id=ri.reservation_id
                WHERE ri.product_id=ANY(%s)
                  AND r.status = ANY(%s)
                GROUP BY ri.product_id
            ) reserved
            GROUP BY product_id
            """,
            (product_ids, product_ids, active_reservation_statuses_sql()),
        )
        reservations = {row[0]: max(float(row[1] or 0), 0) for row in cur.fetchall()}
    conn.close()

    return render_template(
        "index.html",
        products=products,
        warehouse=name,
        package_modes=package_modes,
        warehouses=WAREHOUSES,
        units=sorted(UNITS),
        search_query=search_query,
        reservations=reservations,
    )


@app.route('/packages/<int:product_id>')
@login_required
def packages_for_product(product_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT qty FROM products WHERE id=%s", (product_id,))
    product = cur.fetchone()
    cur.execute(
        """
        SELECT id, number, qty, status, warehouse
        FROM packages
        WHERE product_id=%s
        ORDER BY CASE WHEN status='active' THEN 0 ELSE 1 END, id DESC
        """,
        (product_id,),
    )
    packages = cur.fetchall()
    numbered_qty = sum((row[2] or 0) for row in packages if row[3] == 'active')
    unnumbered_qty = max((product[0] if product else 0) - numbered_qty, 0)
    conn.close()
    return render_template("packages.html", packages=packages, unnumbered_qty=unnumbered_qty)


@app.route('/api/packages/lookup')
@login_required
def package_lookup():
    number = (request.args.get("number") or "").strip()
    warehouse = (request.args.get("warehouse") or "").strip()
    if not number:
        return jsonify({"ok": False, "error": "Podaj numer paczki."}), 400
    conn = db()
    cur = conn.cursor()
    query = """
        SELECT pk.id, pk.number, pk.qty, pk.warehouse, pk.status,
               p.id, p.name, p.unit, p.price_netto, p.vat
        FROM packages pk
        JOIN products p ON p.id=pk.product_id
        WHERE lower(pk.number)=lower(%s) AND pk.status='active' AND pk.qty>0
    """
    params = [number]
    if warehouse:
        query += " AND pk.warehouse=%s"
        params.append(warehouse)
    query += " ORDER BY CASE WHEN pk.status='active' THEN 0 ELSE 1 END, pk.id DESC LIMIT 2"
    cur.execute(query, tuple(params))
    rows = cur.fetchall()
    conn.close()
    if not rows:
        return jsonify({"ok": False, "error": "Nie znaleziono paczki."}), 404
    if len(rows) > 1 and not warehouse:
        return jsonify({"ok": False, "error": "Numer występuje w kilku magazynach. Wybierz magazyn."}), 409
    row = rows[0]
    return jsonify({
        "ok": True,
        "package": {
            "id": row[0], "number": row[1], "qty": row[2], "warehouse": row[3],
            "status": row[4], "product_id": row[5], "product": row[6],
            "unit": row[7], "price_netto": row[8], "vat": row[9],
        },
    })


@app.route('/add_product', methods=['POST'])
@permission_required("inventory_manage")
def add_product():
    name = (request.form.get("name") or "").strip()
    unit = (request.form.get("unit") or "").strip()
    warehouse = (request.form.get("warehouse") or "").strip()
    if not name or len(name) > 200:
        return "Nazwa produktu jest wymagana i może mieć maksymalnie 200 znaków.", 400
    if unit not in UNITS:
        return "Wybierz prawidłową jednostkę.", 400
    if warehouse not in WAREHOUSES:
        return "Wybierz prawidłowy magazyn.", 400
    try:
        price_netto = parse_nonnegative_number(request.form.get("price_netto"), "Cena netto")
        vat = parse_nonnegative_number(request.form.get("vat"), "VAT")
    except ValueError as exc:
        return str(exc), 400
    if vat > 100:
        return "VAT nie może przekraczać 100%.", 400

    conn = db()
    cur = conn.cursor()
    try:
        lock_key = f"product:{warehouse}:{name.casefold()}"
        cur.execute("SELECT pg_advisory_xact_lock(hashtext(%s))", (lock_key,))
        cur.execute(
            "SELECT id FROM products WHERE warehouse=%s AND lower(name)=lower(%s) LIMIT 1",
            (warehouse, name),
        )
        if cur.fetchone():
            raise ValueError("Taki produkt już istnieje w tym magazynie.")
        cur.execute(
            """
            INSERT INTO products(name, qty, unit, warehouse, price_netto, vat)
            VALUES (%s,0,%s,%s,%s,%s)
            """,
            (name, unit, warehouse, price_netto, vat),
        )
        conn.commit()
        cache.clear()
        return redirect(f"/magazyn/{warehouse}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 409
    except Exception:
        conn.rollback()
        logger.exception("Product creation failed.")
        return "Nie udało się dodać produktu.", 500
    finally:
        conn.close()


@app.route('/delete_product/<int:product_id>', methods=['POST'])
@admin_required
def delete_product(product_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM issue_items WHERE product_id=%s", (product_id,))
    linked = cur.fetchone()[0]
    if linked > 0:
        conn.close()
        return "Nie można usunąć produktu użytego w dokumentach.", 400

    cur.execute("DELETE FROM packages WHERE product_id=%s", (product_id,))
    cur.execute("DELETE FROM products WHERE id=%s", (product_id,))
    conn.commit()
    conn.close()
    return ("", 204)


@app.route('/delete_selected', methods=['POST'])
@admin_required
def delete_selected():
    selected = request.form.getlist("selected")
    if not selected:
        return redirect(request.referrer or '/magazyny')

    conn = db()
    cur = conn.cursor()
    for pid_raw in selected:
        try:
            pid = int(pid_raw)
        except ValueError:
            continue

        cur.execute("SELECT COUNT(*) FROM issue_items WHERE product_id=%s", (pid,))
        linked = cur.fetchone()[0]
        if linked > 0:
            continue

        cur.execute("DELETE FROM packages WHERE product_id=%s", (pid,))
        cur.execute("DELETE FROM products WHERE id=%s", (pid,))

    conn.commit()
    conn.close()
    return redirect(request.referrer or '/magazyny')


@app.route('/costs')
@login_required
def costs():
    conn = db()
    cur = conn.cursor()
    cur.execute(
        "SELECT id, name, amount, date, warehouse_source, description, source_warehouse "
        "FROM costs ORDER BY date DESC, id DESC"
    )
    costs_rows = cur.fetchall()
    conn.close()
    return render_template("costs.html", costs=costs_rows, warehouses=WAREHOUSES)


def process_cost():
    name = (request.form.get("name") or "").strip()
    description = (request.form.get("description") or "").strip()
    source_warehouse = (request.form.get("source_warehouse") or "").strip()
    warehouse_source = request.form.get("warehouse_source") == "on"
    try:
        amount = parse_positive_number(request.form.get("amount"), "Kwota / ilość")
        date = normalized_document_date(request.form.get("date"))
    except ValueError as exc:
        return str(exc), 400
    if not name:
        return "Nazwa jest wymagana.", 400
    if warehouse_source and source_warehouse not in WAREHOUSES:
        return "Wybierz prawidłowy magazyn źródłowy.", 400

    conn = db()
    cur = conn.cursor()
    try:
        if warehouse_source:
            cur.execute(
                """
                SELECT id, qty FROM products
                WHERE lower(name)=lower(%s) AND warehouse=%s
                ORDER BY id LIMIT 1 FOR UPDATE
                """,
                (name, source_warehouse),
            )
            product = cur.fetchone()
            if not product:
                raise ValueError(f"Brak produktu '{name}' w magazynie {source_warehouse}.")
            cur.execute(
                "SELECT 1 FROM packages WHERE product_id=%s AND status='active' AND qty>0 LIMIT 1",
                (product[0],),
            )
            if cur.fetchone():
                raise ValueError("Ten produkt jest ewidencjonowany w paczkach. Użyj formularza wydania.")
            cur.execute(
                "UPDATE products SET qty=qty-%s WHERE id=%s AND qty >= %s",
                (amount, product[0], amount),
            )
            if cur.rowcount != 1:
                raise ValueError(f"Brak wystarczającej ilości w magazynie {source_warehouse}.")
        cur.execute(
            """
            INSERT INTO costs(name, amount, date, warehouse_source, description, source_warehouse)
            VALUES (%s,%s,%s,%s,%s,%s)
            """,
            (name, amount, date, warehouse_source, description, source_warehouse or None),
        )
        conn.commit()
        cache.clear()
        return redirect("/costs")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Cost creation failed.")
        return "Nie udało się zapisać kosztu.", 500
    finally:
        conn.close()


@app.route('/add_cost', methods=['POST'])
@login_required
def add_cost():
    return process_cost()

    name = (request.form.get('name') or '').strip()
    description = (request.form.get('description') or '').strip()
    date = request.form.get('date') or datetime.now().strftime("%Y-%m-%d")
    warehouse_source = request.form.get('warehouse_source') == 'on'
    source_warehouse = (request.form.get('source_warehouse') or '').strip()
    try:
        amount = float((request.form.get('amount') or '0').replace(",", "."))
    except Exception:
        return "Nieprawidłowa kwota/ilość.", 400

    if not name or amount <= 0:
        return "Nazwa i kwota/ilość są wymagane.", 400

    conn = db()
    cur = conn.cursor()
    try:
        if warehouse_source:
            if source_warehouse not in WAREHOUSES:
                return "Wybierz prawidłowy magazyn źródłowy.", 400
            cur.execute(
                "SELECT id, qty FROM products WHERE name=%s AND warehouse=%s ORDER BY id LIMIT 1 FOR UPDATE",
                (name, source_warehouse)
            )
            product = cur.fetchone()
            if not product:
                conn.close()
                return f"Brak produktu '{name}' w magazynie głównym ({MAIN_WAREHOUSE}).", 400
            if product[1] < amount:
                conn.close()
                return "Brak wystarczającej ilości w magazynie głównym.", 400

            cur.execute(
                "SELECT 1 FROM packages WHERE product_id=%s AND status='active' AND qty>0 LIMIT 1",
                (product[0],)
            )
            if cur.fetchone():
                return "Ten produkt jest ewidencjonowany w paczkach. Użyj formularza wydania.", 400
            cur.execute(
                "UPDATE products SET qty=qty-%s WHERE id=%s AND qty >= %s",
                (amount, product[0], amount)
            )

        cur.execute(
            """
            INSERT INTO costs(name, amount, date, warehouse_source, description, source_warehouse)
            VALUES (%s,%s,%s,%s,%s,%s)
            """,
            (name, amount, date, warehouse_source, description, source_warehouse or None)
        )
        conn.commit()
        cache.clear()
    except Exception:
        conn.rollback()
        return "Błąd podczas zapisu kosztu.", 400
    finally:
        conn.close()

    return redirect('/costs')


@app.route('/shop/orders')
@login_required
def shop_orders_page():
    return redirect("/sklep")


@app.route('/shop/orders/add', methods=['POST'])
@login_required
def add_shop_order():
    return (
        "Ten formularz został wycofany po ujednoliceniu modułu sklepu. "
        "Otwórz moduł Sklep internetowy i ponów operację.",
        409,
    )


def form_value(values, index, default=""):
    return values[index] if index < len(values) else default


def normalized_document_date(value):
    value = (value or "").strip() or datetime.now().strftime("%Y-%m-%d")
    try:
        datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        raise ValueError("Nieprawidłowa data dokumentu.")
    return value


def normalized_optional_date(value, field_name):
    value = (value or "").strip()
    if not value:
        return ""
    try:
        datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        raise ValueError(f"{field_name}: nieprawidłowa data.")
    return value


def submitted_receipt_rows(forced_warehouse=None):
    fields = {
        "product_id": request.form.getlist("product_id"),
        "product_name": request.form.getlist("product_name"),
        "qty": request.form.getlist("qty"),
        "unit": request.form.getlist("unit"),
        "warehouse": request.form.getlist("warehouse"),
        "package_number": request.form.getlist("package_number"),
        "has_package_number": request.form.getlist("has_package_number"),
        "price_netto": request.form.getlist("price_netto"),
        "price_brutto": request.form.getlist("price_brutto"),
    }
    row_count = max((len(values) for values in fields.values()), default=0)
    rows = []
    for index in range(row_count):
        row = {
            field: form_value(values, index)
            for field, values in fields.items()
        }
        if forced_warehouse:
            row["warehouse"] = forced_warehouse
        row["has_package_number"] = row["has_package_number"] == "1"
        rows.append(row)
    return rows


def render_receipt_form(error=None, forced_warehouse=None, status=200):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT * FROM products ORDER BY warehouse, lower(name), id")
        products = cur.fetchall()
    finally:
        conn.close()
    return (
        render_template(
            "przyjecie.html",
            products=products,
            forced_warehouse=forced_warehouse,
            form_action=(
                "/inwestycja-suwaj/receive_doc"
                if forced_warehouse == INVESTMENT_WAREHOUSE else "/receive_doc"
            ),
            page_title=(
                "Przyjęcie (PZ) – Inwestycja Suwaj"
                if forced_warehouse == INVESTMENT_WAREHOUSE else "Przyjęcie (PZ)"
            ),
            form_error=error,
            submitted_items=submitted_receipt_rows(forced_warehouse) if request.method == "POST" else [],
            submitted_header={
                "doc_number": request.form.get("doc_number", ""),
                "date": request.form.get("date", ""),
                "kontrahent": request.form.get("kontrahent", ""),
            } if request.method == "POST" else {},
        ),
        status,
    )


def collect_document_items(forced_warehouse=None, issuing=False):
    product_ids = request.form.getlist("product_id")
    product_names = request.form.getlist("product_name")
    quantities = request.form.getlist("qty")
    unit_values = request.form.getlist("unit")
    warehouses = request.form.getlist("warehouse")
    package_values = request.form.getlist("package_id" if issuing else "package_number")
    has_package_values = request.form.getlist("has_package_number")
    netto_values = request.form.getlist("price_netto")
    brutto_values = request.form.getlist("price_brutto")
    items = []
    row_count = max(
        len(product_ids), len(product_names), len(quantities), len(unit_values),
        len(warehouses), len(package_values), len(has_package_values),
        len(netto_values), len(brutto_values),
    )
    for index in range(row_count):
        raw_product_id = form_value(product_ids, index)
        raw_product_id = (raw_product_id or "").strip()
        package_value = form_value(package_values, index).strip()
        marker = form_value(has_package_values, index)
        row_started = any(
            str(value or "").strip()
            for value in (
                form_value(product_names, index),
                form_value(quantities, index),
                package_value,
                form_value(netto_values, index),
                form_value(brutto_values, index),
            )
        ) or marker == "1"
        if not raw_product_id:
            if row_started:
                raise ValueError(
                    f"Pozycja {index + 1}: wybierz produkt z listy podpowiedzi."
                )
            continue
        try:
            product_id = int(raw_product_id)
        except ValueError:
            raise ValueError("Wybrano nieprawidłowy produkt.")
        warehouse = forced_warehouse or form_value(warehouses, index).strip()
        if warehouse not in WAREHOUSES:
            raise ValueError("Wybierz prawidłowy magazyn dla każdej pozycji.")
        unit = form_value(unit_values, index).strip()
        if unit and unit not in UNITS:
            raise ValueError(f"Pozycja {index + 1}: wybierz prawidłową jednostkę.")
        qty = parse_positive_number(form_value(quantities, index))
        price_netto = parse_nonnegative_number(form_value(netto_values, index), "Cena netto")
        price_brutto = parse_nonnegative_number(form_value(brutto_values, index), "Cena brutto")
        has_package_number = marker == "1" or (not marker and bool(package_value))
        if not issuing and len(package_value) > 100:
            raise ValueError("Numer paczki może mieć maksymalnie 100 znaków.")
        if not issuing and has_package_number and not package_value:
            raise ValueError("Zaznaczono, że towar posiada numer paczki — wpisz numer paczki.")
        if not issuing and not has_package_number:
            package_value = ""
        items.append({
            "product_id": product_id,
            "warehouse": warehouse,
            "unit": unit,
            "qty": qty,
            "package": package_value,
            "has_package_number": has_package_number,
            "price_netto": price_netto,
            "price_brutto": price_brutto,
        })
    if not items:
        raise ValueError("Dodaj co najmniej jedną kompletną pozycję dokumentu.")
    return items


def create_receipt(forced_warehouse=None):
    try:
        items = collect_document_items(forced_warehouse=forced_warehouse, issuing=False)
        items.sort(key=lambda item: (item["warehouse"], item["product_id"], item["package"].casefold()))
        date = normalized_document_date(request.form.get("date"))
    except ValueError as exc:
        return render_receipt_form(str(exc), forced_warehouse, 400)
    contractor = (request.form.get("kontrahent") or "").strip()
    if not contractor:
        return render_receipt_form("Dostawca jest wymagany.", forced_warehouse, 400)
    if len(contractor) > 200:
        return render_receipt_form(
            "Nazwa dostawcy może mieć maksymalnie 200 znaków.",
            forced_warehouse,
            400,
        )

    conn = db()
    cur = conn.cursor()
    try:
        resolved = []
        request_packages = set()
        for item in items:
            cur.execute(
                "SELECT name, unit, price_netto, vat FROM products WHERE id=%s",
                (item["product_id"],),
            )
            source = cur.fetchone()
            if not source:
                raise ValueError("Wybrany produkt już nie istnieje.")
            cur.execute(
                "SELECT pg_advisory_xact_lock(hashtext(%s))",
                (f"product:{item['warehouse']}:{source[0].casefold()}",),
            )
            cur.execute(
                """
                SELECT id FROM products
                WHERE warehouse=%s AND lower(name)=lower(%s)
                ORDER BY id LIMIT 1 FOR UPDATE
                """,
                (item["warehouse"], source[0]),
            )
            target = cur.fetchone()
            target_id = target[0] if target else None
            if item["package"]:
                package_key = (item["warehouse"].casefold(), item["package"].casefold())
                if package_key in request_packages:
                    raise ValueError(
                        f"Paczka {item['package']} została podana więcej niż raz."
                    )
                request_packages.add(package_key)
                cur.execute(
                    "SELECT pg_advisory_xact_lock(hashtext(%s))",
                    (f"package:{item['warehouse']}:{item['package'].casefold()}",),
                )
                cur.execute(
                    """
                    SELECT 1 FROM packages
                    WHERE warehouse=%s AND lower(number)=lower(%s)
                    LIMIT 1 FOR UPDATE
                    """,
                    (item["warehouse"], item["package"]),
                )
                if cur.fetchone():
                    raise ValueError(
                        f"Paczka {item['package']} już istnieje w magazynie {item['warehouse']}."
                    )
            resolved.append((item, source, target_id))

        cur.execute(
            """
            INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number, movement_type)
            VALUES (%s,%s,%s,%s,%s,'PZ') RETURNING id
            """,
            (date, contractor, forced_warehouse or "", "", ""),
        )
        doc_id = cur.fetchone()[0]
        requested_number = (request.form.get("doc_number") or "").strip()
        if len(requested_number) > 100:
            raise ValueError("Numer dokumentu może mieć maksymalnie 100 znaków.")
        if requested_number:
            cur.execute(
                "SELECT pg_advisory_xact_lock(hashtext(%s))",
                (f"document:{requested_number.casefold()}",),
            )
            cur.execute(
                "SELECT 1 FROM issue_docs WHERE id<>%s AND lower(doc_number)=lower(%s) LIMIT 1",
                (doc_id, requested_number),
            )
            if cur.fetchone():
                raise ValueError("Dokument o takim numerze już istnieje.")
        prefix = "PZ-IS" if forced_warehouse == INVESTMENT_WAREHOUSE else "PZ"
        doc_number = requested_number or f"{prefix}/{doc_id}/{datetime.now().year}"
        cur.execute("UPDATE issue_docs SET doc_number=%s WHERE id=%s", (doc_number, doc_id))

        for item, source, target_id in resolved:
            item_unit = item["unit"] or source[1]
            if target_id:
                cur.execute(
                    "UPDATE products SET qty=qty+%s, unit=%s, price_netto=%s WHERE id=%s",
                    (item["qty"], item_unit, item["price_netto"], target_id),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO products(name, qty, unit, warehouse, price_netto, vat)
                    VALUES (%s,%s,%s,%s,%s,%s) RETURNING id
                    """,
                    (source[0], item["qty"], item_unit, item["warehouse"],
                     item["price_netto"] or source[2] or 0, source[3] or 0),
                )
                target_id = cur.fetchone()[0]

            package_id = None
            if item["package"]:
                cur.execute(
                    """
                    INSERT INTO packages(product_id, number, qty, warehouse, initial_qty, status)
                    VALUES (%s,%s,%s,%s,%s,'active') RETURNING id
                    """,
                    (target_id, item["package"], item["qty"], item["warehouse"], item["qty"]),
                )
                package_id = cur.fetchone()[0]
            cur.execute(
                """
                INSERT INTO issue_items(
                    doc_id, product_id, qty, warehouse, package_id, package_number,
                    price_netto, price_brutto
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (doc_id, target_id, item["qty"], item["warehouse"], package_id,
                 item["package"] or None, item["price_netto"], item["price_brutto"]),
            )
        log_action("document.receipt_created", "issue_doc", doc_id, {"doc_number": doc_number, "items": len(resolved)}, conn)
        conn.commit()
        cache.clear()
        return redirect(f"/doc/{doc_id}")
    except ValueError as exc:
        close_with_rollback(conn)
        return render_receipt_form(str(exc), forced_warehouse, 400)
    except Exception:
        close_with_rollback(conn)
        logger.exception("Receipt creation failed.")
        message = "Nie udało się zapisać przyjęcia. Żadne stany nie zostały zmienione."
        try:
            return render_receipt_form(message, forced_warehouse, 500)
        except Exception:
            logger.exception("Receipt form recovery rendering failed.")
            return render_template(
                "error.html",
                title="Nie udało się zapisać przyjęcia",
                message=message,
            ), 500
    finally:
        conn.close()


def create_issue(forced_warehouse=None):
    try:
        items = collect_document_items(forced_warehouse=forced_warehouse, issuing=True)
        items.sort(
            key=lambda item: (
                item["warehouse"],
                item["product_id"],
                int(item["package"]) if item["package"].isdigit() else 0,
            )
        )
        date = normalized_document_date(request.form.get("date"))
    except ValueError as exc:
        return str(exc), 400
    contractor = (request.form.get("kontrahent") or "").strip()
    if not contractor:
        return "Kontrahent jest wymagany.", 400
    if len(contractor) > 200:
        return "Nazwa kontrahenta może mieć maksymalnie 200 znaków.", 400
    movement_type = (request.form.get("movement_type") or "WZ").strip().upper()
    if movement_type not in {"WZ", "RW"}:
        return "Nieprawidłowy typ dokumentu wydania.", 400

    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number, movement_type)
            VALUES (%s,%s,%s,%s,%s,%s) RETURNING id
            """,
            (date, contractor, forced_warehouse or "", "", "", movement_type),
        )
        doc_id = cur.fetchone()[0]
        requested_number = (request.form.get("doc_number") or "").strip()
        if len(requested_number) > 100:
            raise ValueError("Numer dokumentu może mieć maksymalnie 100 znaków.")
        if requested_number:
            cur.execute(
                "SELECT pg_advisory_xact_lock(hashtext(%s))",
                (f"document:{requested_number.casefold()}",),
            )
            cur.execute(
                "SELECT 1 FROM issue_docs WHERE id<>%s AND lower(doc_number)=lower(%s) LIMIT 1",
                (doc_id, requested_number),
            )
            if cur.fetchone():
                raise ValueError("Dokument o takim numerze już istnieje.")
        prefix = f"{movement_type}-IS" if forced_warehouse == INVESTMENT_WAREHOUSE else movement_type
        doc_number = requested_number or f"{prefix}/{doc_id}/{datetime.now().year}"
        cur.execute("UPDATE issue_docs SET doc_number=%s WHERE id=%s", (doc_number, doc_id))

        for item in items:
            cur.execute(
                """
                SELECT id FROM products
                WHERE id=%s AND warehouse=%s
                FOR UPDATE
                """,
                (item["product_id"], item["warehouse"]),
            )
            if not cur.fetchone():
                raise ValueError("Produkt nie istnieje w wybranym magazynie.")

            package_id = None
            package_number = None
            if item["package"]:
                try:
                    package_id = int(item["package"])
                except ValueError:
                    raise ValueError("Wybrano nieprawidłową paczkę.")
                cur.execute(
                    """
                    SELECT number, qty FROM packages
                    WHERE id=%s AND product_id=%s AND warehouse=%s AND status='active'
                    FOR UPDATE
                    """,
                    (package_id, item["product_id"], item["warehouse"]),
                )
                package = cur.fetchone()
                if not package:
                    raise ValueError("Paczka nie należy do wybranego produktu lub magazynu.")
                package_number = package[0]
                if package[1] + 1e-9 < item["qty"]:
                    raise ValueError(
                        f"W paczce {package_number} jest tylko {package[1]}."
                    )
                cur.execute(
                    """
                    UPDATE packages
                    SET qty=qty-%s,
                        status=CASE WHEN qty-%s <= 0 THEN 'issued' ELSE 'active' END,
                        archived_at=CASE WHEN qty-%s <= 0 THEN NOW() ELSE NULL END
                    WHERE id=%s AND qty >= %s
                    """,
                    (item["qty"], item["qty"], item["qty"], package_id, item["qty"]),
                )
                if cur.rowcount != 1:
                    raise ValueError("Stan paczki zmienił się. Odśwież stronę i spróbuj ponownie.")
            else:
                cur.execute(
                    """
                    SELECT COALESCE(SUM(qty),0) FROM packages
                    WHERE product_id=%s AND warehouse=%s AND status='active' AND qty>0
                    """,
                    (item["product_id"], item["warehouse"]),
                )
                numbered_qty = cur.fetchone()[0] or 0
                cur.execute("SELECT qty FROM products WHERE id=%s AND warehouse=%s", (item["product_id"], item["warehouse"]))
                total_qty = (cur.fetchone() or (0,))[0] or 0
                if total_qty - numbered_qty + 1e-9 < item["qty"]:
                    raise ValueError("Brak wystarczającej ilości bez numeru paczki. Wybierz paczkę albo zmniejsz ilość.")

            cur.execute(
                """
                UPDATE products SET qty=qty-%s
                WHERE id=%s AND warehouse=%s AND qty >= %s
                """,
                (item["qty"], item["product_id"], item["warehouse"], item["qty"]),
            )
            if cur.rowcount != 1:
                raise ValueError(f"Brak wystarczającego stanu w magazynie {item['warehouse']}.")
            cur.execute(
                """
                INSERT INTO issue_items(
                    doc_id, product_id, qty, warehouse, package_id, package_number,
                    price_netto, price_brutto
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (doc_id, item["product_id"], item["qty"], item["warehouse"], package_id,
                 package_number, item["price_netto"], item["price_brutto"]),
            )
        log_action("document.issue_created", "issue_doc", doc_id, {"doc_number": doc_number, "movement_type": movement_type, "items": len(items)}, conn)
        conn.commit()
        cache.clear()
        return redirect(f"/doc/{doc_id}")
    except ValueError as exc:
        close_with_rollback(conn)
        return str(exc), 400
    except Exception:
        close_with_rollback(conn)
        logger.exception("Issue creation failed.")
        return "Nie udało się zapisać wydania. Żadne stany nie zostały zmienione.", 500
    finally:
        conn.close()


# 📥 PRZYJĘCIE
@app.route('/przyjecie')
@login_required
def przyjecie():
    return render_receipt_form()


# 📥 ZAPIS PRZYJĘCIA
@app.route('/receive_doc', methods=['POST'])
@login_required
def receive_doc():
    return create_receipt()

    conn = db()
    cur = conn.cursor()

    date = datetime.now().strftime("%Y-%m-%d")

    cur.execute("""
        INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number)
        VALUES (%s,%s,%s,%s,%s) RETURNING id
    """, (date, request.form.get('kontrahent'), "PZ", "", "PZ"))

    doc_id = cur.fetchone()[0]

    product_ids = request.form.getlist('product_id')
    qtys = request.form.getlist('qty')
    warehouses = request.form.getlist('warehouse')
    package_numbers = request.form.getlist('package_number')
    prices_netto = request.form.getlist('price_netto')
    prices_brutto = request.form.getlist('price_brutto')

    for i in range(len(product_ids)):
        if not product_ids[i]:
            continue

        pid = int(product_ids[i])
        wh = warehouses[i]

        try:
            qty = float(qtys[i].replace(",", "."))
        except:
            qty = 0

        if qty <= 0:
            continue
        try:
            price_netto = float(prices_netto[i].replace(",", "."))
        except:
            price_netto = 0
        try:
            price_brutto = float(prices_brutto[i].replace(",", "."))
        except:
            price_brutto = 0

        # ✅ stan +
        cur.execute("""
            UPDATE products
            SET qty = qty + %s
            WHERE id=%s AND warehouse=%s
        """, (qty, pid, wh))

        # zapis pozycji
        cur.execute("""
            INSERT INTO issue_items(doc_id, product_id, qty, warehouse, price_netto, price_brutto)
            VALUES (%s,%s,%s,%s,%s,%s)
        """, (doc_id, pid, qty, wh, price_netto, price_brutto))

        # pakiet
        if package_numbers[i]:
            cur.execute("""
                INSERT INTO packages(product_id, number, qty, warehouse)
                VALUES (%s,%s,%s,%s)
            """, (pid, package_numbers[i], qty, wh))

    conn.commit()
    conn.close()

    return redirect('/historia')


def process_excel_import(upload):
    if not upload or not upload.filename:
        return "Brak pliku do importu.", 400
    if not upload.filename.lower().endswith(".xlsx"):
        return "Dozwolony jest tylko plik .xlsx.", 400
    try:
        dataframe = pd.read_excel(io.BytesIO(upload.read()), engine="openpyxl")
    except Exception:
        logger.warning("Excel file parsing failed.", exc_info=True)
        return "Nie udało się odczytać pliku Excel.", 400
    required = {"name", "qty", "unit", "warehouse"}
    if not required.issubset(dataframe.columns):
        return f"Brak wymaganych kolumn: {', '.join(sorted(required))}.", 400

    rows = []
    try:
        for row_number, row in dataframe.iterrows():
            raw_name = row.get("name", "")
            raw_unit = row.get("unit", "")
            raw_warehouse = row.get("warehouse", "")
            name = "" if pd.isna(raw_name) else str(raw_name).strip()
            unit = "" if pd.isna(raw_unit) else str(raw_unit).strip()
            warehouse = "" if pd.isna(raw_warehouse) else str(raw_warehouse).strip()
            if not name or len(name) > 200 or unit not in UNITS or warehouse not in WAREHOUSES:
                raise ValueError(
                    f"Wiersz {row_number + 2}: nieprawidłowa nazwa, jednostka lub magazyn."
                )
            qty = parse_positive_number(row.get("qty"), f"Wiersz {row_number + 2}, ilość")
            raw_package_number = row.get("package_number", "")
            package_number = (
                "" if pd.isna(raw_package_number) else str(raw_package_number).strip()
            )
            if len(package_number) > 100:
                raise ValueError(f"Wiersz {row_number + 2}: numer paczki jest za długi.")
            rows.append((name, qty, unit, warehouse, package_number))
    except ValueError as exc:
        return str(exc), 400
    if not rows:
        return "Plik nie zawiera żadnych pozycji.", 400
    rows.sort(key=lambda row: (row[3], row[0].casefold(), row[4].casefold()))

    conn = db()
    cur = conn.cursor()
    try:
        date = datetime.now().strftime("%Y-%m-%d")
        cur.execute(
            """
            INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number, movement_type)
            VALUES (%s,%s,%s,%s,%s,'PZ') RETURNING id
            """,
            (date, f"Import: {upload.filename}", "", "", ""),
        )
        doc_id = cur.fetchone()[0]
        cur.execute(
            "UPDATE issue_docs SET doc_number=%s WHERE id=%s",
            (f"PZ-IMPORT/{doc_id}/{datetime.now().year}", doc_id),
        )
        for name, qty, unit, warehouse, package_number in rows:
            cur.execute(
                "SELECT pg_advisory_xact_lock(hashtext(%s))",
                (f"product:{warehouse}:{name.casefold()}",),
            )
            cur.execute(
                """
                SELECT id FROM products
                WHERE warehouse=%s AND lower(name)=lower(%s)
                ORDER BY id LIMIT 1 FOR UPDATE
                """,
                (warehouse, name),
            )
            existing = cur.fetchone()
            if existing:
                product_id = existing[0]
                cur.execute(
                    "UPDATE products SET qty=qty+%s, unit=%s WHERE id=%s",
                    (qty, unit, product_id),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO products(name, qty, unit, warehouse, price_netto, vat)
                    VALUES (%s,%s,%s,%s,0,23) RETURNING id
                    """,
                    (name, qty, unit, warehouse),
                )
                product_id = cur.fetchone()[0]
            package_id = None
            if package_number:
                cur.execute(
                    "SELECT pg_advisory_xact_lock(hashtext(%s))",
                    (f"package:{warehouse}:{package_number.casefold()}",),
                )
                cur.execute(
                    """
                    SELECT 1 FROM packages
                    WHERE warehouse=%s AND lower(number)=lower(%s)
                    """,
                    (warehouse, package_number),
                )
                if cur.fetchone():
                    raise ValueError(f"Paczka {package_number} już istnieje w magazynie {warehouse}.")
                cur.execute(
                    """
                    INSERT INTO packages(product_id, number, qty, warehouse, initial_qty, status)
                    VALUES (%s,%s,%s,%s,%s,'active') RETURNING id
                    """,
                    (product_id, package_number, qty, warehouse, qty),
                )
                package_id = cur.fetchone()[0]
            cur.execute(
                """
                INSERT INTO issue_items(doc_id, product_id, qty, warehouse, package_id, package_number)
                VALUES (%s,%s,%s,%s,%s,%s)
                """,
                (doc_id, product_id, qty, warehouse, package_id, package_number or None),
            )
        conn.commit()
        cache.clear()
        return redirect(f"/doc/{doc_id}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Excel import failed.")
        return "Błąd importu. Żadne dane nie zostały zapisane.", 500
    finally:
        conn.close()


@app.route('/import_excel', methods=['POST'])
@login_required
def import_excel():
    file = request.files.get("excel_file")
    return process_excel_import(file)

    if not file or not file.filename:
        return "Brak pliku do importu.", 400
    if not file.filename.lower().endswith(".xlsx"):
        return "Dozwolony jest tylko plik .xlsx", 400

    try:
        file_bytes = file.read()
        df = pd.read_excel(io.BytesIO(file_bytes), engine="openpyxl")
    except Exception:
        return "Nie udało się odczytać pliku Excel.", 400

    required_columns = ["name", "qty", "unit", "warehouse"]
    if not all(col in df.columns for col in required_columns):
        return f"Brak wymaganych kolumn: {', '.join(required_columns)}", 400

    conn = db()
    cur = conn.cursor()
    try:
        for _, row in df.iterrows():
            name = str(row.get("name", "")).strip()
            unit = str(row.get("unit", "")).strip()
            warehouse = str(row.get("warehouse", "")).strip()
            if not name or not unit or not warehouse:
                continue

            try:
                qty = float(str(row.get("qty", "0")).replace(",", "."))
            except Exception:
                continue
            if qty <= 0:
                continue

            cur.execute(
                "SELECT id FROM products WHERE name=%s AND warehouse=%s LIMIT 1",
                (name, warehouse)
            )
            existing = cur.fetchone()
            if existing:
                cur.execute(
                    "UPDATE products SET qty = qty + %s, unit=%s WHERE id=%s",
                    (qty, unit, existing[0])
                )
            else:
                cur.execute(
                    """
                    INSERT INTO products(name, qty, unit, warehouse, price_netto, vat)
                    VALUES (%s,%s,%s,%s,%s,%s)
                    """,
                    (name, qty, unit, warehouse, 0, 23)
                )

        conn.commit()
    except Exception:
        conn.rollback()
        return "Błąd importu. Sprawdź dane w pliku.", 400
    finally:
        conn.close()

    return redirect('/magazyny')


def stage_general_import(cur, filename, sheets):
    metadata = [
        {
            "name": sheet["name"],
            "entity_type": sheet["entity_type"],
            "columns": sheet["columns"],
            "mapping": sheet["mapping"],
            "row_count": len(sheet["rows"]),
        }
        for sheet in sheets
    ]
    cur.execute(
        """
        INSERT INTO general_imports(filename, imported_by, detected_sheets)
        VALUES (%s,%s,%s::jsonb) RETURNING id
        """,
        (filename, session.get("user"), json.dumps(metadata, ensure_ascii=False)),
    )
    import_id = cur.fetchone()[0]
    for sheet in sheets:
        entity_type = sheet["entity_type"]
        if entity_type == "ignored":
            continue
        for row in sheet["rows"]:
            data = dict(row["normalized_data"])
            errors = (
                validate_import_row(entity_type, data, UNITS, WAREHOUSES)
                if entity_type in ENTITY_FIELDS
                else ["Wybierz typ danych i przypisz kolumny."]
            )
            duplicate = (
                detect_import_duplicate(cur, entity_type, data)
                if entity_type in ENTITY_FIELDS and not errors else None
            )
            resolution = "skip" if duplicate else "new"
            cur.execute(
                """
                INSERT INTO general_import_rows(
                    import_id,sheet_name,row_number,entity_type,source_data,
                    normalized_data,duplicate_data,resolution,validation_errors
                ) VALUES (%s,%s,%s,%s,%s::jsonb,%s::jsonb,%s::jsonb,%s,%s::jsonb)
                """,
                (
                    import_id,
                    sheet["name"],
                    row["row_number"],
                    entity_type,
                    json.dumps(row["source_data"], ensure_ascii=False),
                    json.dumps(data, ensure_ascii=False),
                    json.dumps(duplicate, ensure_ascii=False) if duplicate else None,
                    resolution,
                    json.dumps(errors, ensure_ascii=False),
                ),
            )
    return import_id


def stage_issue_import(cur, filename, sheets):
    metadata = []
    prepared = []
    allow_general_stock = issue_import_allow_general_stock(cur)
    for sheet in sheets:
        if sheet["entity_type"] == "ignored" or not sheet["rows"]:
            continue
        mapping = issue_mapping(sheet["columns"])
        selected = issue_sheet_selected(sheet)
        metadata.append(
            {
                "name": sheet["name"],
                "columns": sheet["columns"],
                "mapping": mapping,
                "row_count": len(sheet["rows"]),
                "selected": selected,
            }
        )
        for row in sheet["rows"]:
            data = (
                normalize_issue_row(row["source_data"], mapping)
                if selected
                else {field: "" for field in ISSUE_IMPORT_FIELDS}
            )
            duplicate = detect_issue_import_duplicate(cur, data) if selected else None
            resolution = "skip" if duplicate else "new"
            errors, _context = (
                issue_import_row_context(
                    cur,
                    data,
                    allow_general_stock,
                    duplicate,
                    resolution,
                )
                if selected
                else (["Wybierz arkusz i przypisz kolumny."], {})
            )
            prepared.append(
                (
                    sheet["name"],
                    row["row_number"],
                    row["source_data"],
                    data,
                    duplicate,
                    resolution,
                    selected,
                    errors,
                )
            )
    cur.execute(
        """
        INSERT INTO issue_imports(filename,imported_by,detected_sheets)
        VALUES (%s,%s,%s::jsonb) RETURNING id
        """,
        (
            filename,
            session.get("user"),
            json.dumps(metadata, ensure_ascii=False),
        ),
    )
    import_id = cur.fetchone()[0]
    for (
        sheet_name,
        row_number,
        source_data,
        data,
        duplicate,
        resolution,
        included,
        errors,
    ) in prepared:
        cur.execute(
            """
            INSERT INTO issue_import_rows(
                import_id,sheet_name,row_number,source_data,normalized_data,
                duplicate_data,resolution,included,validation_errors
            ) VALUES (%s,%s,%s,%s::jsonb,%s::jsonb,%s::jsonb,%s,%s,%s::jsonb)
            """,
            (
                import_id,
                sheet_name,
                row_number,
                json.dumps(source_data, ensure_ascii=False),
                json.dumps(data, ensure_ascii=False),
                json.dumps(duplicate, ensure_ascii=False) if duplicate else None,
                resolution,
                included,
                json.dumps(errors, ensure_ascii=False),
            ),
        )
    return import_id


def apply_issue_import_row(cur, import_id, row, data, context, document_cache):
    row_id = row[0]
    duplicate = row[5] or {}
    resolution = row[6]
    requested_number = str(data.get("doc_number") or "").strip()
    cache_key = (requested_number.casefold(), resolution)
    doc_created = False
    doc_prior = None
    if cache_key in document_cache:
        doc_id, doc_created, doc_prior = document_cache[cache_key]
        document_cache[cache_key] = (doc_id, doc_created, None)
    elif duplicate and resolution == "update":
        doc_id = duplicate["id"]
        cur.execute(
            """
            SELECT date,kontrahent,warehouse,doc_number
            FROM issue_docs WHERE id=%s FOR UPDATE
            """,
            (doc_id,),
        )
        previous_doc = cur.fetchone()
        if not previous_doc:
            raise ValueError("Nie znaleziono dokumentu wybranego do aktualizacji.")
        doc_prior = {
            "date": previous_doc[0],
            "contractor": previous_doc[1],
            "warehouse": previous_doc[2],
            "doc_number": previous_doc[3],
        }
        cur.execute(
            """
            UPDATE issue_docs
            SET date=%s,kontrahent=%s,warehouse=%s,movement_type='WZ'
            WHERE id=%s
            """,
            (
                data["date"],
                str(data["contractor"]).strip(),
                str(data["warehouse"]).strip(),
                doc_id,
            ),
        )
        document_cache[cache_key] = (doc_id, False, None)
    else:
        doc_number = requested_number
        if duplicate and resolution == "new":
            doc_number = unique_import_identifier(
                cur,
                "issue_docs",
                "doc_number",
                requested_number,
                import_id,
                row_id,
            )
        cur.execute(
            """
            INSERT INTO issue_docs(
                date,kontrahent,warehouse,image,doc_number,movement_type
            ) VALUES (%s,%s,%s,'',%s,'WZ') RETURNING id
            """,
            (
                data["date"],
                str(data["contractor"]).strip(),
                str(data["warehouse"]).strip(),
                doc_number,
            ),
        )
        doc_id = cur.fetchone()[0]
        doc_created = True
        document_cache[cache_key] = (doc_id, True, None)
    product = context["product"]
    package = context["package"]
    existing_item = context["existing_item"]
    qty = float(data["qty"])
    old_qty = float(existing_item[1] or 0) if existing_item else 0.0
    delta = qty - old_qty
    if delta > 0:
        cur.execute(
            "UPDATE products SET qty=qty-%s WHERE id=%s AND qty>=%s",
            (delta, product[0], delta),
        )
        if cur.rowcount != 1:
            raise ValueError(f"Brak stanu produktu {data['product_name']}.")
        if package:
            cur.execute(
                """
                UPDATE packages
                SET qty=qty-%s,
                    status=CASE WHEN qty-%s<=0 THEN 'issued' ELSE 'active' END,
                    archived_at=CASE WHEN qty-%s<=0 THEN NOW() ELSE NULL END
                WHERE id=%s AND qty>=%s
                """,
                (delta, delta, delta, package[0], delta),
            )
            if cur.rowcount != 1:
                raise ValueError(f"Brak stanu w paczce {package[1]}.")
    elif delta < 0:
        cur.execute(
            "UPDATE products SET qty=qty+%s WHERE id=%s",
            (-delta, product[0]),
        )
        if package:
            cur.execute(
                """
                UPDATE packages
                SET qty=qty+%s,status='active',archived_at=NULL
                WHERE id=%s
                """,
                (-delta, package[0]),
            )
    prior_data = {
        "doc_created": doc_created,
        "doc_prior": doc_prior,
    }
    if existing_item:
        item_id = existing_item[0]
        prior_data["item"] = {
            "qty": old_qty,
            "package_id": existing_item[2],
            "package_number": existing_item[3],
            "dimension": existing_item[4],
            "species": existing_item[5],
            "notes": existing_item[6],
        }
        cur.execute(
            """
            UPDATE issue_items
            SET qty=%s,warehouse=%s,package_id=%s,package_number=%s,
                dimension=%s,species=%s,notes=%s,price_netto=%s,price_brutto=%s
            WHERE id=%s
            """,
            (
                qty,
                data["warehouse"],
                package[0] if package else None,
                package[1] if package else None,
                str(data.get("dimension") or "").strip() or None,
                str(data.get("species") or "").strip() or None,
                str(data.get("notes") or "").strip() or None,
                float(product[3] or 0),
                float(product[3] or 0) * (1 + float(product[4] or 0) / 100),
                item_id,
            ),
        )
        action = "updated_item"
        outcome = "updated"
    else:
        cur.execute(
            """
            INSERT INTO issue_items(
                doc_id,product_id,qty,warehouse,package_id,package_number,
                price_netto,price_brutto,dimension,species,notes
            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING id
            """,
            (
                doc_id,
                product[0],
                qty,
                data["warehouse"],
                package[0] if package else None,
                package[1] if package else None,
                float(product[3] or 0),
                float(product[3] or 0) * (1 + float(product[4] or 0) / 100),
                str(data.get("dimension") or "").strip() or None,
                str(data.get("species") or "").strip() or None,
                str(data.get("notes") or "").strip() or None,
            ),
        )
        item_id = cur.fetchone()[0]
        action = "added_item"
        outcome = "added"
    cur.execute(
        """
        INSERT INTO issue_import_effects(
            import_id,row_id,action,doc_id,item_id,product_id,package_id,qty,prior_data
        ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb)
        """,
        (
            import_id,
            row_id,
            action,
            doc_id,
            item_id,
            product[0],
            package[0] if package else None,
            qty,
            json.dumps(prior_data, ensure_ascii=False, default=str),
        ),
    )
    issue_doc_history(
        cur,
        doc_id,
        "import wydania z Excela",
        f"Import #{import_id}, wiersz {row[2]}",
    )
    return outcome


def undo_issue_import(cur, import_id, actor):
    cur.execute(
        """
        SELECT id,row_id,action,doc_id,item_id,product_id,package_id,qty,prior_data
        FROM issue_import_effects
        WHERE import_id=%s ORDER BY id DESC FOR UPDATE
        """,
        (import_id,),
    )
    effects = cur.fetchall()
    if not effects:
        raise ValueError("Import nie zawiera operacji możliwych do cofnięcia.")
    created_docs = set()
    restored_docs = set()
    locked_docs = set()
    for effect in effects:
        action = effect[2]
        doc_id = effect[3]
        item_id = effect[4]
        product_id = effect[5]
        package_id = effect[6]
        prior = effect[8] or {}
        if doc_id not in locked_docs:
            cur.execute(
                "SELECT voided_at FROM issue_docs WHERE id=%s FOR UPDATE",
                (doc_id,),
            )
            document = cur.fetchone()
            if not document:
                raise ValueError(
                    "Nie można cofnąć importu: dokument został usunięty."
                )
            if document[0]:
                raise ValueError(
                    "Nie można cofnąć importu: jeden z dokumentów został już anulowany."
                )
            locked_docs.add(doc_id)
        cur.execute(
            "SELECT qty FROM issue_items WHERE id=%s FOR UPDATE",
            (item_id,),
        )
        current_item = cur.fetchone()
        if not current_item:
            raise ValueError(
                "Nie można cofnąć importu: jedna z pozycji została usunięta."
            )
        current_qty = float(current_item[0] or 0)
        if action == "added_item":
            cur.execute(
                "UPDATE products SET qty=qty+%s WHERE id=%s",
                (current_qty, product_id),
            )
            if package_id:
                cur.execute(
                    """
                    UPDATE packages
                    SET qty=qty+%s,status='active',archived_at=NULL
                    WHERE id=%s
                    """,
                    (current_qty, package_id),
                )
            if prior.get("doc_created"):
                created_docs.add(doc_id)
            else:
                cur.execute("DELETE FROM issue_items WHERE id=%s", (item_id,))
        elif action == "updated_item":
            old_item = prior.get("item") or {}
            old_qty = float(old_item.get("qty") or 0)
            restore_delta = current_qty - old_qty
            if restore_delta >= 0:
                cur.execute(
                    "UPDATE products SET qty=qty+%s WHERE id=%s",
                    (restore_delta, product_id),
                )
                if package_id:
                    cur.execute(
                        """
                        UPDATE packages
                        SET qty=qty+%s,status='active',archived_at=NULL
                        WHERE id=%s
                        """,
                        (restore_delta, package_id),
                    )
            else:
                needed = -restore_delta
                cur.execute(
                    "UPDATE products SET qty=qty-%s WHERE id=%s AND qty>=%s",
                    (needed, product_id, needed),
                )
                if cur.rowcount != 1:
                    raise ValueError(
                        "Nie można odtworzyć poprzedniego stanu produktu."
                    )
                if package_id:
                    cur.execute(
                        "UPDATE packages SET qty=qty-%s WHERE id=%s AND qty>=%s",
                        (needed, package_id, needed),
                    )
                    if cur.rowcount != 1:
                        raise ValueError(
                            "Nie można odtworzyć poprzedniego stanu paczki."
                        )
            cur.execute(
                """
                UPDATE issue_items
                SET qty=%s,package_id=%s,package_number=%s,dimension=%s,
                    species=%s,notes=%s
                WHERE id=%s
                """,
                (
                    old_qty,
                    old_item.get("package_id"),
                    old_item.get("package_number"),
                    old_item.get("dimension"),
                    old_item.get("species"),
                    old_item.get("notes"),
                    item_id,
                ),
            )
        doc_prior = prior.get("doc_prior")
        if doc_prior and doc_id not in restored_docs:
            cur.execute(
                """
                UPDATE issue_docs
                SET date=%s,kontrahent=%s,warehouse=%s,doc_number=%s
                WHERE id=%s
                """,
                (
                    doc_prior.get("date"),
                    doc_prior.get("contractor"),
                    doc_prior.get("warehouse"),
                    doc_prior.get("doc_number"),
                    doc_id,
                ),
            )
            restored_docs.add(doc_id)
    for doc_id in created_docs:
        cur.execute(
            """
            UPDATE issue_docs SET voided_at=NOW(),voided_by=%s
            WHERE id=%s AND voided_at IS NULL
            """,
            (actor, doc_id),
        )
        issue_doc_history(
            cur,
            doc_id,
            "cofnięto import wydania",
            f"Import #{import_id}",
        )
    cur.execute(
        """
        UPDATE issue_imports
        SET status='undone',undone_at=NOW(),undone_by=%s
        WHERE id=%s
        """,
        (actor, import_id),
    )


def import_run_query(cur, import_id, for_update=False):
    cur.execute(
        """
        SELECT id,filename,imported_by,status,detected_sheets,summary,errors,
               created_at,completed_at
        FROM general_imports WHERE id=%s
        """ + (" FOR UPDATE" if for_update else ""),
        (import_id,),
    )
    return cur.fetchone()


def import_rows_query(cur, import_id, for_update=False):
    cur.execute(
        """
        SELECT id,sheet_name,row_number,entity_type,source_data,normalized_data,
               duplicate_data,resolution,included,validation_errors
        FROM general_import_rows
        WHERE import_id=%s ORDER BY sheet_name,row_number,id
        """ + (" FOR UPDATE" if for_update else ""),
        (import_id,),
    )
    return cur.fetchall()


def get_or_create_import_product(cur, data):
    name = str(data.get("product_name") or data.get("name") or "").strip()
    warehouse = str(data.get("warehouse") or "").strip()
    unit = str(data.get("unit") or "szt").strip() or "szt"
    cur.execute(
        """
        SELECT id,qty FROM products
        WHERE lower(name)=lower(%s) AND warehouse=%s
        ORDER BY id LIMIT 1 FOR UPDATE
        """,
        (name, warehouse),
    )
    product = cur.fetchone()
    if product:
        return product[0], product[1] or 0
    cur.execute(
        """
        INSERT INTO products(name,qty,unit,warehouse,price_netto,vat)
        VALUES (%s,0,%s,%s,%s,%s) RETURNING id
        """,
        (
            name,
            unit if unit in UNITS else "szt",
            warehouse,
            float(data.get("price_netto") or 0),
            float(data.get("vat") or 0),
        ),
    )
    return cur.fetchone()[0], 0


def apply_import_product(cur, row, data):
    duplicate = row[6] or {}
    resolution = row[7]
    values = (
        str(data.get("name") or "").strip(),
        float(data.get("qty") or 0),
        str(data.get("unit") or "").strip(),
        str(data.get("warehouse") or "").strip(),
        float(data.get("price_netto") or 0),
        float(data.get("vat") or 0),
    )
    if duplicate and resolution == "update":
        cur.execute(
            """
            UPDATE products
            SET name=%s,qty=%s,unit=%s,warehouse=%s,price_netto=%s,vat=%s
            WHERE id=%s
            """,
            values + (duplicate["id"],),
        )
        if cur.rowcount == 1:
            return "updated"
    cur.execute(
        """
        INSERT INTO products(name,qty,unit,warehouse,price_netto,vat)
        VALUES (%s,%s,%s,%s,%s,%s)
        """,
        values,
    )
    return "added"


def apply_import_package(cur, row, data, product_snapshots, import_id):
    duplicate = row[6] or {}
    resolution = row[7]
    product_id, _ = get_or_create_import_product(cur, data)
    qty = float(data.get("qty") or 0)
    package_number = str(data.get("package_number") or "").strip()
    product_key = duplicate_identity(
        "product",
        {"name": data.get("product_name"), "warehouse": data.get("warehouse")},
    )
    update_total = product_key not in product_snapshots
    if duplicate and resolution == "update":
        cur.execute("SELECT qty,product_id FROM packages WHERE id=%s FOR UPDATE", (duplicate["id"],))
        old = cur.fetchone()
        if old:
            old_qty = float(old[0] or 0)
            old_product_id = old[1]
            cur.execute(
                """
                UPDATE packages
                SET product_id=%s,number=%s,qty=%s,warehouse=%s,initial_qty=%s,
                    status=CASE WHEN %s>0 THEN 'active' ELSE 'issued' END,
                    archived_at=CASE WHEN %s>0 THEN NULL ELSE NOW() END
                WHERE id=%s
                """,
                (
                    product_id, package_number, qty, data["warehouse"], qty,
                    qty, qty, duplicate["id"],
                ),
            )
            if update_total:
                if old_product_id == product_id:
                    delta = qty - old_qty
                    cur.execute(
                        "UPDATE products SET qty=qty+%s WHERE id=%s AND qty+%s>=0",
                        (delta, product_id, delta),
                    )
                    if cur.rowcount != 1:
                        raise ValueError("Aktualizacja paczki spowodowałaby ujemny stan produktu.")
                else:
                    cur.execute(
                        "UPDATE products SET qty=qty-%s WHERE id=%s AND qty>=%s",
                        (old_qty, old_product_id, old_qty),
                    )
                    if cur.rowcount != 1:
                        raise ValueError("Nie można odpiąć paczki od poprzedniego produktu.")
                    cur.execute(
                        "UPDATE products SET qty=qty+%s WHERE id=%s",
                        (qty, product_id),
                    )
            return "updated"
    if duplicate and resolution == "new":
        package_number = unique_import_identifier(
            cur, "packages", "number", package_number, import_id, row[0]
        )
    cur.execute(
        """
        INSERT INTO packages(product_id,number,qty,warehouse,initial_qty,status)
        VALUES (%s,%s,%s,%s,%s,%s)
        """,
        (
            product_id, package_number, qty, data["warehouse"], qty,
            "active" if qty > 0 else "issued",
        ),
    )
    if update_total:
        cur.execute("UPDATE products SET qty=qty+%s WHERE id=%s", (qty, product_id))
    return "added"


def apply_import_contractor(cur, row, data):
    duplicate = row[6] or {}
    values = (
        str(data.get("contractor") or "").strip(),
        str(data.get("nip") or "").strip() or None,
        str(data.get("email") or "").strip() or None,
        str(data.get("phone") or "").strip() or None,
        str(data.get("address") or "").strip() or None,
    )
    if duplicate and row[7] == "update":
        cur.execute(
            """
            UPDATE contractors
            SET name=%s,nip=%s,email=%s,phone=%s,address=%s,updated_at=NOW()
            WHERE id=%s
            """,
            values + (duplicate["id"],),
        )
        if cur.rowcount == 1:
            return "updated"
    cur.execute(
        "INSERT INTO contractors(name,nip,email,phone,address) VALUES (%s,%s,%s,%s,%s)",
        values,
    )
    return "added"


def apply_import_document(cur, row, data, import_id, document_cache):
    entity_type = row[3]
    duplicate = row[6] or {}
    resolution = row[7]
    movement = str(
        data.get("movement_type") or ("PZ" if entity_type == "receipt" else "WZ")
    ).upper()
    requested_number = str(data.get("doc_number") or "").strip()
    cache_key = (movement, requested_number.casefold()) if requested_number else None
    if duplicate and resolution == "update":
        cur.execute(
            """
            UPDATE issue_docs
            SET date=%s,kontrahent=%s,movement_type=%s
            WHERE id=%s
            """,
            (data.get("date"), data.get("contractor"), movement, duplicate["id"]),
        )
        return "updated"
    if cache_key and cache_key in document_cache:
        doc_id = document_cache[cache_key]
    else:
        doc_number = requested_number
        if duplicate and resolution == "new":
            doc_number = unique_import_identifier(
                cur, "issue_docs", "doc_number", doc_number, import_id, row[0]
            )
        cur.execute(
            """
            INSERT INTO issue_docs(date,kontrahent,warehouse,image,doc_number,movement_type)
            VALUES (%s,%s,%s,'',%s,%s) RETURNING id
            """,
            (
                data.get("date"),
                str(data.get("contractor") or "").strip(),
                str(data.get("warehouse") or "").strip(),
                doc_number or "",
                movement,
            ),
        )
        doc_id = cur.fetchone()[0]
        if not doc_number:
            doc_number = f"{movement}-IMPORT/{import_id}/{doc_id}"
            cur.execute("UPDATE issue_docs SET doc_number=%s WHERE id=%s", (doc_number, doc_id))
        if cache_key:
            document_cache[cache_key] = doc_id

    if not str(data.get("product_name") or "").strip() or data.get("qty") in ("", None):
        return "added"
    product_id, product_qty = get_or_create_import_product(cur, data)
    qty = float(data["qty"])
    package_id = None
    package_number = str(data.get("package_number") or "").strip() or None
    if movement == "PZ":
        cur.execute("UPDATE products SET qty=qty+%s WHERE id=%s", (qty, product_id))
        if package_number:
            package_number = unique_import_identifier(
                cur, "packages", "number", package_number, import_id, row[0]
            )
            cur.execute(
                """
                INSERT INTO packages(product_id,number,qty,warehouse,initial_qty,status)
                VALUES (%s,%s,%s,%s,%s,'active') RETURNING id
                """,
                (product_id, package_number, qty, data.get("warehouse"), qty),
            )
            package_id = cur.fetchone()[0]
    else:
        if product_qty + 1e-9 < qty:
            raise ValueError(
                f"Brak stanu produktu {data.get('product_name')} w magazynie {data.get('warehouse')}."
            )
        if package_number:
            cur.execute(
                """
                SELECT id,qty FROM packages
                WHERE product_id=%s AND warehouse=%s AND lower(number)=lower(%s)
                  AND status='active'
                ORDER BY id LIMIT 1 FOR UPDATE
                """,
                (product_id, data.get("warehouse"), package_number),
            )
            package = cur.fetchone()
            if not package or float(package[1] or 0) + 1e-9 < qty:
                raise ValueError(f"Brak wymaganej ilości w paczce {package_number}.")
            package_id = package[0]
            cur.execute(
                """
                UPDATE packages SET qty=qty-%s,
                    status=CASE WHEN qty-%s<=0 THEN 'issued' ELSE 'active' END,
                    archived_at=CASE WHEN qty-%s<=0 THEN NOW() ELSE NULL END
                WHERE id=%s
                """,
                (qty, qty, qty, package_id),
            )
        cur.execute("UPDATE products SET qty=qty-%s WHERE id=%s", (qty, product_id))
    cur.execute(
        """
        INSERT INTO issue_items(
            doc_id,product_id,qty,warehouse,package_id,package_number,
            price_netto,price_brutto,dimension,species,notes
        ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
        """,
        (
            doc_id, product_id, qty, data.get("warehouse"), package_id, package_number,
            float(data.get("price_netto") or 0),
            float(data.get("price_netto") or 0) * (1 + float(data.get("vat") or 0) / 100),
            str(data.get("dimension") or "").strip() or None,
            str(data.get("species") or "").strip() or None,
            str(data.get("notes") or "").strip() or None,
        ),
    )
    return "added"


def apply_import_shop_order(cur, row, data, import_id, order_cache=None):
    order_cache = order_cache if order_cache is not None else {}
    duplicate = row[6] or {}
    order_number = str(data.get("order_number") or "").strip()
    cache_key = order_number.casefold()
    values = (
        data.get("date"),
        str(data.get("contractor") or "").strip(),
        str(data.get("address") or "").strip(),
        str(data.get("phone") or "").strip(),
        str(data.get("email") or "").strip(),
        float(data.get("shipping_cost") or 0),
        str(data.get("payment_method") or "").strip(),
        str(data.get("payment_status") or "Oczekuje na płatność").strip(),
        str(data.get("status") or "Nowe zamówienie").strip(),
        str(data.get("doc_number") or "").strip(),
        str(data.get("tracking_number") or "").strip(),
        str(data.get("nip") or "").strip(),
    )
    outcome = "added"
    order_id = None
    if duplicate and row[7] == "update":
        cur.execute(
            """
            UPDATE shop_orders SET order_date=%s,customer_name=%s,delivery_address=%s,
                phone=%s,email=%s,shipping_cost=%s,payment_method=%s,payment_status=%s,
                status=%s,sales_document_number=%s,tracking_number=%s,nip=%s,updated_at=NOW()
            WHERE id=%s
            """,
            values + (duplicate["id"],),
        )
        if cur.rowcount == 1:
            order_id = duplicate["id"]
            outcome = "updated"
    if order_id is None and cache_key in order_cache:
        order_id = order_cache[cache_key]
    if order_id is None:
        if duplicate and row[7] == "new":
            order_number = unique_import_identifier(
                cur, "shop_orders", "order_number", order_number, import_id, row[0]
            )
        cur.execute(
            """
            INSERT INTO shop_orders(
                order_number,order_date,customer_name,delivery_address,phone,email,
                shipping_cost,payment_method,payment_status,status,sales_document_number,
                tracking_number,nip,created_by
            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id
            """,
            (order_number,) + values + (session.get("user"),),
        )
        order_id = cur.fetchone()[0]
    order_cache[cache_key] = order_id
    if str(data.get("product_name") or "").strip() and data.get("qty") not in ("", None):
        product_id, _ = get_or_create_import_product(cur, data)
        qty = float(data["qty"])
        price = float(data.get("price_netto") or 0)
        vat = float(data.get("vat") or 0)
        cur.execute(
            """
            INSERT INTO shop_order_items(
                order_id,product_id,product_name,qty,price_netto,price_brutto,vat,
                warehouse,reserved_qty,issued_qty
            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,0,0)
            """,
            (
                order_id, product_id, data["product_name"], qty, price,
                price * (1 + vat / 100), vat, data.get("warehouse") or "Inne",
            ),
        )
    ensure_shop_accounting_row(cur, order_id)
    return outcome


def apply_import_accounting(cur, row, data):
    duplicate = row[6] or detect_import_duplicate(cur, "accounting", data) or {}
    if not duplicate:
        raise ValueError(f"Nie znaleziono zamówienia {data.get('order_number')} dla księgowości.")
    order_id = duplicate["id"]
    ensure_shop_accounting_row(cur, order_id)
    cur.execute(
        """
        UPDATE shop_accounting SET payment_method=%s,paid=%s,amount_paid=%s,
            amount_due=%s,invoice_number=%s,receipt_number=%s,proforma_number=%s,
            payment_received_date=CASE WHEN %s THEN NULLIF(%s,'')::date
                                       ELSE payment_received_date END,
            updated_by=%s,updated_at=NOW()
        WHERE order_id=%s
        """,
        (
            str(data.get("payment_method") or "").strip() or None,
            bool(data.get("paid")),
            float(data.get("amount_paid") or 0),
            float(data.get("amount_due") or 0),
            str(data.get("invoice_number") or "").strip(),
            str(data.get("receipt_number") or "").strip(),
            str(data.get("proforma_number") or "").strip(),
            bool(data.get("date")),
            str(data.get("date") or ""),
            session.get("user"),
            order_id,
        ),
    )
    sync_accounting_payment_status(cur, order_id)
    return "updated"


@app.route('/kontrahenci')
@login_required
def contractors_page():
    if session.get("role") not in {"admin", "warehouse", "accounting", "sales"}:
        return "Brak uprawnień.", 403
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id,name,nip,email,phone,address,active,created_at,updated_at
        FROM contractors ORDER BY lower(name),id
        """
    )
    contractors = cur.fetchall()
    conn.close()
    return render_template("contractors.html", contractors=contractors)


@app.route('/kontrahenci/<int:contractor_id>/edit', methods=['POST'])
@login_required
def edit_contractor(contractor_id):
    if session.get("role") not in {"admin", "warehouse", "accounting", "sales"}:
        return "Brak uprawnień.", 403
    name = (request.form.get("name") or "").strip()
    if not name or len(name) > 200:
        return "Nazwa kontrahenta jest wymagana i może mieć maksymalnie 200 znaków.", 400
    values = (
        name,
        (request.form.get("nip") or "").strip()[:30] or None,
        (request.form.get("email") or "").strip()[:254] or None,
        (request.form.get("phone") or "").strip()[:50] or None,
        (request.form.get("address") or "").strip()[:500] or None,
        contractor_id,
    )
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            UPDATE contractors SET name=%s,nip=%s,email=%s,phone=%s,address=%s,
                updated_at=NOW() WHERE id=%s
            """,
            values,
        )
        if cur.rowcount != 1:
            return "Nie znaleziono kontrahenta.", 404
        conn.commit()
        return redirect("/kontrahenci")
    except Exception:
        conn.rollback()
        logger.exception("Contractor update failed.")
        return "Nie udało się zaktualizować kontrahenta.", 500
    finally:
        conn.close()


@app.route('/kontrahenci/add', methods=['POST'])
@login_required
def add_contractor():
    if session.get("role") not in {"admin", "warehouse", "accounting", "sales"}:
        return "Brak uprawnień.", 403
    name = (request.form.get("name") or "").strip()
    if not name or len(name) > 200:
        return "Nazwa kontrahenta jest wymagana i może mieć maksymalnie 200 znaków.", 400
    values = (
        name,
        (request.form.get("nip") or "").strip()[:30] or None,
        (request.form.get("email") or "").strip()[:254] or None,
        (request.form.get("phone") or "").strip()[:50] or None,
        (request.form.get("address") or "").strip()[:500] or None,
    )
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            INSERT INTO contractors(name,nip,email,phone,address)
            VALUES (%s,%s,%s,%s,%s)
            """,
            values,
        )
        conn.commit()
        return redirect("/kontrahenci")
    except Exception:
        conn.rollback()
        logger.exception("Contractor creation failed.")
        return "Nie udało się dodać kontrahenta.", 500
    finally:
        conn.close()


@app.route('/import-ogolny')
@login_required
def general_import_page():
    denied = import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    if session.get("role") == "admin":
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at,completed_at
            FROM general_imports ORDER BY id DESC LIMIT 100
            """
        )
    else:
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at,completed_at
            FROM general_imports WHERE lower(imported_by)=lower(%s)
            ORDER BY id DESC LIMIT 100
            """,
            (session.get("user"),),
        )
    history = cur.fetchall()
    conn.close()
    return render_template(
        "general_import.html",
        history=history,
        allowed_groups=general_import_groups(),
    )


@app.route('/import-ogolny/upload', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def general_import_upload():
    denied = import_access_error()
    if denied:
        return denied
    upload = request.files.get("excel_file")
    if not upload or not upload.filename:
        return "Wybierz plik Excel.", 400
    if not upload.filename.lower().endswith(".xlsx"):
        return "Dozwolony jest wyłącznie plik .xlsx.", 400
    filename = secure_filename(upload.filename)[:200] or "import.xlsx"
    try:
        sheets = parse_workbook(upload.read())
    except ValueError as exc:
        logger.info("General Excel import rejected: %s", exc)
        return f"Nie udało się odczytać pliku: {exc}", 400
    except Exception:
        logger.warning("General Excel import parsing failed.", exc_info=True)
        return "Nie udało się odczytać pliku. Sprawdź, czy jest to prawidłowy plik .xlsx.", 400
    staged_sheets = [sheet for sheet in sheets if sheet["entity_type"] != "ignored" and sheet["rows"]]
    if not staged_sheets:
        return "Plik nie zawiera danych możliwych do przygotowania.", 400
    recognized = [sheet for sheet in staged_sheets if sheet["entity_type"] in ENTITY_FIELDS]
    allowed_groups = general_import_groups()
    forbidden = {
        ENTITY_LABELS[sheet["entity_type"]]
        for sheet in recognized
        if ENTITY_GROUPS[sheet["entity_type"]] not in allowed_groups
    }
    if forbidden:
        return (
            "Plik zawiera dane bez wymaganych uprawnień: " + ", ".join(sorted(forbidden)),
            403,
        )
    conn = db()
    cur = conn.cursor()
    try:
        import_id = stage_general_import(cur, filename, sheets)
        conn.commit()
        return redirect(f"/import-ogolny/{import_id}")
    except Exception:
        conn.rollback()
        logger.exception("General import staging failed.")
        return "Nie udało się przygotować podglądu importu.", 500
    finally:
        conn.close()


@app.route('/products/<int:product_id>/edit', methods=['POST'])
@permission_required("inventory_manage")
def edit_product(product_id):
    name = (request.form.get("name") or "").strip()
    unit = (request.form.get("unit") or "").strip()
    warehouse = (request.form.get("warehouse") or "").strip()
    if not name or len(name) > 200 or unit not in UNITS or warehouse not in WAREHOUSES:
        return "Nieprawidłowa nazwa, jednostka lub magazyn.", 400
    try:
        qty = parse_nonnegative_number(request.form.get("qty"), "Ilość")
        price_netto = parse_nonnegative_number(request.form.get("price_netto"), "Cena netto")
        vat = parse_nonnegative_number(request.form.get("vat"), "VAT")
    except ValueError as exc:
        return str(exc), 400
    if vat > 100:
        return "VAT nie może przekraczać 100%.", 400
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT warehouse FROM products WHERE id=%s FOR UPDATE",
            (product_id,),
        )
        existing = cur.fetchone()
        if not existing:
            return "Nie znaleziono produktu.", 404
        cur.execute(
            """
            SELECT COALESCE(SUM(qty),0),COUNT(*)
            FROM packages WHERE product_id=%s AND status='active'
            """,
            (product_id,),
        )
        package_qty, package_count = cur.fetchone()
        if package_count and existing[0] != warehouse:
            raise ValueError("Nie można zmienić magazynu produktu posiadającego paczki.")
        if qty + 1e-9 < float(package_qty or 0):
            raise ValueError("Stan produktu nie może być mniejszy niż suma aktywnych paczek.")
        cur.execute(
            """
            SELECT 1 FROM products
            WHERE id<>%s AND warehouse=%s AND lower(name)=lower(%s) LIMIT 1
            """,
            (product_id, warehouse, name),
        )
        if cur.fetchone():
            raise ValueError("Taki produkt już istnieje w wybranym magazynie.")
        cur.execute(
            """
            UPDATE products SET name=%s,qty=%s,unit=%s,warehouse=%s,price_netto=%s,vat=%s
            WHERE id=%s
            """,
            (name, qty, unit, warehouse, price_netto, vat, product_id),
        )
        conn.commit()
        cache.clear()
        return redirect(f"/magazyn/{warehouse}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 409
    except Exception:
        conn.rollback()
        logger.exception("Product update failed.")
        return "Nie udało się zaktualizować produktu.", 500
    finally:
        conn.close()


@app.route('/packages/<int:package_id>/edit', methods=['POST'])
@permission_required("inventory_manage")
def edit_package(package_id):
    number = (request.form.get("number") or "").strip()
    if not number or len(number) > 100:
        return "Numer paczki jest wymagany i może mieć maksymalnie 100 znaków.", 400
    try:
        qty = parse_nonnegative_number(request.form.get("qty"), "Ilość")
    except ValueError as exc:
        return str(exc), 400
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT product_id,qty,warehouse FROM packages WHERE id=%s FOR UPDATE",
            (package_id,),
        )
        package = cur.fetchone()
        if not package:
            return "Nie znaleziono paczki.", 404
        cur.execute(
            """
            SELECT 1 FROM packages
            WHERE id<>%s AND warehouse=%s AND lower(number)=lower(%s) LIMIT 1
            """,
            (package_id, package[2], number),
        )
        if cur.fetchone():
            raise ValueError("Paczka o takim numerze już istnieje w tym magazynie.")
        delta = qty - float(package[1] or 0)
        cur.execute(
            """
            UPDATE products SET qty=qty+%s
            WHERE id=%s AND qty+%s>=0
            """,
            (delta, package[0], delta),
        )
        if cur.rowcount != 1:
            raise ValueError("Zmiana paczki spowodowałaby ujemny stan produktu.")
        cur.execute(
            """
            UPDATE packages SET number=%s,qty=%s,initial_qty=GREATEST(initial_qty,%s),
                status=CASE WHEN %s>0 THEN 'active' ELSE 'issued' END,
                archived_at=CASE WHEN %s>0 THEN NULL ELSE NOW() END
            WHERE id=%s
            """,
            (number, qty, qty, qty, qty, package_id),
        )
        conn.commit()
        cache.clear()
        return ("", 204)
    except ValueError as exc:
        conn.rollback()
        return str(exc), 409
    except Exception:
        conn.rollback()
        logger.exception("Package update failed.")
        return "Nie udało się zaktualizować paczki.", 500
    finally:
        conn.close()


@app.route('/import-ogolny/<int:import_id>')
@login_required
def general_import_preview(import_id):
    denied = import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    run = import_run_query(cur, import_id)
    if not import_run_allowed(run):
        conn.close()
        return "Nie znaleziono importu.", 404
    rows = import_rows_query(cur, import_id)
    conn.close()
    return render_template(
        "general_import_preview.html",
        run=run,
        rows=rows,
        entity_fields=ENTITY_FIELDS,
        entity_groups=ENTITY_GROUPS,
        entity_labels=ENTITY_LABELS,
        field_labels=FIELD_LABELS,
        units=sorted(UNITS),
        warehouses=WAREHOUSES,
        allowed_groups=general_import_groups(),
    )


@app.route('/import-ogolny/<int:import_id>/mapping', methods=['POST'])
@login_required
def general_import_mapping(import_id):
    denied = import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = import_run_query(cur, import_id, for_update=True)
        if not import_run_allowed(run) or run[3] != "draft":
            return "Import nie jest dostępny do edycji.", 409
        metadata = list(run[4] or [])
        allowed_groups = general_import_groups()
        for index, sheet in enumerate(metadata):
            entity_type = request.form.get(f"entity_{index}", sheet.get("entity_type", "unknown"))
            if entity_type not in ENTITY_FIELDS:
                sheet["entity_type"] = entity_type
                sheet["mapping"] = {}
                cur.execute(
                    """
                    UPDATE general_import_rows
                    SET entity_type=%s,normalized_data='{}'::jsonb,
                        validation_errors='["Wybierz typ danych i przypisz kolumny."]'::jsonb,
                        duplicate_data=NULL,resolution='skip'
                    WHERE import_id=%s AND sheet_name=%s
                    """,
                    (entity_type, import_id, sheet["name"]),
                )
                continue
            if ENTITY_GROUPS[entity_type] not in allowed_groups:
                return "Brak uprawnień do wybranego typu danych.", 403
            mapping = {}
            for field in ENTITY_FIELDS[entity_type]:
                column = request.form.get(f"map_{index}_{field}", "").strip()
                if column and column in sheet.get("columns", []):
                    mapping[field] = column
            sheet["entity_type"] = entity_type
            sheet["mapping"] = mapping
            cur.execute(
                """
                SELECT id,source_data FROM general_import_rows
                WHERE import_id=%s AND sheet_name=%s FOR UPDATE
                """,
                (import_id, sheet["name"]),
            )
            for row_id, source_data in cur.fetchall():
                data = normalize_import_row(entity_type, source_data or {}, mapping)
                errors = validate_import_row(entity_type, data, UNITS, WAREHOUSES)
                duplicate = detect_import_duplicate(cur, entity_type, data) if not errors else None
                cur.execute(
                    """
                    UPDATE general_import_rows SET entity_type=%s,normalized_data=%s::jsonb,
                        validation_errors=%s::jsonb,duplicate_data=%s::jsonb,
                        resolution=%s
                    WHERE id=%s
                    """,
                    (
                        entity_type,
                        json.dumps(data, ensure_ascii=False),
                        json.dumps(errors, ensure_ascii=False),
                        json.dumps(duplicate, ensure_ascii=False) if duplicate else None,
                        "skip" if duplicate else "new",
                        row_id,
                    ),
                )
        cur.execute(
            "UPDATE general_imports SET detected_sheets=%s::jsonb WHERE id=%s",
            (json.dumps(metadata, ensure_ascii=False), import_id),
        )
        conn.commit()
        return redirect(f"/import-ogolny/{import_id}?mapping=saved")
    except Exception:
        conn.rollback()
        logger.exception("General import mapping update failed.")
        return "Nie udało się zapisać mapowania kolumn.", 500
    finally:
        conn.close()


@app.route('/import-ogolny/<int:import_id>/prepare', methods=['POST'])
@login_required
def general_import_prepare(import_id):
    denied = import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = import_run_query(cur, import_id, for_update=True)
        if not import_run_allowed(run) or run[3] != "draft":
            return "Import nie jest dostępny do edycji.", 409
        rows = import_rows_query(cur, import_id, for_update=True)
        allowed_groups = general_import_groups()
        for row in rows:
            row_id, _, _, entity_type = row[:4]
            included = request.form.get(f"included_{row_id}") == "on"
            if entity_type not in ENTITY_FIELDS:
                cur.execute(
                    "UPDATE general_import_rows SET included=%s WHERE id=%s",
                    (included, row_id),
                )
                continue
            if ENTITY_GROUPS.get(entity_type) not in allowed_groups:
                return "Brak uprawnień do jednego z wierszy.", 403
            data = dict(row[5] or {})
            for field in ENTITY_FIELDS.get(entity_type, ()):
                key = f"row_{row_id}_{field}"
                if key in request.form:
                    data[field] = request.form.get(key, "").strip()
            errors = validate_import_row(entity_type, data, UNITS, WAREHOUSES)
            duplicate = detect_import_duplicate(cur, entity_type, data) if not errors else None
            resolution = request.form.get(f"resolution_{row_id}", row[7])
            if resolution not in {"skip", "update", "new"}:
                resolution = "skip" if duplicate else "new"
            if resolution == "update" and not duplicate:
                resolution = "new"
            cur.execute(
                """
                UPDATE general_import_rows SET normalized_data=%s::jsonb,
                    validation_errors=%s::jsonb,duplicate_data=%s::jsonb,
                    resolution=%s,included=%s
                WHERE id=%s
                """,
                (
                    json.dumps(data, ensure_ascii=False),
                    json.dumps(errors, ensure_ascii=False),
                    json.dumps(duplicate, ensure_ascii=False) if duplicate else None,
                    resolution,
                    included,
                    row_id,
                ),
            )
        conn.commit()
        return redirect(f"/import-ogolny/{import_id}?prepared=1")
    except Exception:
        conn.rollback()
        logger.exception("General import preview update failed.")
        return "Nie udało się zapisać korekt.", 500
    finally:
        conn.close()


@app.route('/import-ogolny/<int:import_id>/confirm', methods=['POST'])
@login_required
@limiter.limit("5 per hour")
def general_import_confirm(import_id):
    denied = import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT pg_advisory_xact_lock(%s)", (88000000 + import_id,))
        run = import_run_query(cur, import_id, for_update=True)
        if not import_run_allowed(run) or run[3] != "draft":
            return "Import został już zakończony albo nie istnieje.", 409
        rows = import_rows_query(cur, import_id, for_update=True)
        allowed_groups = general_import_groups()
        validation_failed = False
        for row in rows:
            if not row[8]:
                continue
            data = dict(row[5] or {})
            errors = validate_import_row(row[3], data, UNITS, WAREHOUSES)
            if ENTITY_GROUPS.get(row[3]) not in allowed_groups:
                errors.append("Brak uprawnień do tego typu danych.")
            current_duplicate = (
                detect_import_duplicate(cur, row[3], data)
                if row[3] in ENTITY_FIELDS and not errors else None
            )
            staged_duplicate = row[6] or None
            duplicate_changed = (
                (current_duplicate or {}).get("id")
                != (staged_duplicate or {}).get("id")
            )
            resolution = row[7]
            if duplicate_changed:
                errors.append(
                    "Stan duplikatów zmienił się od czasu podglądu. "
                    "Sprawdź wiersz i ponownie wybierz decyzję."
                )
                resolution = "skip"
            cur.execute(
                """
                UPDATE general_import_rows
                SET normalized_data=%s::jsonb,validation_errors=%s::jsonb,
                    duplicate_data=%s::jsonb,resolution=%s
                WHERE id=%s
                """,
                (
                    json.dumps(data, ensure_ascii=False),
                    json.dumps(errors, ensure_ascii=False),
                    json.dumps(current_duplicate, ensure_ascii=False)
                    if current_duplicate else None,
                    resolution,
                    row[0],
                ),
            )
            validation_failed = validation_failed or bool(errors)
        if validation_failed:
            conn.commit()
            return redirect(f"/import-ogolny/{import_id}?validation=failed")

        counts = {"added": 0, "updated": 0, "skipped": 0, "errors": 0}
        product_snapshots = {
            duplicate_identity("product", row[5] or {})
            for row in rows
            if row[8] and row[3] == "product" and row[7] != "skip"
        }
        document_cache = {}
        order_cache = {}
        order = {
            "contractor": 0, "product": 1, "package": 2, "receipt": 3,
            "document": 3, "issue": 3, "shop_order": 4, "accounting": 5,
        }
        for row in sorted(rows, key=lambda item: (order.get(item[3], 99), item[0])):
            if not row[8] or row[7] == "skip":
                counts["skipped"] += 1
                continue
            data = dict(row[5] or {})
            if row[3] == "product":
                outcome = apply_import_product(cur, row, data)
            elif row[3] == "package":
                outcome = apply_import_package(cur, row, data, product_snapshots, import_id)
            elif row[3] == "contractor":
                outcome = apply_import_contractor(cur, row, data)
            elif row[3] in {"issue", "receipt", "document"}:
                outcome = apply_import_document(cur, row, data, import_id, document_cache)
            elif row[3] == "shop_order":
                outcome = apply_import_shop_order(
                    cur, row, data, import_id, order_cache
                )
            elif row[3] == "accounting":
                outcome = apply_import_accounting(cur, row, data)
            else:
                counts["skipped"] += 1
                continue
            counts[outcome] += 1
        cur.execute(
            """
            UPDATE general_imports
            SET status='completed',summary=%s::jsonb,errors='[]'::jsonb,completed_at=NOW()
            WHERE id=%s
            """,
            (json.dumps(counts, ensure_ascii=False), import_id),
        )
        log_action("general_import.completed", "general_import", import_id, counts, conn)
        conn.commit()
        cache.clear()
        return redirect(f"/import-ogolny/{import_id}?completed=1")
    except ValueError as exc:
        conn.rollback()
        logger.warning("General import validation failed during commit.", exc_info=True)
        error_conn = None
        try:
            error_conn = db()
            error_cur = error_conn.cursor()
            error_cur.execute(
                "UPDATE general_imports SET errors=%s::jsonb WHERE id=%s",
                (json.dumps([str(exc)], ensure_ascii=False), import_id),
            )
            error_conn.commit()
        except Exception:
            logger.exception("General import error history update failed.")
        finally:
            if error_conn:
                error_conn.close()
        return redirect(f"/import-ogolny/{import_id}?commit_error=1")
    except Exception:
        conn.rollback()
        logger.exception("General import commit failed.")
        return "Import nie został zapisany. Żadne dane nie zostały zmienione.", 500
    finally:
        conn.close()


# 📤 WYDANIE
@app.route('/import-wydan')
@login_required
def issue_import_page():
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    if session.get("role") == "admin":
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at,
                   completed_at,undone_at,undone_by
            FROM issue_imports ORDER BY id DESC LIMIT 100
            """
        )
    else:
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at,
                   completed_at,undone_at,undone_by
            FROM issue_imports
            WHERE lower(imported_by)=lower(%s)
            ORDER BY id DESC LIMIT 100
            """,
            (session.get("user"),),
        )
    history = cur.fetchall()
    allow_general_stock = issue_import_allow_general_stock(cur)
    cur.execute(
        """
        SELECT id FROM issue_imports
        WHERE status='completed' AND undone_at IS NULL
        ORDER BY completed_at DESC,id DESC LIMIT 1
        """
    )
    latest_undo = cur.fetchone()
    conn.close()
    return render_template(
        "issue_import.html",
        history=history,
        allow_general_stock=allow_general_stock,
        latest_undo_id=latest_undo[0] if latest_undo else None,
    )


@app.route('/import-wydan/settings', methods=['POST'])
@login_required
@admin_required
def issue_import_settings():
    allowed = request.form.get("allow_general_stock") == "on"
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            INSERT INTO app_settings(key,value,updated_by,updated_at)
            VALUES ('issue_import_allow_general_stock',%s,%s,NOW())
            ON CONFLICT (key) DO UPDATE SET
                value=EXCLUDED.value,updated_by=EXCLUDED.updated_by,
                updated_at=NOW()
            """,
            ("true" if allowed else "false", session.get("user")),
        )
        conn.commit()
        return redirect("/import-wydan?settings=saved")
    except Exception:
        conn.rollback()
        logger.exception("Issue import settings update failed.")
        return "Nie udało się zapisać ustawień importu.", 500
    finally:
        conn.close()


@app.route('/import-wydan/upload', methods=['POST'])
@login_required
@limiter.limit("10 per hour")
def issue_import_upload():
    denied = issue_import_access_error()
    if denied:
        return denied
    upload = request.files.get("excel_file")
    if not upload or not upload.filename:
        return "Wybierz plik Excel.", 400
    if not upload.filename.lower().endswith(".xlsx"):
        return "Dozwolony jest wyłącznie plik .xlsx.", 400
    filename = secure_filename(upload.filename)[:200] or "wydania.xlsx"
    try:
        sheets = parse_workbook(upload.read())
    except ValueError as exc:
        return f"Nie udało się odczytać pliku: {exc}", 400
    except Exception:
        logger.warning("Issue import workbook parsing failed.", exc_info=True)
        return "Nie udało się odczytać pliku Excel.", 400
    if not any(
        sheet["entity_type"] != "ignored" and sheet["rows"]
        for sheet in sheets
    ):
        return "Plik nie zawiera danych do przygotowania.", 400
    conn = db()
    cur = conn.cursor()
    try:
        import_id = stage_issue_import(cur, filename, sheets)
        conn.commit()
        return redirect(f"/import-wydan/{import_id}")
    except Exception:
        conn.rollback()
        logger.exception("Issue import staging failed.")
        return "Nie udało się przygotować podglądu importu.", 500
    finally:
        conn.close()


@app.route('/import-wydan/<int:import_id>')
@login_required
def issue_import_preview(import_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    run = issue_import_run_query(cur, import_id)
    if not issue_import_run_allowed(run):
        conn.close()
        return "Nie znaleziono importu.", 404
    rows = issue_import_rows_query(cur, import_id)
    cur.execute(
        """
        SELECT pk.id,pk.number,pk.qty,pk.warehouse,p.name,p.unit
        FROM packages pk
        JOIN products p ON p.id=pk.product_id
        WHERE pk.status='active' AND pk.qty>0
        ORDER BY pk.warehouse,lower(p.name),lower(pk.number)
        """
    )
    packages = cur.fetchall()
    allow_general_stock = issue_import_allow_general_stock(cur)
    conn.close()
    return render_template(
        "issue_import_preview.html",
        run=run,
        rows=rows,
        fields=ISSUE_IMPORT_FIELDS,
        field_labels=ISSUE_IMPORT_LABELS,
        units=sorted(UNITS),
        warehouses=WAREHOUSES,
        packages=packages,
        allow_general_stock=allow_general_stock,
    )


@app.route('/import-wydan/<int:import_id>/mapping', methods=['POST'])
@login_required
def issue_import_mapping(import_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = issue_import_run_query(cur, import_id, for_update=True)
        if not issue_import_run_allowed(run):
            return "Nie znaleziono importu.", 404
        if run[3] != "draft":
            raise ValueError("Zakończonego importu nie można zmieniać.")
        metadata = list(run[4] or [])
        allow_general_stock = issue_import_allow_general_stock(cur)
        for index, sheet in enumerate(metadata):
            selected = request.form.get(f"selected_{index}") == "on"
            mapping = {
                field: request.form.get(f"map_{index}_{field}", "").strip()
                for field in ISSUE_IMPORT_FIELDS
                if request.form.get(f"map_{index}_{field}", "").strip()
            }
            sheet["selected"] = selected
            sheet["mapping"] = mapping
            cur.execute(
                """
                SELECT id,source_data,resolution
                FROM issue_import_rows
                WHERE import_id=%s AND sheet_name=%s
                ORDER BY id FOR UPDATE
                """,
                (import_id, sheet["name"]),
            )
            for row_id, source_data, current_resolution in cur.fetchall():
                data = (
                    normalize_issue_row(source_data or {}, mapping)
                    if selected
                    else {field: "" for field in ISSUE_IMPORT_FIELDS}
                )
                duplicate = (
                    detect_issue_import_duplicate(cur, data) if selected else None
                )
                resolution = (
                    current_resolution
                    if duplicate and current_resolution in {"skip", "update", "new"}
                    else ("skip" if duplicate else "new")
                )
                errors, _context = (
                    issue_import_row_context(
                        cur,
                        data,
                        allow_general_stock,
                        duplicate,
                        resolution,
                    )
                    if selected
                    else (["Arkusz nie jest wybrany do importu."], {})
                )
                cur.execute(
                    """
                    UPDATE issue_import_rows
                    SET normalized_data=%s::jsonb,duplicate_data=%s::jsonb,
                        resolution=%s,included=%s,validation_errors=%s::jsonb
                    WHERE id=%s
                    """,
                    (
                        json.dumps(data, ensure_ascii=False),
                        json.dumps(duplicate, ensure_ascii=False)
                        if duplicate
                        else None,
                        resolution,
                        selected,
                        json.dumps(errors, ensure_ascii=False),
                        row_id,
                    ),
                )
        cur.execute(
            "UPDATE issue_imports SET detected_sheets=%s::jsonb WHERE id=%s",
            (json.dumps(metadata, ensure_ascii=False), import_id),
        )
        conn.commit()
        return redirect(f"/import-wydan/{import_id}?mapping=saved")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Issue import mapping update failed.")
        return "Nie udało się zapisać przypisania kolumn.", 500
    finally:
        conn.close()


@app.route('/import-wydan/<int:import_id>/prepare', methods=['POST'])
@login_required
def issue_import_prepare(import_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = issue_import_run_query(cur, import_id, for_update=True)
        if not issue_import_run_allowed(run):
            return "Nie znaleziono importu.", 404
        if run[3] != "draft":
            raise ValueError("Zakończonego importu nie można zmieniać.")
        rows = issue_import_rows_query(cur, import_id, for_update=True)
        allow_general_stock = issue_import_allow_general_stock(cur)
        for row in rows:
            included = request.form.get(f"included_{row[0]}") == "on"
            resolution = request.form.get(
                f"resolution_{row[0]}", row[6]
            )
            if resolution not in {"skip", "update", "new"}:
                resolution = "skip"
            data = {
                field: request.form.get(f"row_{row[0]}_{field}", "").strip()
                for field in ISSUE_IMPORT_FIELDS
            }
            data["package_id"] = request.form.get(
                f"row_{row[0]}_package_id", ""
            ).strip()
            duplicate = detect_issue_import_duplicate(cur, data)
            errors, _context = (
                issue_import_row_context(
                    cur,
                    data,
                    allow_general_stock,
                    duplicate,
                    resolution,
                )
                if included and resolution != "skip"
                else ([], {})
            )
            cur.execute(
                """
                UPDATE issue_import_rows
                SET normalized_data=%s::jsonb,duplicate_data=%s::jsonb,
                    resolution=%s,included=%s,validation_errors=%s::jsonb
                WHERE id=%s
                """,
                (
                    json.dumps(data, ensure_ascii=False),
                    json.dumps(duplicate, ensure_ascii=False)
                    if duplicate
                    else None,
                    resolution,
                    included,
                    json.dumps(errors, ensure_ascii=False),
                    row[0],
                ),
            )
        conn.commit()
        return redirect(f"/import-wydan/{import_id}?prepared=1")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Issue import preparation failed.")
        return "Nie udało się zapisać korekt.", 500
    finally:
        conn.close()


@app.route('/import-wydan/<int:import_id>/rows', methods=['POST'])
@login_required
def issue_import_add_row(import_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = issue_import_run_query(cur, import_id, for_update=True)
        if not issue_import_run_allowed(run):
            return "Nie znaleziono importu.", 404
        if run[3] != "draft":
            raise ValueError("Zakończonego importu nie można zmieniać.")
        data = {
            field: request.form.get(field, "").strip()
            for field in ISSUE_IMPORT_FIELDS
        }
        data["package_id"] = request.form.get("package_id", "").strip()
        duplicate = detect_issue_import_duplicate(cur, data)
        resolution = "skip" if duplicate else "new"
        errors, _context = issue_import_row_context(
            cur,
            data,
            issue_import_allow_general_stock(cur),
            duplicate,
            resolution,
        )
        cur.execute(
            """
            SELECT COALESCE(MAX(row_number),1)+1
            FROM issue_import_rows WHERE import_id=%s
            """,
            (import_id,),
        )
        row_number = cur.fetchone()[0]
        cur.execute(
            """
            INSERT INTO issue_import_rows(
                import_id,sheet_name,row_number,source_data,normalized_data,
                duplicate_data,resolution,included,validation_errors
            ) VALUES (%s,'Dodane ręcznie',%s,%s::jsonb,%s::jsonb,%s::jsonb,%s,TRUE,%s::jsonb)
            """,
            (
                import_id,
                row_number,
                json.dumps(data, ensure_ascii=False),
                json.dumps(data, ensure_ascii=False),
                json.dumps(duplicate, ensure_ascii=False)
                if duplicate
                else None,
                resolution,
                json.dumps(errors, ensure_ascii=False),
            ),
        )
        conn.commit()
        return redirect(f"/import-wydan/{import_id}?row=added")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Issue import row creation failed.")
        return "Nie udało się dodać wiersza.", 500
    finally:
        conn.close()


@app.route(
    '/import-wydan/<int:import_id>/rows/<int:row_id>/delete',
    methods=['POST'],
)
@login_required
def issue_import_delete_row(import_id, row_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = issue_import_run_query(cur, import_id, for_update=True)
        if not issue_import_run_allowed(run):
            return "Nie znaleziono importu.", 404
        if run[3] != "draft":
            raise ValueError("Zakończonego importu nie można zmieniać.")
        cur.execute(
            "DELETE FROM issue_import_rows WHERE id=%s AND import_id=%s",
            (row_id, import_id),
        )
        if cur.rowcount != 1:
            raise ValueError("Nie znaleziono wiersza.")
        conn.commit()
        return redirect(f"/import-wydan/{import_id}?row=deleted")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Issue import row deletion failed.")
        return "Nie udało się usunąć wiersza.", 500
    finally:
        conn.close()


@app.route('/import-wydan/<int:import_id>/confirm', methods=['POST'])
@login_required
def issue_import_confirm(import_id):
    denied = issue_import_access_error()
    if denied:
        return denied
    conn = db()
    cur = conn.cursor()
    try:
        run = issue_import_run_query(cur, import_id, for_update=True)
        if not issue_import_run_allowed(run):
            return "Nie znaleziono importu.", 404
        if run[3] != "draft":
            raise ValueError("Ten import został już zakończony.")
        rows = issue_import_rows_query(cur, import_id, for_update=True)
        allow_general_stock = issue_import_allow_general_stock(cur)
        prepared = []
        validation_failed = False
        for row in rows:
            if not row[7] or row[6] == "skip":
                continue
            data = dict(row[4] or {})
            duplicate = detect_issue_import_duplicate(cur, data)
            errors, context = issue_import_row_context(
                cur,
                data,
                allow_general_stock,
                duplicate,
                row[6],
            )
            cur.execute(
                """
                UPDATE issue_import_rows
                SET duplicate_data=%s::jsonb,validation_errors=%s::jsonb
                WHERE id=%s
                """,
                (
                    json.dumps(duplicate, ensure_ascii=False)
                    if duplicate
                    else None,
                    json.dumps(errors, ensure_ascii=False),
                    row[0],
                ),
            )
            if errors:
                validation_failed = True
            effective_row = list(row)
            effective_row[5] = duplicate
            prepared.append((tuple(effective_row), data, context))
        if validation_failed:
            conn.commit()
            return redirect(f"/import-wydan/{import_id}?validation=failed")
        counts = {"issues": 0, "updated": 0, "skipped": 0, "errors": 0}
        counts["skipped"] = sum(
            1 for row in rows if not row[7] or row[6] == "skip"
        )
        document_cache = {}
        for row, data, context in prepared:
            outcome = apply_issue_import_row(
                cur,
                import_id,
                row,
                data,
                context,
                document_cache,
            )
            counts["issues"] += 1
            if outcome == "updated":
                counts["updated"] += 1
        cur.execute(
            """
            UPDATE issue_imports
            SET status='completed',summary=%s::jsonb,errors='[]'::jsonb,
                completed_at=NOW()
            WHERE id=%s
            """,
            (json.dumps(counts, ensure_ascii=False), import_id),
        )
        log_action(
            "issue_import.completed",
            "issue_import",
            import_id,
            counts,
            conn,
        )
        conn.commit()
        cache.clear()
        return redirect(f"/import-wydan/{import_id}?completed=1")
    except ValueError as exc:
        conn.rollback()
        try:
            cur = conn.cursor()
            cur.execute(
                "UPDATE issue_imports SET errors=%s::jsonb WHERE id=%s",
                (json.dumps([str(exc)], ensure_ascii=False), import_id),
            )
            conn.commit()
        except Exception:
            conn.rollback()
            logger.exception("Issue import validation history update failed.")
        return redirect(f"/import-wydan/{import_id}?commit_error=1")
    except Exception:
        conn.rollback()
        logger.exception("Issue import commit failed.")
        user_message = (
            "Nie udało się zapisać importu. Żadne stany magazynowe nie zostały "
            "zmienione."
        )
        try:
            cur = conn.cursor()
            cur.execute(
                "UPDATE issue_imports SET errors=%s::jsonb WHERE id=%s",
                (json.dumps([user_message], ensure_ascii=False), import_id),
            )
            conn.commit()
        except Exception:
            conn.rollback()
        return redirect(f"/import-wydan/{import_id}?commit_error=1")
    finally:
        conn.close()


@app.route('/import-wydan/undo-last', methods=['POST'])
@login_required
@admin_required
def issue_import_undo_last():
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT id FROM issue_imports
            WHERE status='completed' AND undone_at IS NULL
            ORDER BY completed_at DESC,id DESC
            LIMIT 1 FOR UPDATE
            """
        )
        row = cur.fetchone()
        if not row:
            raise ValueError("Brak zakończonego importu do cofnięcia.")
        undo_issue_import(cur, row[0], session.get("user", "system"))
        log_action(
            "issue_import.undone",
            "issue_import",
            row[0],
            {"status": "undone"},
            conn,
        )
        conn.commit()
        cache.clear()
        return redirect(f"/import-wydan?undone={row[0]}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 409
    except Exception:
        conn.rollback()
        logger.exception("Issue import undo failed.")
        return "Nie udało się cofnąć importu. Żadne stany nie zostały zmienione.", 500
    finally:
        conn.close()


@app.route('/import-wydan/history.<fmt>')
@login_required
def issue_import_history_export(fmt):
    denied = issue_import_access_error()
    if denied:
        return denied
    if fmt not in {"xlsx", "pdf"}:
        return "Nieprawidłowy format eksportu.", 400
    conn = db()
    cur = conn.cursor()
    if session.get("role") == "admin":
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at
            FROM issue_imports ORDER BY id DESC
            """
        )
    else:
        cur.execute(
            """
            SELECT id,filename,imported_by,status,summary,errors,created_at
            FROM issue_imports WHERE lower(imported_by)=lower(%s)
            ORDER BY id DESC
            """,
            (session.get("user"),),
        )
    rows = cur.fetchall()
    conn.close()
    if fmt == "xlsx":
        data = issue_history_xlsx(rows)
        mimetype = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        data = issue_history_pdf(rows)
        mimetype = "application/pdf"
    return Response(
        data,
        mimetype=mimetype,
        headers={
            "Content-Disposition": (
                f'attachment; filename="historia-importow-wydan.{fmt}"'
            )
        },
    )


@app.route('/wydanie')
@login_required
def wydanie():
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM products ORDER BY warehouse, lower(name), id")
    products = cur.fetchall()

    conn.close()

    return render_template("wydanie.html", products=products)


@app.route('/inwestycja-suwaj', endpoint='inwestycja_suwaj_page_view')
@login_required
def inwestycja_suwaj():
    return render_template("inwestycja_suwaj.html", warehouse=INVESTMENT_WAREHOUSE)


@app.route('/inwestycja-suwaj/magazyn')
@login_required
def inwestycja_suwaj_magazyn():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM products WHERE warehouse=%s", (INVESTMENT_WAREHOUSE,))
    products = cur.fetchall()
    conn.close()
    return render_template(
        "index.html",
        products=products,
        warehouse=INVESTMENT_WAREHOUSE,
        package_modes={},
        warehouses=WAREHOUSES,
        units=sorted(UNITS),
    )


@app.route('/inwestycja-suwaj/przyjecie')
@login_required
def inwestycja_suwaj_przyjecie():
    return render_receipt_form(forced_warehouse=INVESTMENT_WAREHOUSE)


@app.route('/inwestycja-suwaj/receive_doc', methods=['POST'])
@login_required
def inwestycja_suwaj_receive_doc():
    return create_receipt(INVESTMENT_WAREHOUSE)

    conn = db()
    cur = conn.cursor()

    date = datetime.now().strftime("%Y-%m-%d")
    kontrahent = request.form.get('kontrahent')

    cur.execute("SELECT COUNT(*) FROM issue_docs WHERE warehouse=%s", (INVESTMENT_WAREHOUSE,))
    num = cur.fetchone()[0] + 1
    doc_number = f"PZ-IS/{num}/{datetime.now().year}"

    cur.execute("""
        INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number)
        VALUES (%s,%s,%s,%s,%s) RETURNING id
    """, (date, kontrahent, INVESTMENT_WAREHOUSE, "", doc_number))
    doc_id = cur.fetchone()[0]

    product_ids = request.form.getlist('product_id')
    qtys = request.form.getlist('qty')
    package_numbers = request.form.getlist('package_number')
    prices_netto = request.form.getlist('price_netto')
    prices_brutto = request.form.getlist('price_brutto')

    for i in range(len(product_ids)):
        if not product_ids[i]:
            continue
        pid = int(product_ids[i])
        try:
            qty = float(qtys[i].replace(",", "."))
        except:
            qty = 0
        if qty <= 0:
            continue
        try:
            price_netto = float(prices_netto[i].replace(",", "."))
        except:
            price_netto = 0
        try:
            price_brutto = float(prices_brutto[i].replace(",", "."))
        except:
            price_brutto = 0

        cur.execute("""
            UPDATE products
            SET qty = qty + %s
            WHERE id=%s AND warehouse=%s
        """, (qty, pid, INVESTMENT_WAREHOUSE))
        updated = cur.rowcount

        # jeśli produktu nie ma jeszcze w magazynie inwestycji, sklonuj kartotekę i dodaj stan
        if updated == 0:
            cur.execute("""
                SELECT name, unit, price_netto, vat
                FROM products
                WHERE id=%s
            """, (pid,))
            source = cur.fetchone()

            if source:
                cur.execute("""
                    INSERT INTO products(name, qty, unit, warehouse, price_netto, vat)
                    VALUES (%s,%s,%s,%s,%s,%s)
                    RETURNING id
                """, (source[0], qty, source[1], INVESTMENT_WAREHOUSE, source[2], source[3]))
                pid = cur.fetchone()[0]
            else:
                continue

        cur.execute("""
            INSERT INTO issue_items(doc_id, product_id, qty, warehouse, price_netto, price_brutto)
            VALUES (%s,%s,%s,%s,%s,%s)
        """, (doc_id, pid, qty, INVESTMENT_WAREHOUSE, price_netto, price_brutto))

        if i < len(package_numbers) and package_numbers[i]:
            cur.execute("""
                INSERT INTO packages(product_id, number, qty, warehouse)
                VALUES (%s,%s,%s,%s)
            """, (pid, package_numbers[i], qty, INVESTMENT_WAREHOUSE))

    conn.commit()
    conn.close()

    return redirect('/inwestycja-suwaj/magazyn')


@app.route('/inwestycja-suwaj/wydanie')
@login_required
def inwestycja_suwaj_wydanie():
    conn = db()
    cur = conn.cursor()

    cur.execute(
        "SELECT * FROM products WHERE warehouse=%s ORDER BY lower(name), id",
        (INVESTMENT_WAREHOUSE,),
    )
    products = cur.fetchall()
    conn.close()

    return render_template(
        "wydanie.html",
        products=products,
        forced_warehouse=INVESTMENT_WAREHOUSE,
        form_action="/inwestycja-suwaj/issue_doc",
        page_title="📄 Wydanie (WZ) – Inwestycja Suwaj"
    )


@app.route('/inwestycja-suwaj/issue_doc', methods=['POST'])
@login_required
def inwestycja_suwaj_issue_doc():
    return create_issue(INVESTMENT_WAREHOUSE)

    conn = db()
    cur = conn.cursor()

    date = datetime.now().strftime("%Y-%m-%d")
    kontrahent = request.form.get('kontrahent')

    cur.execute("SELECT COUNT(*) FROM issue_docs WHERE warehouse=%s", (INVESTMENT_WAREHOUSE,))
    num = cur.fetchone()[0] + 1
    doc_number = f"WZ-IS/{num}/{datetime.now().year}"

    cur.execute("""
        INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number)
        VALUES (%s,%s,%s,%s,%s) RETURNING id
    """, (date, kontrahent, INVESTMENT_WAREHOUSE, "", doc_number))
    doc_id = cur.fetchone()[0]

    product_ids = request.form.getlist('product_id')
    qtys = request.form.getlist('qty')
    package_ids = request.form.getlist('package_id')
    prices_netto = request.form.getlist('price_netto')
    prices_brutto = request.form.getlist('price_brutto')

    for i in range(len(product_ids)):
        if not product_ids[i]:
            continue

        pid = int(product_ids[i])
        try:
            qty = float(qtys[i].replace(",", "."))
        except:
            qty = 0
        if qty <= 0:
            continue
        try:
            price_netto = float(prices_netto[i].replace(",", "."))
        except:
            price_netto = 0
        try:
            price_brutto = float(prices_brutto[i].replace(",", "."))
        except:
            price_brutto = 0

        pkg = package_ids[i] if package_ids[i] else None
        if pkg:
            pkg = int(pkg)
            cur.execute("""
                UPDATE packages
                SET qty = qty - %s
                WHERE id=%s AND (warehouse=%s OR warehouse IS NULL) AND qty >= %s
            """, (qty, pkg, INVESTMENT_WAREHOUSE, qty))
            if cur.rowcount == 0:
                conn.rollback()
                conn.close()
                return "Brak w paczce"

        cur.execute("""
            UPDATE products
            SET qty = qty - %s
            WHERE id=%s AND warehouse=%s AND qty >= %s
        """, (qty, pid, INVESTMENT_WAREHOUSE, qty))

        if cur.rowcount == 0:
            conn.rollback()
            conn.close()
            return f"Brak stanu w magazynie {INVESTMENT_WAREHOUSE}"

        cur.execute("""
            INSERT INTO issue_items(doc_id, product_id, qty, warehouse, package_id, price_netto, price_brutto)
            VALUES (%s,%s,%s,%s,%s,%s,%s)
        """, (doc_id, pid, qty, INVESTMENT_WAREHOUSE, pkg, price_netto, price_brutto))

    cur.execute("DELETE FROM packages WHERE qty <= 0")
    conn.commit()
    conn.close()

    return redirect(f"/doc/{doc_id}")


# 📤 ZAPIS WYDANIA (PRO)
@app.route('/issue_doc', methods=['POST'])
@login_required
def issue_doc():
    return create_issue()

    conn = db()
    cur = conn.cursor()

    date = datetime.now().strftime("%Y-%m-%d")
    kontrahent = request.form.get('kontrahent')

    cur.execute("SELECT COUNT(*) FROM issue_docs")
    num = cur.fetchone()[0] + 1
    doc_number = f"WZ/{num}/{datetime.now().year}"

    cur.execute("""
        INSERT INTO issue_docs(date, kontrahent, warehouse, image, doc_number)
        VALUES (%s,%s,%s,%s,%s) RETURNING id
    """, (date, kontrahent, "", "", doc_number))

    doc_id = cur.fetchone()[0]

    product_ids = request.form.getlist('product_id')
    qtys = request.form.getlist('qty')
    warehouses = request.form.getlist('warehouse')
    package_ids = request.form.getlist('package_id')
    prices_netto = request.form.getlist('price_netto')
    prices_brutto = request.form.getlist('price_brutto')

    for i in range(len(product_ids)):
        if not product_ids[i]:
            continue

        pid = int(product_ids[i])
        wh = warehouses[i]

        try:
            qty = float(qtys[i].replace(",", "."))
        except:
            qty = 0

        if qty <= 0:
            continue
        try:
            price_netto = float(prices_netto[i].replace(",", "."))
        except:
            price_netto = 0
        try:
            price_brutto = float(prices_brutto[i].replace(",", "."))
        except:
            price_brutto = 0

        pkg = package_ids[i] if package_ids[i] else None

        # pakiet
        if pkg:
            pkg = int(pkg)

            cur.execute("""
                UPDATE packages
                SET qty = qty - %s
                WHERE id=%s AND warehouse=%s AND qty >= %s
            """, (qty, pkg, wh, qty))
            if cur.rowcount == 0:
                conn.rollback()
                conn.close()
                return "Brak w paczce"

        # ✅ stan -
        cur.execute("""
            UPDATE products
            SET qty = qty - %s
            WHERE id=%s AND warehouse=%s AND qty >= %s
        """, (qty, pid, wh, qty))
        if cur.rowcount == 0:
            conn.rollback()
            conn.close()
            return f"Brak stanu w magazynie {wh}"

        cur.execute("""
            INSERT INTO issue_items(doc_id, product_id, qty, warehouse, package_id, price_netto, price_brutto)
            VALUES (%s,%s,%s,%s,%s,%s,%s)
        """, (doc_id, pid, qty, wh, pkg, price_netto, price_brutto))

    # czyść puste paczki
    cur.execute("DELETE FROM packages WHERE qty <= 0")

    conn.commit()
    conn.close()

    return redirect(f"/doc/{doc_id}")


# 📄 SZCZEGÓŁ
@app.route('/doc/<int:id>')
@login_required
def doc_detail(id):
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM issue_docs WHERE id=%s", (id,))
    doc = cur.fetchone()
    if not doc:
        conn.close()
        return "Nie znaleziono dokumentu.", 404

    cur.execute("""
        SELECT p.name, i.qty, COALESCE(i.warehouse, p.warehouse),
               COALESCE(i.package_number, pk.number),
               i.dimension,i.species,i.notes
        FROM issue_items i
        JOIN products p ON p.id = i.product_id
        LEFT JOIN packages pk ON pk.id = i.package_id
        WHERE i.doc_id=%s
    """, (id,))
    items = cur.fetchall()
    issue_document = is_issue_document(doc[6] if len(doc) > 6 else None, doc[5])
    photos = []
    history = []
    if issue_document:
        cur.execute(
            """
            SELECT id, filename, content_type, note, added_by, created_at
            FROM issue_doc_photos
            WHERE doc_id=%s
            ORDER BY created_at DESC, id DESC
            """,
            (id,),
        )
        photos = cur.fetchall()
        cur.execute(
            """
            SELECT user_email, action, details, created_at
            FROM issue_doc_history
            WHERE doc_id=%s
            ORDER BY created_at DESC, id DESC
            """,
            (id,),
        )
        history = cur.fetchall()

    conn.close()

    return render_template(
        "doc_detail.html",
        doc=doc,
        items=items,
        issue_document=issue_document,
        photos=photos,
        history=history,
    )



@app.route('/doc/<int:id>/photos', methods=['POST'])
@login_required
def add_issue_doc_photos(id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT movement_type, doc_number FROM issue_docs WHERE id=%s", (id,))
        doc = cur.fetchone()
        if not doc:
            return "Nie znaleziono dokumentu.", 404
        if not is_issue_document(doc[0], doc[1]):
            return "Zdjęcia można dodawać tylko do dokumentów WZ/RW.", 400
        saved = save_issue_photos(cur, id, request.files.getlist("photos"), request.form.get("photo_note"))
        if not saved:
            return "Dodaj co najmniej jedno zdjęcie.", 400
        conn.commit()
        cache.clear()
        return redirect(f"/doc/{id}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Issue document photo upload failed.")
        return "Nie udało się zapisać zdjęć.", 500
    finally:
        conn.close()


@app.route('/doc-photo/<int:photo_id>')
@login_required
def view_issue_doc_photo(photo_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT filename, content_type, data FROM issue_doc_photos WHERE id=%s", (photo_id,))
    photo = cur.fetchone()
    conn.close()
    if not photo:
        return "Nie znaleziono zdjęcia.", 404
    return Response(bytes(photo[2]), mimetype=photo[1], headers={"Content-Disposition": f"inline; filename=\"{photo[0]}\""})


@app.route('/doc-photo/<int:photo_id>/delete', methods=['POST'])
@admin_required
def delete_issue_doc_photo(photo_id):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT doc_id, filename FROM issue_doc_photos WHERE id=%s FOR UPDATE", (photo_id,))
    photo = cur.fetchone()
    if not photo:
        conn.close()
        return "Nie znaleziono zdjęcia.", 404
    cur.execute("DELETE FROM issue_doc_photos WHERE id=%s", (photo_id,))
    issue_doc_history(cur, photo[0], "usunięto zdjęcie", photo[1])
    conn.commit()
    cache.clear()
    conn.close()
    return redirect(f"/doc/{photo[0]}")

@app.route('/doc/<int:id>/edit', methods=['POST'])
@login_required
def edit_doc(id):
    contractor = (request.form.get("kontrahent") or "").strip()
    try:
        date = normalized_document_date(request.form.get("date"))
    except ValueError as exc:
        return str(exc), 400
    if not contractor:
        return "Kontrahent jest wymagany.", 400
    conn = db()
    cur = conn.cursor()
    cur.execute("""
        UPDATE issue_docs
        SET kontrahent=%s, date=%s
        WHERE id=%s AND voided_at IS NULL
    """, (contractor, date, id))
    conn.commit()
    conn.close()
    cache.clear()
    return redirect(f"/doc/{id}")


def void_document(document_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT movement_type, doc_number, voided_at
            FROM issue_docs WHERE id=%s FOR UPDATE
            """,
            (document_id,),
        )
        document = cur.fetchone()
        if not document:
            return redirect("/historia")
        if document[2]:
            return redirect(f"/doc/{document_id}")
        document_number = str(document[1] or "").upper()
        movement_type = document[0] or (
            "RW" if document_number.startswith("RW")
            else "WZ" if document_number.startswith("WZ")
            else "PZ"
        )
        cur.execute(
            """
            SELECT product_id, qty, warehouse, package_id
            FROM issue_items WHERE doc_id=%s ORDER BY id FOR UPDATE
            """,
            (document_id,),
        )
        items = cur.fetchall()
        if movement_type in {"WZ", "RW"}:
            for product_id, qty, warehouse, package_id in items:
                cur.execute(
                    "UPDATE products SET qty=qty+%s WHERE id=%s AND warehouse=%s",
                    (qty, product_id, warehouse),
                )
                if package_id:
                    cur.execute(
                        """
                        UPDATE packages
                        SET qty=qty+%s, status='active', archived_at=NULL
                        WHERE id=%s
                        """,
                        (qty, package_id),
                    )
        else:
            for product_id, qty, warehouse, package_id in items:
                cur.execute(
                    """
                    UPDATE products SET qty=qty-%s
                    WHERE id=%s AND warehouse=%s AND qty >= %s
                    """,
                    (qty, product_id, warehouse, qty),
                )
                if cur.rowcount != 1:
                    raise ValueError("Nie można anulować przyjęcia: część towaru została już wydana.")
                if package_id:
                    cur.execute(
                        "SELECT qty FROM packages WHERE id=%s FOR UPDATE",
                        (package_id,),
                    )
                    package = cur.fetchone()
                    if not package or package[0] + 1e-9 < qty:
                        raise ValueError("Nie można anulować przyjęcia: paczka została już częściowo wydana.")
                    cur.execute(
                        """
                        UPDATE packages
                        SET qty=0, status='cancelled', archived_at=NOW()
                        WHERE id=%s
                        """,
                        (package_id,),
                    )
        cur.execute(
            "UPDATE issue_docs SET voided_at=NOW(), voided_by=%s WHERE id=%s",
            (session.get("user"), document_id),
        )
        log_action("document.voided", "issue_doc", document_id, {"movement_type": movement_type, "doc_number": document[1]}, conn)
        conn.commit()
        cache.clear()
        return redirect(f"/doc/{document_id}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 409
    except Exception:
        conn.rollback()
        logger.exception("Document cancellation failed.")
        return "Nie udało się anulować dokumentu.", 500
    finally:
        conn.close()


@app.route('/doc/<int:id>/delete', methods=['POST'])
@admin_required
def delete_doc(id):
    return void_document(id)

    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT id, date, kontrahent, warehouse, image, doc_number FROM issue_docs WHERE id=%s", (id,))
    doc = cur.fetchone()
    if not doc:
        conn.close()
        return redirect('/historia')

    cur.execute("""
        SELECT product_id, qty, warehouse, package_id
        FROM issue_items
        WHERE doc_id=%s
    """, (id,))
    items = cur.fetchall()

    is_issue_doc = str(doc[5] or "").startswith("WZ")

    for product_id, qty, warehouse, package_id in items:
        wh = warehouse or ""
        sign = 1 if is_issue_doc else -1
        cur.execute("""
            UPDATE products
            SET qty = qty + %s
            WHERE id=%s AND warehouse=%s
        """, (sign * qty, product_id, wh))

        if package_id and is_issue_doc:
            cur.execute("UPDATE packages SET qty = qty + %s WHERE id=%s", (qty, package_id))

    cur.execute("DELETE FROM issue_items WHERE doc_id=%s", (id,))
    cur.execute("DELETE FROM issue_docs WHERE id=%s", (id,))
    cur.execute("DELETE FROM packages WHERE qty <= 0")
    conn.commit()
    conn.close()
    return redirect('/historia')


def backup_bucket():
    bucket_name = FIREBASE_CONFIG.get("storageBucket")
    if not bucket_name:
        raise RuntimeError("Brak FIREBASE_STORAGE_BUCKET.")
    return storage.bucket(bucket_name)


def perform_database_backup(created_by):
    guard_conn = db()
    try:
        guard_cur = guard_conn.cursor()
        guard_cur.execute("SELECT pg_try_advisory_lock(67431031)")
        if not guard_cur.fetchone()[0]:
            raise RuntimeError("Inna kopia zapasowa jest już wykonywana.")
        return _perform_database_backup(created_by)
    finally:
        try:
            if "guard_cur" in locals():
                guard_cur.execute("SELECT pg_advisory_unlock(67431031)")
                guard_conn.commit()
        except Exception:
            logger.warning("Database backup advisory lock cleanup failed.", exc_info=True)
        finally:
            guard_conn.close()


def _perform_database_backup(created_by):
    ensure_db_initialized()
    log_conn = db()
    try:
        log_cur = log_conn.cursor()
        log_cur.execute(
            """
            INSERT INTO backup_runs(created_by, status)
            VALUES (%s, 'running') RETURNING id
            """,
            (created_by,),
        )
        run_id = log_cur.fetchone()[0]
        log_conn.commit()
    finally:
        log_conn.close()
    try:
        export_conn = db()
        try:
            compressed, checksum, _ = export_database(export_conn)
        finally:
            export_conn.close()
        encrypted = encrypt_backup(
            compressed,
            os.environ.get("BACKUP_ENCRYPTION_KEY", ""),
        )
        timestamp = datetime.utcnow().strftime("%Y/%m/%Y-%m-%dT%H-%M-%SZ")
        object_name = f"database-backups/{timestamp}-{run_id}.json.gz.enc"
        blob = backup_bucket().blob(object_name)
        blob.metadata = {
            "sha256": checksum,
            "format": "magazyn-app-backup-v1",
            "created-by": created_by[:100],
        }
        blob.upload_from_string(
            encrypted,
            content_type="application/octet-stream",
        )
        update_conn = db()
        try:
            update_cur = update_conn.cursor()
            update_cur.execute(
                """
                UPDATE backup_runs
                SET status='completed', object_name=%s, size_bytes=%s, checksum=%s
                WHERE id=%s
                """,
                (object_name, len(encrypted), checksum, run_id),
            )
            update_conn.commit()
        finally:
            update_conn.close()
        return {
            "id": run_id,
            "object_name": object_name,
            "size_bytes": len(encrypted),
            "checksum": checksum,
        }
    except Exception as exc:
        logger.exception("Database backup failed.")
        try:
            error_conn = db()
            try:
                error_cur = error_conn.cursor()
                error_cur.execute(
                    "UPDATE backup_runs SET status='failed', error=%s WHERE id=%s",
                    (str(exc)[:1000], run_id),
                )
                error_conn.commit()
            finally:
                error_conn.close()
        except Exception:
            logger.exception("Backup failure logging failed.")
        raise


def valid_backup_object_name(object_name):
    return (
        object_name.startswith("database-backups/")
        and object_name.endswith(".json.gz.enc")
        and ".." not in object_name
        and "\\" not in object_name
    )


@app.route('/admin/backups')
@admin_required
def backups_page():
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, created_at, created_by, status, object_name, size_bytes, checksum, error
        FROM backup_runs ORDER BY id DESC LIMIT 100
        """
    )
    runs = cur.fetchall()
    conn.close()
    return render_template(
        "backups.html",
        runs=runs,
        restore_enabled=os.environ.get("ALLOW_BACKUP_RESTORE", "").lower() == "true",
    )


@app.route('/admin/backups/create', methods=['POST'])
@admin_required
@limiter.limit("2 per hour")
def create_manual_backup():
    try:
        perform_database_backup(session.get("user") or "administrator")
    except Exception as exc:
        logger.exception("Manual database backup failed.")
        return redirect(
            "/admin/backups?failed=1&error="
            + quote_plus(str(exc)[:300] or "Nie udało się utworzyć kopii zapasowej.")
        )
    return redirect("/admin/backups?created=1")


@app.route('/admin/backups/download')
@admin_required
@limiter.limit("20 per hour")
def download_backup():
    object_name = (request.args.get("name") or "").strip()
    if not valid_backup_object_name(object_name):
        return "Nieprawidłowa nazwa kopii.", 400
    try:
        encrypted = backup_bucket().blob(object_name).download_as_bytes()
    except Exception:
        logger.exception("Backup download failed.")
        return "Nie udało się pobrać kopii.", 502
    filename = object_name.rsplit("/", 1)[-1]
    return Response(
        encrypted,
        mimetype="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.route('/admin/backups/restore', methods=['POST'])
@admin_required
@limiter.limit("1 per hour")
def restore_backup():
    if os.environ.get("ALLOW_BACKUP_RESTORE", "").lower() != "true":
        return "Przywracanie jest wyłączone. Włącz je czasowo w Renderze.", 403
    if request.form.get("confirmation") != "PRZYWRÓĆ":
        return "Wpisz PRZYWRÓĆ, aby potwierdzić.", 400
    object_name = (request.form.get("object_name") or "").strip()
    if not valid_backup_object_name(object_name):
        return "Nieprawidłowa nazwa kopii.", 400
    try:
        encrypted = backup_bucket().blob(object_name).download_as_bytes()
        compressed = decrypt_backup(
            encrypted,
            os.environ.get("BACKUP_ENCRYPTION_KEY", ""),
        )
        payload = parse_backup(compressed)
        restore_conn = db()
        try:
            restore_database(restore_conn, payload)
        except Exception:
            restore_conn.rollback()
            raise
        finally:
            restore_conn.close()
        cache.clear()
    except Exception:
        logger.exception("Database restore failed.")
        return "Nie udało się przywrócić kopii. Baza nie została zmieniona.", 500
    return redirect("/admin/backups?restored=1")


# 📊 HISTORIA
@app.route('/ksiegowosc')
@login_required
@accounting_required
def accounting_dashboard():
    filters = accounting_filters_from_request()
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT id FROM shop_orders")
        for (order_id,) in cur.fetchall():
            ensure_shop_accounting_row(cur, order_id)
        conn.commit()
        data = accounting_dashboard_data(cur, filters)
    finally:
        conn.close()
    return render_template(
        "accounting.html",
        payment_methods=ACCOUNTING_PAYMENT_METHODS,
        filters=filters,
        statuses=SHOP_STATUS_FLOW,
        **data,
    )


def accounting_filters_from_request():
    return {
        key: (request.args.get(key) or "").strip()
        for key in (
            "payment_status",
            "payment_method",
            "proforma_status",
            "invoice_status",
            "receipt_status",
            "sales_document_status",
            "order_status",
            "salesperson",
            "client",
            "proforma_number",
            "invoice_number",
            "receipt_number",
            "date",
        )
    }


def accounting_query_parts(filters):
    where = []
    params = []
    if filters.get("payment_status"):
        where.append("o.payment_status=%s")
        params.append(filters["payment_status"])
    if filters.get("payment_method"):
        where.append("a.payment_method=%s")
        params.append(filters["payment_method"])
    if filters.get("order_status"):
        where.append("o.status=%s")
        params.append(filters["order_status"])
    if filters.get("proforma_status") == "issued":
        where.append("a.proforma_issued=TRUE")
    elif filters.get("proforma_status") == "not_issued":
        where.append("a.proforma_issued=FALSE")
    elif filters.get("proforma_status") == "cancelled":
        where.append(
            "EXISTS (SELECT 1 FROM accounting_documents ad "
            "WHERE ad.order_id=o.id AND ad.document_type='proforma' "
            "AND ad.status='cancelled')"
        )
    document_filters = {
        "invoice_status": ("invoice_issued", "invoice"),
        "receipt_status": ("receipt_issued", "receipt"),
        "sales_document_status": ("sales_document_generated", "sales"),
    }
    for key, (column, document_type) in document_filters.items():
        value = filters.get(key)
        if value == "issued":
            where.append(f"a.{column}=TRUE")
        elif value == "not_issued":
            where.append(f"a.{column}=FALSE")
        elif value == "cancelled":
            where.append(
                "EXISTS (SELECT 1 FROM accounting_documents ad "
                f"WHERE ad.order_id=o.id AND ad.document_type='{document_type}' "
                "AND ad.status='cancelled')"
            )
    text_filters = (
        ("client", "o.customer_name"),
        ("proforma_number", "a.proforma_number"),
        ("invoice_number", "a.invoice_number"),
        ("receipt_number", "a.receipt_number"),
    )
    for key, column in text_filters:
        if filters.get(key):
            where.append(f"lower(COALESCE({column},'')) LIKE %s")
            params.append(f"%{filters[key].lower()}%")
    if filters.get("salesperson"):
        like = f"%{filters['salesperson'].lower()}%"
        where.append(
            "(lower(COALESCE(o.salesperson_name,'')) LIKE %s "
            "OR lower(COALESCE(o.salesperson_email,'')) LIKE %s "
            "OR lower(COALESCE(a.salesperson,'')) LIKE %s)"
        )
        params.extend([like, like, like])
    if filters.get("date"):
        where.append("o.order_date=%s")
        params.append(filters["date"])
    where_sql = (" WHERE " + " AND ".join(where)) if where else ""
    base = f"""
        FROM shop_orders o
        JOIN shop_accounting a ON a.order_id=o.id
        LEFT JOIN shop_sales_documents sd ON sd.order_id=o.id
        LEFT JOIN LATERAL (
            SELECT COALESCE(SUM(i.qty*i.price_brutto),0)+COALESCE(o.shipping_cost,0)
                   AS order_value
            FROM shop_order_items i WHERE i.order_id=o.id
        ) value ON TRUE
        {where_sql}
    """
    return base, tuple(params)


def accounting_dashboard_data(cur, filters):
    base, params = accounting_query_parts(filters)
    cur.execute(
        f"""
        SELECT o.id,o.order_number,o.order_date,o.customer_name,o.payment_status,o.status,
               a.payment_method,a.proforma_number,a.invoice_number,a.receipt_number,
               COALESCE(o.salesperson_name,o.salesperson_email,a.salesperson),
               value.order_value,a.amount_paid,a.amount_due,a.proforma_issued,
               a.invoice_issued,a.receipt_issued,a.sales_document_generated,
               a.document_sent,a.ready_to_ship,a.settled,a.paid,a.partial_payment,
               a.waiting_for_payment,sd.id,sd.voided_at,a.payment_due_date,
               a.document_to_warehouse,a.accounting_notes
        {base}
        ORDER BY o.order_date DESC,o.id DESC
        """,
        params,
    )
    orders = cur.fetchall()
    cur.execute(
        f"""
        SELECT
            COALESCE(SUM(value.order_value) FILTER (WHERE a.proforma_issued),0),
            COALESCE(SUM(value.order_value) FILTER (WHERE o.status<>'Anulowane'),0),
            COALESCE(SUM(a.amount_paid),0),
            COALESCE(SUM(a.amount_due),0),
            COUNT(*) FILTER (WHERE a.proforma_issued),
            COUNT(*) FILTER (WHERE a.paid),
            COUNT(*) FILTER (WHERE NOT a.paid),
            COUNT(*) FILTER (WHERE a.invoice_issued),
            COUNT(*) FILTER (WHERE a.receipt_issued),
            COUNT(*) FILTER (WHERE a.sales_document_generated),
            COUNT(*)
        {base}
        """,
        params,
    )
    totals = cur.fetchone()
    cur.execute(
        f"""
        SELECT COALESCE(o.salesperson_name,o.salesperson_email,a.salesperson,'Nieprzypisany'),
               COUNT(*),COALESCE(SUM(value.order_value),0),
               COALESCE(SUM(a.amount_paid),0),COALESCE(SUM(a.amount_due),0)
        {base}
        GROUP BY COALESCE(o.salesperson_name,o.salesperson_email,a.salesperson,'Nieprzypisany')
        ORDER BY 1
        """,
        params,
    )
    by_salesperson = cur.fetchall()
    cur.execute(
        f"""
        SELECT COALESCE(NULLIF(a.payment_method,''),'Brak'),COUNT(*),
               COALESCE(SUM(value.order_value),0),COALESCE(SUM(a.amount_paid),0),
               COALESCE(SUM(a.amount_due),0)
        {base}
        GROUP BY COALESCE(NULLIF(a.payment_method,''),'Brak') ORDER BY 1
        """,
        params,
    )
    by_payment = cur.fetchall()
    cur.execute(
        f"""
        SELECT o.payment_status,COUNT(*),COALESCE(SUM(value.order_value),0),
               COALESCE(SUM(a.amount_paid),0),COALESCE(SUM(a.amount_due),0)
        {base}
        GROUP BY o.payment_status ORDER BY 1
        """,
        params,
    )
    by_status = cur.fetchall()
    lists = {
        "waiting": [row for row in orders if row[23] or row[4] == "Oczekuje na płatność"],
        "paid_no_invoice": [row for row in orders if row[21] and not row[15]],
        "invoiced": [row for row in orders if row[15]],
        "ready_warehouse": [row for row in orders if row[19]],
        "unsettled": [row for row in orders if not row[20]],
        "finished": [row for row in orders if row[20] or row[5] in {"Zakończone", "Dostarczone"}],
    }
    return {
        "orders": orders,
        "totals": totals,
        "by_salesperson": by_salesperson,
        "by_payment": by_payment,
        "by_status": by_status,
        "lists": lists,
    }


def accounting_excel_text(value):
    text = str(value or "")
    return "'" + text if text[:1] in {"=", "+", "-", "@"} else text


def accounting_report_xlsx(data, filters):
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Podsumowanie"
    summary.sheet_view.showGridLines = False
    summary["A1"] = "Raport księgowości Primadera"
    summary["A1"].font = Font(size=16, bold=True, color="422D00")
    summary.merge_cells("A1:D1")
    summary["A2"] = "Wygenerowano"
    summary["B2"] = datetime.now()
    summary["B2"].number_format = "yyyy-mm-dd hh:mm"
    labels = (
        "Suma wartości proform","Suma sprzedaży","Suma zapłacona",
        "Suma do zapłaty","Liczba proform","Liczba opłaconych",
        "Liczba nieopłaconych","Liczba faktur","Liczba paragonów",
        "Liczba dokumentów sprzedaży","Liczba zamówień",
    )
    totals = data["totals"]
    for row_index, (label, value) in enumerate(zip(labels, totals), start=4):
        summary.cell(row_index, 1, label)
        summary.cell(row_index, 2, value or 0)
        if row_index <= 7:
            summary.cell(row_index, 2).number_format = '#,##0.00 "zł"'
    summary["D2"] = "Aktywne filtry"
    summary["D2"].font = Font(bold=True)
    filter_row = 3
    for key, value in filters.items():
        if value:
            summary.cell(filter_row, 4, accounting_excel_text(key))
            summary.cell(filter_row, 5, accounting_excel_text(value))
            filter_row += 1
    for cell in summary["A"]:
        cell.font = Font(bold=True) if cell.row >= 4 else cell.font
    summary.column_dimensions["A"].width = 34
    summary.column_dimensions["B"].width = 22
    summary.column_dimensions["D"].width = 24
    summary.column_dimensions["E"].width = 28

    orders_sheet = workbook.create_sheet("Zamówienia")
    orders_sheet.sheet_view.showGridLines = False
    headers = (
        "Zamówienie","Data","Klient","Handlowiec","Status zamówienia",
        "Status płatności","Sposób płatności","Proforma","Faktura","Paragon",
        "Wartość","Zapłacono","Do zapłaty","Termin płatności",
        "Dokument sprzedaży","Wysłany","Gotowe do wysyłki","Rozliczone",
    )
    orders_sheet.append(headers)
    for order in data["orders"]:
        orders_sheet.append(
            [
                accounting_excel_text(order[1]),order[2],
                accounting_excel_text(order[3]),accounting_excel_text(order[10]),
                accounting_excel_text(order[5]),accounting_excel_text(order[4]),
                accounting_excel_text(order[6]),accounting_excel_text(order[7]),
                accounting_excel_text(order[8]),accounting_excel_text(order[9]),
                float(order[11] or 0),float(order[12] or 0),float(order[13] or 0),
                order[26],"Tak" if order[17] else "Nie",
                "Tak" if order[18] else "Nie","Tak" if order[19] else "Nie",
                "Tak" if order[20] else "Nie",
            ]
        )
    header_fill = PatternFill("solid", fgColor="FFC067")
    for cell in orders_sheet[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="422D00")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    orders_sheet.freeze_panes = "A2"
    orders_sheet.auto_filter.ref = orders_sheet.dimensions
    widths = (20,13,28,24,22,22,18,18,18,18,16,16,16,16,18,13,18,14)
    for index, width in enumerate(widths, start=1):
        orders_sheet.column_dimensions[get_column_letter(index)].width = width
    for row in orders_sheet.iter_rows(min_row=2):
        row[1].number_format = "yyyy-mm-dd"
        if row[13].value:
            row[13].number_format = "yyyy-mm-dd"
        for index in (10,11,12):
            row[index].number_format = '#,##0.00 "zł"'
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def accounting_report_pdf(data, filters):
    body_font, bold_font = shop_pdf_fonts()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,pagesize=landscape(A4),leftMargin=10*mm,rightMargin=10*mm,
        topMargin=10*mm,bottomMargin=10*mm,title="Raport księgowości Primadera",
        author="Primadera",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AccountingTitle",parent=styles["Title"],fontName=bold_font,fontSize=16,
        leading=20,textColor=colors.HexColor("#422D00"),alignment=TA_CENTER,
    )
    cell_style = ParagraphStyle(
        "AccountingCell",parent=styles["Normal"],fontName=body_font,
        fontSize=6.8,leading=8,
    )
    header_style = ParagraphStyle(
        "AccountingHeader",parent=cell_style,fontName=bold_font,
        textColor=colors.HexColor("#422D00"),alignment=TA_CENTER,
    )
    totals = data["totals"]
    story = [
        Paragraph("Raport księgowości Primadera", title_style),
        Spacer(1, 3*mm),
        Paragraph(
            escape(
                f"Sprzedaż: {float(totals[1] or 0):.2f} zł | "
                f"Zapłacono: {float(totals[2] or 0):.2f} zł | "
                f"Do zapłaty: {float(totals[3] or 0):.2f} zł | "
                f"Zamówienia: {int(totals[10] or 0)}"
            ),
            cell_style,
        ),
        Spacer(1, 3*mm),
    ]
    table_data = [[
        Paragraph(label, header_style)
        for label in (
            "Zamówienie","Data","Klient","Handlowiec","Płatność","Sposób",
            "Proforma","Faktura","Paragon","Wartość","Zapłacono","Do zapłaty",
        )
    ]]
    for order in data["orders"]:
        table_data.append([
            Paragraph(escape(str(order[1] or "")),cell_style),
            Paragraph(escape(str(order[2] or "")),cell_style),
            Paragraph(escape(str(order[3] or "")),cell_style),
            Paragraph(escape(str(order[10] or "")),cell_style),
            Paragraph(escape(str(order[4] or "")),cell_style),
            Paragraph(escape(str(order[6] or "")),cell_style),
            Paragraph(escape(str(order[7] or "")),cell_style),
            Paragraph(escape(str(order[8] or "")),cell_style),
            Paragraph(escape(str(order[9] or "")),cell_style),
            Paragraph(f"{float(order[11] or 0):.2f}",cell_style),
            Paragraph(f"{float(order[12] or 0):.2f}",cell_style),
            Paragraph(f"{float(order[13] or 0):.2f}",cell_style),
        ])
    table = Table(
        table_data,repeatRows=1,
        colWidths=[25*mm,18*mm,31*mm,27*mm,25*mm,21*mm,21*mm,21*mm,21*mm,20*mm,20*mm,20*mm],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#FFC067")),
        ("GRID",(0,0),(-1,-1),0.3,colors.HexColor("#D1D5DB")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("ALIGN",(9,1),(-1,-1),"RIGHT"),
        ("TOPPADDING",(0,0),(-1,-1),3),
        ("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]))
    story.append(table)
    document.build(story)
    return output.getvalue()


@app.route('/ksiegowosc/export.<fmt>')
@login_required
@accounting_required
def accounting_export(fmt):
    if fmt not in {"xlsx","pdf"}:
        return "Nieprawidłowy format eksportu.", 400
    filters = accounting_filters_from_request()
    conn = db()
    cur = conn.cursor()
    try:
        data = accounting_dashboard_data(cur, filters)
    finally:
        conn.close()
    if fmt == "xlsx":
        content = accounting_report_xlsx(data, filters)
        mimetype = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    else:
        content = accounting_report_pdf(data, filters)
        mimetype = "application/pdf"
    filename = f"ksiegowosc-{datetime.now().strftime('%Y%m%d-%H%M')}.{fmt}"
    return Response(
        content,mimetype=mimetype,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.route('/ksiegowosc/orders/<int:order_id>')
@login_required
@accounting_required
def accounting_order_detail(order_id):
    conn = db()
    cur = conn.cursor()
    try:
        ensure_shop_accounting_row(cur, order_id)
        conn.commit()
        cur.execute(
            """
            SELECT id,order_number,order_date,customer_name,delivery_address,phone,
                   email,shipping_cost,payment_method,payment_status,status,
                   tracking_number,notes,nip,salesperson_name,salesperson_email
            FROM shop_orders WHERE id=%s
            """,
            (order_id,),
        )
        order = cur.fetchone()
        if not order:
            return "Nie znaleziono zamówienia.", 404
        cur.execute(
            """
            SELECT id,product_id,product_name,qty,price_netto,price_brutto,vat,
                   warehouse,reserved_qty,issued_qty
            FROM shop_order_items WHERE order_id=%s ORDER BY id
            """,
            (order_id,),
        )
        items = cur.fetchall()
        cur.execute("SELECT * FROM shop_accounting WHERE order_id=%s", (order_id,))
        accounting = cur.fetchone()
        cur.execute(
            """
            SELECT id,document_type,document_number,status,void_reason,voided_by,
                   voided_at,restored_by,restored_at,updated_at
            FROM accounting_documents WHERE order_id=%s
            ORDER BY CASE document_type
                WHEN 'proforma' THEN 1 WHEN 'invoice' THEN 2
                WHEN 'receipt' THEN 3 ELSE 4 END
            """,
            (order_id,),
        )
        documents = cur.fetchall()
        document_warnings = {
            document[0]: accounting_document_warnings(cur, order_id, document[1])
            for document in documents
        }
        cur.execute(
            """
            SELECT entity_type,field_name,previous_value,new_value,changed_by,changed_at
            FROM accounting_history WHERE order_id=%s
            ORDER BY changed_at DESC,id DESC LIMIT 300
            """,
            (order_id,),
        )
        accounting_history = cur.fetchall()
        cur.execute(
            "SELECT id,voided_at FROM shop_sales_documents WHERE order_id=%s",
            (order_id,),
        )
        sales_document = cur.fetchone()
        return render_template(
            "accounting_order.html",
            order=order,
            items=items,
            accounting=accounting,
            documents=documents,
            document_warnings=document_warnings,
            accounting_history=accounting_history,
            sales_document=sales_document,
            payment_methods=ACCOUNTING_PAYMENT_METHODS,
            can_ship=order_can_be_shipped(accounting),
        )
    finally:
        conn.close()


@app.route('/ksiegowosc/orders/<int:order_id>/edit', methods=['POST'])
@login_required
@accounting_required
def edit_accounting_order(order_id):
    conn = db()
    cur = conn.cursor()
    try:
        customer_name = (request.form.get("customer_name") or "").strip()
        delivery_address = (request.form.get("delivery_address") or "").strip()
        if not customer_name or len(customer_name) > 200:
            raise ValueError("Podaj prawidłową nazwę klienta.")
        if not delivery_address or len(delivery_address) > 500:
            raise ValueError("Podaj prawidłowy adres klienta.")
        shipping_cost = parse_nonnegative_number(
            request.form.get("shipping_cost"), "Koszt wysyłki"
        )
        payment_method = (request.form.get("payment_method") or "").strip()
        if payment_method and payment_method not in ACCOUNTING_PAYMENT_METHODS:
            raise ValueError("Nieprawidłowy sposób płatności.")
        payment_due_date = normalized_optional_date(
            request.form.get("payment_due_date"), "Termin płatności"
        )
        cur.execute(
            """
            SELECT customer_name,delivery_address,phone,email,nip,shipping_cost,
                   payment_method,notes
            FROM shop_orders WHERE id=%s FOR UPDATE
            """,
            (order_id,),
        )
        before_order = cur.fetchone()
        if not before_order:
            return "Nie znaleziono zamówienia.", 404
        new_order_values = (
            customer_name,delivery_address,
            (request.form.get("phone") or "").strip()[:50],
            (request.form.get("email") or "").strip()[:254],
            (request.form.get("nip") or "").strip()[:30],
            shipping_cost,payment_method or None,
            (request.form.get("notes") or "").strip()[:2000],
        )
        order_fields = (
            "customer_name","delivery_address","phone","email","nip",
            "shipping_cost","payment_method","notes",
        )
        for field, previous, new in zip(order_fields, before_order, new_order_values):
            accounting_history_change(
                cur, order_id, "order", order_id, field, previous, new
            )
        cur.execute(
            """
            UPDATE shop_orders SET customer_name=%s,delivery_address=%s,phone=%s,
                email=%s,nip=%s,shipping_cost=%s,payment_method=%s,notes=%s,
                updated_at=NOW() WHERE id=%s
            """,
            (*new_order_values, order_id),
        )
        item_ids = request.form.getlist("item_id")
        names = request.form.getlist("product_name")
        quantities = request.form.getlist("qty")
        net_prices = request.form.getlist("price_netto")
        gross_prices = request.form.getlist("price_brutto")
        vat_values = request.form.getlist("vat")
        if not item_ids:
            raise ValueError("Zamówienie musi zawierać co najmniej jedną pozycję.")
        for index, item_id_raw in enumerate(item_ids):
            try:
                item_id = int(item_id_raw)
            except (TypeError, ValueError):
                raise ValueError(f"Pozycja {index + 1}: nieprawidłowy identyfikator.")
            product_name = form_value(names, index).strip()
            if not product_name or len(product_name) > 200:
                raise ValueError(f"Pozycja {index + 1}: podaj nazwę produktu.")
            qty = parse_positive_number(form_value(quantities, index))
            net = parse_nonnegative_number(
                form_value(net_prices, index), f"Pozycja {index + 1}: cena netto"
            )
            gross = parse_nonnegative_number(
                form_value(gross_prices, index), f"Pozycja {index + 1}: cena brutto"
            )
            vat = parse_nonnegative_number(
                form_value(vat_values, index), f"Pozycja {index + 1}: VAT"
            )
            if vat > 100:
                raise ValueError(f"Pozycja {index + 1}: VAT nie może przekraczać 100%.")
            cur.execute(
                """
                SELECT product_id,product_name,qty,price_netto,price_brutto,vat,
                       reserved_qty,issued_qty
                FROM shop_order_items
                WHERE id=%s AND order_id=%s FOR UPDATE
                """,
                (item_id, order_id),
            )
            before_item = cur.fetchone()
            if not before_item:
                raise ValueError(f"Pozycja {index + 1} nie istnieje.")
            if qty + 1e-9 < float(before_item[7] or 0):
                raise ValueError(
                    f"Pozycja {index + 1}: ilość nie może być mniejsza od już wydanej."
                )
            cur.execute(
                """
                SELECT p.qty-COALESCE((
                    SELECT SUM(i.reserved_qty-i.issued_qty)
                    FROM shop_order_items i
                    JOIN shop_orders o ON o.id=i.order_id
                    WHERE i.product_id=p.id AND i.order_id<>%s
                      AND o.status NOT IN ('Anulowane','Zakończone')
                ),0)-COALESCE((
                    SELECT SUM(ri.qty)
                    FROM reservation_items ri
                    JOIN reservations r ON r.id=ri.reservation_id
                    WHERE ri.product_id=p.id AND r.status = ANY(%s)
                ),0)
                FROM products p WHERE p.id=%s
                """,
                (order_id, active_reservation_statuses_sql(), before_item[0]),
            )
            available_row = cur.fetchone()
            if not available_row or float(available_row[0] or 0) + 1e-9 < qty:
                raise ValueError(
                    f"Pozycja {index + 1}: brak wystarczającego stanu magazynowego."
                )
            item_fields = ("product_name","qty","price_netto","price_brutto","vat")
            new_values = (product_name,qty,net,gross,vat)
            for field, previous, new in zip(item_fields, before_item[1:6], new_values):
                accounting_history_change(
                    cur, order_id, "item", item_id, field, previous, new
                )
            cur.execute(
                """
                UPDATE shop_order_items SET product_name=%s,qty=%s,price_netto=%s,
                    price_brutto=%s,vat=%s,reserved_qty=%s
                WHERE id=%s AND order_id=%s
                """,
                (product_name,qty,net,gross,vat,qty,item_id,order_id),
            )
        ensure_shop_accounting_row(cur, order_id)
        cur.execute(
            """
            SELECT accounting_notes,payment_due_date,amount_paid,paid
            FROM shop_accounting WHERE order_id=%s FOR UPDATE
            """,
            (order_id,),
        )
        before_accounting = cur.fetchone()
        accounting_notes = (request.form.get("accounting_notes") or "").strip()[:2000]
        accounting_history_change(
            cur,order_id,"accounting",order_id,"accounting_notes",
            before_accounting[0],accounting_notes,
        )
        accounting_history_change(
            cur,order_id,"accounting",order_id,"payment_due_date",
            before_accounting[1],payment_due_date,
        )
        cur.execute(
            """
            SELECT COALESCE(SUM(i.qty*i.price_brutto),0)+o.shipping_cost
            FROM shop_orders o LEFT JOIN shop_order_items i ON i.order_id=o.id
            WHERE o.id=%s GROUP BY o.id
            """,
            (order_id,),
        )
        order_value = float(cur.fetchone()[0] or 0)
        amount_due = 0 if before_accounting[3] else max(
            order_value-float(before_accounting[2] or 0),0
        )
        cur.execute(
            """
            UPDATE shop_accounting SET payment_method=%s,
                payment_due_date=NULLIF(%s,'')::date,accounting_notes=%s,
                amount_due=%s,updated_by=%s,updated_at=NOW() WHERE order_id=%s
            """,
            (
                payment_method or None,payment_due_date,accounting_notes,amount_due,
                session.get("user"),order_id,
            ),
        )
        cur.execute(
            "SELECT 1 FROM shop_sales_documents WHERE order_id=%s AND voided_at IS NULL",
            (order_id,),
        )
        if cur.fetchone():
            generate_shop_sales_document(cur, order_id, session.get("user"))
        shop_history(cur, order_id, "edytowano dane księgowe dokumentu")
        conn.commit()
        cache.clear()
        return redirect(f"/ksiegowosc/orders/{order_id}?edited=1")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Accounting document edit failed.")
        return "Nie udało się zapisać zmian dokumentu.", 500
    finally:
        conn.close()


@app.route('/ksiegowosc/orders/<int:order_id>', methods=['POST'])
@login_required
@accounting_required
def update_order_accounting(order_id):
    conn = db(); cur = conn.cursor()
    try:
        ensure_shop_accounting_row(cur, order_id)
        cur.execute("SELECT * FROM shop_accounting WHERE order_id=%s FOR UPDATE", (order_id,))
        before = cur.fetchone()
        if not before:
            return "Nie znaleziono zamówienia.", 404
        bool_values = {field: request.form.get(field) == "on" for field in ACCOUNTING_BOOL_FIELDS}
        payment_method = (request.form.get("payment_method") or "").strip()
        if payment_method and payment_method not in ACCOUNTING_PAYMENT_METHODS:
            raise ValueError("Nieprawidłowy sposób płatności.")
        amount_paid = parse_nonnegative_number(request.form.get("amount_paid"), "Kwota zapłacona")
        amount_due = parse_nonnegative_number(request.form.get("amount_due"), "Kwota pozostała do zapłaty")
        document_issue_date = normalized_optional_date(
            request.form.get("document_issue_date"),
            "Data wystawienia dokumentu",
        )
        payment_received_date = normalized_optional_date(
            request.form.get("payment_received_date"),
            "Data otrzymania płatności",
        )
        payment_due_date = normalized_optional_date(
            request.form.get("payment_due_date"),
            "Termin płatności",
        )
        salesperson = (request.form.get("salesperson") or "").strip()[:200]
        proforma_number = request.form.get("proforma_number", "").strip()[:100]
        invoice_number = request.form.get("invoice_number", "").strip()[:100]
        receipt_number = request.form.get("receipt_number", "").strip()[:100]
        accounting_notes = request.form.get("accounting_notes", "").strip()[:2000]
        if bool_values["proforma_issued"] and not proforma_number:
            raise ValueError("Podaj numer wystawionej proformy.")
        if bool_values["invoice_issued"] and not invoice_number:
            raise ValueError("Podaj numer wystawionej faktury.")
        if bool_values["receipt_issued"] and not receipt_number:
            raise ValueError("Podaj numer wystawionego paragonu.")
        cur.execute(
            """
            SELECT document_type FROM accounting_documents
            WHERE order_id=%s AND status='cancelled'
            """,
            (order_id,),
        )
        cancelled_types = {row[0] for row in cur.fetchall()}
        for field, document_type in (
            ("proforma_issued", "proforma"),
            ("invoice_issued", "invoice"),
            ("receipt_issued", "receipt"),
            ("sales_document_generated", "sales"),
        ):
            if bool_values[field] and document_type in cancelled_types:
                raise ValueError(
                    "Anulowany dokument musi zostać przywrócony przez administratora."
                )
        if bool_values["sales_document_generated"]:
            cur.execute(
                """
                SELECT 1 FROM shop_sales_documents
                WHERE order_id=%s AND voided_at IS NULL
                """,
                (order_id,),
            )
            if not cur.fetchone():
                raise ValueError("Najpierw wygeneruj dokument sprzedaży.")
        if bool_values["document_sent"] and not bool_values["sales_document_generated"]:
            raise ValueError("Nie można wysłać dokumentu, który nie został wygenerowany.")
        if bool_values["paid"]:
            bool_values["partial_payment"] = False
            bool_values["waiting_for_payment"] = False
            amount_due = 0
        elif payment_method == "Pobranie":
            bool_values["waiting_for_payment"] = False
        if bool_values["ready_to_ship"] and not (
            bool_values["paid"] or payment_method == "Pobranie"
        ):
            raise ValueError(
                "Wydanie jest zablokowane do czasu zapłaty lub wyboru płatności za pobraniem."
            )
        cur.execute(
            """
            UPDATE shop_accounting SET
                proforma_issued=%s, reserved_by_proforma=%s, waiting_for_payment=%s,
                partial_payment=%s, paid=%s, payment_method=%s, invoice_issued=%s,
                receipt_issued=%s, invoice_sent=%s, document_to_warehouse=%s,
                sales_document_generated=%s,document_sent=%s,
                ready_to_ship=%s, settled=%s, proforma_number=%s, invoice_number=%s,
                receipt_number=%s, document_issue_date=NULLIF(%s,'')::date,
                payment_received_date=NULLIF(%s,'')::date,
                payment_due_date=NULLIF(%s,'')::date,amount_paid=%s,
                amount_due=%s, accounting_notes=%s, salesperson=%s,
                updated_by=%s, updated_at=NOW()
            WHERE order_id=%s
            """,
            (
                bool_values["proforma_issued"], bool_values["reserved_by_proforma"],
                bool_values["waiting_for_payment"], bool_values["partial_payment"],
                bool_values["paid"], payment_method or None, bool_values["invoice_issued"],
                bool_values["receipt_issued"], bool_values["invoice_sent"],
                bool_values["document_to_warehouse"],
                bool_values["sales_document_generated"],bool_values["document_sent"],
                bool_values["ready_to_ship"],bool_values["settled"],
                proforma_number,invoice_number,receipt_number,
                document_issue_date,payment_received_date,payment_due_date,
                amount_paid,amount_due,accounting_notes,
                salesperson, session.get("user"), order_id,
            ),
        )
        sync_accounting_payment_status(cur, order_id)
        for accounting_field, stage_key in {
            "paid": "paid",
            "proforma_issued": "proforma_issued",
            "invoice_issued": "invoice_issued",
            "receipt_issued": "receipt_issued",
            "sales_document_generated": "sales_document_generated",
            "document_sent": "document_sent",
        }.items():
            update_shop_stage(
                cur,
                order_id,
                stage_key,
                bool_values[accounting_field],
                session.get("user", "system"),
            )
        if invoice_number:
            cur.execute(
                "UPDATE shop_orders SET sales_document_number=%s WHERE id=%s",
                (invoice_number, order_id),
            )
        document_values = (
            ("proforma", proforma_number, bool_values["proforma_issued"]),
            ("invoice", invoice_number, bool_values["invoice_issued"]),
            ("receipt", receipt_number, bool_values["receipt_issued"]),
        )
        for document_type, number, issued in document_values:
            if number or issued:
                sync_accounting_document(
                    cur,
                    order_id,
                    document_type,
                    number,
                    issued,
                    session.get("user"),
                    {
                        "document_issue_date": document_issue_date,
                        "payment_due_date": payment_due_date,
                    },
                )
        labels = []
        before_bool = dict(zip(
            ACCOUNTING_BOOL_FIELDS,
            [
                before[1],before[2],before[3],before[4],before[5],before[7],
                before[8],before[9],before[10],before[24],before[25],before[11],
                before[12],
            ],
        ))
        for field in ACCOUNTING_BOOL_FIELDS:
            if bool(before_bool[field]) != bool_values[field]:
                labels.append(f"{ACCOUNTING_FIELD_LABELS[field]}: {'tak' if bool_values[field] else 'nie'}")
                accounting_history_change(
                    cur,
                    order_id,
                    "accounting",
                    order_id,
                    field,
                    bool(before_bool[field]),
                    bool_values[field],
                )
        scalar_changes = {
            "payment_method": (before[6], payment_method),
            "proforma_number": (before[13], proforma_number),
            "invoice_number": (before[14], invoice_number),
            "receipt_number": (before[15], receipt_number),
            "document_issue_date": (before[16], document_issue_date),
            "payment_received_date": (before[17], payment_received_date),
            "amount_paid": (before[18], amount_paid),
            "amount_due": (before[19], amount_due),
            "accounting_notes": (before[20], accounting_notes),
            "payment_due_date": (before[26], payment_due_date),
        }
        for field, (previous, new) in scalar_changes.items():
            if str(previous or "") != str(new or ""):
                labels.append(f"{field}: {previous or 'brak'} → {new or 'brak'}")
                accounting_history_change(
                    cur, order_id, "accounting", order_id, field, previous, new
                )
        if (before[21] or "") != salesperson:
            labels.append(f"Handlowiec: {before[21] or 'brak'} → {salesperson or 'brak'}")
            accounting_history_change(
                cur,
                order_id,
                "accounting",
                order_id,
                "salesperson",
                before[21],
                salesperson,
            )
        if labels:
            shop_history(cur, order_id, "zmieniono księgowość", "; ".join(labels))
        conn.commit()
        cache.clear()
        return redirect(f"/ksiegowosc/orders/{order_id}?saved=1")
    except ValueError as exc:
        conn.rollback(); return str(exc), 400
    except Exception:
        conn.rollback(); logger.exception("Accounting update failed."); return "Nie udało się zapisać księgowości.", 500
    finally:
        conn.close()


@app.route('/ksiegowosc/documents/<int:document_id>/delete', methods=['POST'])
@login_required
@accounting_required
def soft_delete_accounting_document(document_id):
    reason = (request.form.get("reason") or "").strip()
    if len(reason) < 3:
        return "Podaj powód usunięcia dokumentu.", 400
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT order_id,document_type,document_number,status
            FROM accounting_documents WHERE id=%s FOR UPDATE
            """,
            (document_id,),
        )
        document = cur.fetchone()
        if not document:
            return "Nie znaleziono dokumentu.", 404
        order_id, document_type, _number, status = document
        if status == "cancelled":
            return redirect(f"/ksiegowosc/orders/{order_id}")
        warnings = accounting_document_warnings(cur, order_id, document_type)
        if warnings and request.form.get("confirm_links") != "on":
            return (
                "Dokument jest powiązany z: " + ", ".join(warnings)
                + ". Potwierdź usunięcie.",
                409,
            )
        cur.execute(
            """
            UPDATE accounting_documents SET status='cancelled',void_reason=%s,
                voided_by=%s,voided_at=NOW(),updated_by=%s,updated_at=NOW()
            WHERE id=%s
            """,
            (reason,session.get("user"),session.get("user"),document_id),
        )
        column_by_type = {
            "proforma": "proforma_issued",
            "invoice": "invoice_issued",
            "receipt": "receipt_issued",
            "sales": "sales_document_generated",
        }
        stage_by_type = {
            "proforma": "proforma_issued",
            "invoice": "invoice_issued",
            "receipt": "receipt_issued",
            "sales": "sales_document_generated",
        }
        column = column_by_type[document_type]
        cur.execute(
            f"""
            UPDATE shop_accounting SET {column}=FALSE,updated_by=%s,updated_at=NOW()
            WHERE order_id=%s
            """,
            (session.get("user"),order_id),
        )
        if document_type == "sales":
            cur.execute(
                """
                UPDATE shop_sales_documents SET voided_at=NOW(),voided_by=%s,
                    void_reason=%s WHERE order_id=%s AND voided_at IS NULL
                """,
                (session.get("user"),reason,order_id),
            )
            cur.execute(
                """
                UPDATE shop_accounting SET document_sent=FALSE
                WHERE order_id=%s
                """,
                (order_id,),
            )
            update_shop_stage(
                cur,order_id,"document_sent",False,session.get("user","system")
            )
        update_shop_stage(
            cur,order_id,stage_by_type[document_type],False,
            session.get("user","system"),
        )
        accounting_history_change(
            cur,order_id,"document",document_id,"status",status,"cancelled"
        )
        accounting_history_change(
            cur,order_id,"document",document_id,"void_reason","",reason
        )
        shop_history(
            cur,order_id,"anulowano dokument",
            f"{document_type}: {reason}",
        )
        conn.commit()
        cache.clear()
        return redirect(f"/ksiegowosc/orders/{order_id}?deleted=1")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Accounting document soft delete failed.")
        return "Nie udało się anulować dokumentu.", 500
    finally:
        conn.close()


@app.route('/ksiegowosc/documents/<int:document_id>/restore', methods=['POST'])
@admin_required
def restore_accounting_document(document_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT order_id,document_type,status
            FROM accounting_documents WHERE id=%s FOR UPDATE
            """,
            (document_id,),
        )
        document = cur.fetchone()
        if not document:
            return "Nie znaleziono dokumentu.", 404
        order_id, document_type, status = document
        if status != "cancelled":
            return redirect(f"/ksiegowosc/orders/{order_id}")
        if document_type == "sales":
            cur.execute(
                "SELECT 1 FROM shop_sales_documents WHERE order_id=%s",
                (order_id,),
            )
            if not cur.fetchone():
                raise ValueError("Brak pliku dokumentu sprzedaży do przywrócenia.")
            cur.execute(
                """
                UPDATE shop_sales_documents SET voided_at=NULL,voided_by=NULL,
                    void_reason=NULL,restored_at=NOW(),restored_by=%s
                WHERE order_id=%s
                """,
                (session.get("user"),order_id),
            )
        cur.execute(
            """
            UPDATE accounting_documents SET status='active',restored_by=%s,
                restored_at=NOW(),updated_by=%s,updated_at=NOW()
            WHERE id=%s
            """,
            (session.get("user"),session.get("user"),document_id),
        )
        column_by_type = {
            "proforma": "proforma_issued",
            "invoice": "invoice_issued",
            "receipt": "receipt_issued",
            "sales": "sales_document_generated",
        }
        stage_by_type = {
            "proforma": "proforma_issued",
            "invoice": "invoice_issued",
            "receipt": "receipt_issued",
            "sales": "sales_document_generated",
        }
        cur.execute(
            f"""
            UPDATE shop_accounting SET {column_by_type[document_type]}=TRUE,
                updated_by=%s,updated_at=NOW() WHERE order_id=%s
            """,
            (session.get("user"),order_id),
        )
        update_shop_stage(
            cur,order_id,stage_by_type[document_type],True,
            session.get("user","system"),
        )
        accounting_history_change(
            cur,order_id,"document",document_id,"status","cancelled","active"
        )
        shop_history(cur,order_id,"przywrócono dokument",document_type)
        conn.commit()
        cache.clear()
        return redirect(f"/ksiegowosc/orders/{order_id}?restored=1")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Accounting document restore failed.")
        return "Nie udało się przywrócić dokumentu.", 500
    finally:
        conn.close()


@app.route('/sklep')
@login_required
def shop_orders():
    denied = require_shop_permission("view")
    if denied:
        return denied
    return render_shop_dashboard()


def shop_dashboard_data(cur, filters=None):
    filters = filters or {}
    q = (filters.get("q") or "").strip()
    params = []
    where = []
    if current_user_role() == "sales":
        where.append("lower(COALESCE(o.salesperson_email,''))=lower(%s)")
        params.append(session.get("user", ""))
    if q:
        like = f"%{q.lower()}%"
        where.append("(lower(o.order_number) LIKE %s OR lower(o.customer_name) LIKE %s OR lower(COALESCE(o.tracking_number,'')) LIKE %s OR lower(COALESCE(o.sales_document_number,'')) LIKE %s OR EXISTS (SELECT 1 FROM shop_order_items i WHERE i.order_id=o.id AND lower(i.product_name) LIKE %s))")
        params += [like]*5
    if filters.get("status"):
        where.append("o.status=%s")
        params.append(filters["status"])
    if filters.get("client"):
        where.append("lower(o.customer_name) LIKE %s")
        params.append(f"%{filters['client'].lower()}%")
    if filters.get("salesperson"):
        where.append(
            "(lower(COALESCE(o.salesperson_name,'')) LIKE %s "
            "OR lower(COALESCE(o.salesperson_email,'')) LIKE %s "
            "OR lower(COALESCE(a.salesperson,'')) LIKE %s)"
        )
        salesperson_like = f"%{filters['salesperson'].lower()}%"
        params.extend([salesperson_like] * 3)
    if filters.get("date"):
        where.append("o.order_date=%s")
        params.append(filters["date"])
    boolean_filters = {
        "invoice": "COALESCE(a.invoice_issued,FALSE)",
        "receipt": "COALESCE(a.receipt_issued,FALSE)",
        "paid": "COALESCE(a.paid,FALSE)",
        "shipped": "COALESCE((o.stages->>'shipped')::boolean,FALSE)",
        "packed": "COALESCE((o.stages->>'packed')::boolean,FALSE)",
    }
    for key, expression in boolean_filters.items():
        value = filters.get(key)
        if value in {"yes", "no"}:
            where.append(f"{expression}=%s")
            params.append(value == "yes")
    sql_where = (" WHERE " + " AND ".join(where)) if where else ""
    cur.execute(
        f"""
        SELECT o.id,o.order_number,o.order_date,o.customer_name,o.status,
               o.payment_status,o.sales_document_number,o.tracking_number,
               COALESCE(SUM(i.qty*i.price_brutto),0)+o.shipping_cost AS total,
               sd.id,sd.document_number,o.stages,
               COALESCE(a.invoice_issued,FALSE),
               COALESCE(a.receipt_issued,FALSE),
               COALESCE(a.paid,FALSE),
               COALESCE(o.salesperson_name,o.salesperson_email,a.salesperson)
        FROM shop_orders o
        LEFT JOIN shop_order_items i ON i.order_id=o.id
        LEFT JOIN shop_sales_documents sd ON sd.order_id=o.id
        LEFT JOIN shop_accounting a ON a.order_id=o.id
        {sql_where}
        GROUP BY o.id,sd.id,a.order_id
        ORDER BY o.order_date DESC,o.id DESC
        """,
        tuple(params),
    )
    orders = cur.fetchall()
    if current_user_role() == "sales":
        cur.execute(
            """
            SELECT type,message,created_at FROM shop_notifications
            WHERE resolved=FALSE AND lower(COALESCE(recipient_email,''))=lower(%s)
            ORDER BY created_at DESC LIMIT 8
            """,
            (session.get("user", ""),),
        )
    else:
        cur.execute(
            """
            SELECT type,message,created_at FROM shop_notifications
            WHERE resolved=FALSE ORDER BY created_at DESC LIMIT 8
            """
        )
    notifications = cur.fetchall()
    cur.execute(
        f"""
        SELECT
            COUNT(*) FILTER (
                WHERE NOT COALESCE(a.paid,FALSE)
                  AND o.status NOT IN ('Zakończone','Anulowane')
            ),
            COUNT(*) FILTER (WHERE COALESCE(a.invoice_issued,FALSE)),
            COUNT(*) FILTER (WHERE COALESCE(a.receipt_issued,FALSE)),
            COUNT(*) FILTER (
                WHERE COALESCE((o.stages->>'packed')::boolean,FALSE)
            ),
            COUNT(*) FILTER (
                WHERE COALESCE((o.stages->>'shipped')::boolean,FALSE)
            ),
            COUNT(*) FILTER (
                WHERE COALESCE((o.stages->>'completed')::boolean,FALSE)
                   OR o.status='Zakończone'
            )
        FROM shop_orders o
        LEFT JOIN shop_accounting a ON a.order_id=o.id
        {sql_where}
        """,
        tuple(params),
    )
    reports = cur.fetchone()
    cur.execute("SELECT id,name,qty,unit,warehouse,price_netto,vat FROM products ORDER BY warehouse,lower(name),id")
    products = cur.fetchall()
    return orders, notifications, reports, products


def render_shop_dashboard(form_error=None, status_code=200, form_data=None):
    filters = {
        key: (request.args.get(key) or "").strip()
        for key in (
            "q",
            "status",
            "client",
            "salesperson",
            "date",
            "invoice",
            "receipt",
            "paid",
            "shipped",
            "packed",
        )
    }
    conn = db()
    cur = conn.cursor()
    try:
        orders, notifications, reports, products = shop_dashboard_data(cur, filters)
        cur.execute(
            """
            SELECT email,trim(concat_ws(' ',first_name,last_name))
            FROM users WHERE role='sales' AND status='active'
            ORDER BY lower(last_name),lower(first_name),lower(email)
            """
        )
        salespeople = cur.fetchall()
    finally:
        conn.close()
    return (
        render_template(
            "shop.html",
            orders=orders,
            notifications=notifications,
            reports=reports,
            products=products,
            statuses=SHOP_STATUS_FLOW,
            role_labels=SHOP_ROLE_LABELS,
            payment_methods=ACCOUNTING_PAYMENT_METHODS,
            filters=filters,
            form_error=form_error,
            form_data=form_data,
            salespeople=salespeople,
        ),
        status_code,
    )


def reservation_filters_from_request():
    return {
        key: (request.args.get(key) or "").strip()
        for key in ("client", "salesperson", "warehouse", "status", "date", "package", "product")
    }


@app.route('/rezerwacje')
@login_required
def reservations_page():
    filters = reservation_filters_from_request()
    where = []
    params = []
    if filters["client"]:
        where.append("lower(r.customer_name) LIKE %s")
        params.append(f"%{filters['client'].lower()}%")
    if filters["salesperson"]:
        where.append("(lower(COALESCE(r.salesperson_email,'')) LIKE %s OR lower(COALESCE(r.salesperson_name,'')) LIKE %s)")
        params.extend([f"%{filters['salesperson'].lower()}%"] * 2)
    if filters["warehouse"]:
        where.append("r.warehouse=%s")
        params.append(filters["warehouse"])
    if filters["status"]:
        where.append("r.status=%s")
        params.append(filters["status"])
    if filters["date"]:
        where.append("r.reservation_date=%s")
        params.append(filters["date"])
    if filters["package"]:
        where.append("EXISTS (SELECT 1 FROM reservation_items ri WHERE ri.reservation_id=r.id AND lower(COALESCE(ri.package_number,'')) LIKE %s)")
        params.append(f"%{filters['package'].lower()}%")
    if filters["product"]:
        where.append("EXISTS (SELECT 1 FROM reservation_items ri WHERE ri.reservation_id=r.id AND lower(ri.product_name) LIKE %s)")
        params.append(f"%{filters['product'].lower()}%")
    where_sql = "WHERE " + " AND ".join(where) if where else ""
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            f"""
            SELECT r.id,r.reservation_number,r.reservation_date,r.customer_name,
                   COALESCE(r.salesperson_name,r.salesperson_email,''),r.warehouse,
                   COALESCE(SUM(ri.qty*ri.price_brutto),0) AS value,
                   COALESCE(SUM(ri.qty),0) AS qty,r.status,r.pdf_generated_at,r.notes
            FROM reservations r
            LEFT JOIN reservation_items ri ON ri.reservation_id=r.id
            {where_sql}
            GROUP BY r.id
            ORDER BY r.created_at DESC,r.id DESC
            LIMIT 300
            """,
            tuple(params),
        )
        reservations = cur.fetchall()
        cur.execute(
            "SELECT id,name,qty,unit,warehouse,price_netto,vat FROM products ORDER BY warehouse,lower(name),id"
        )
        products = cur.fetchall()
        product_ids = [p[0] for p in products]
        packages_by_product = {}
        if product_ids:
            cur.execute(
                """
                SELECT id,product_id,number,qty,warehouse
                FROM packages
                WHERE product_id=ANY(%s) AND status='active' AND qty>0
                ORDER BY warehouse,lower(number),id
                """,
                (product_ids,),
            )
            for package in cur.fetchall():
                packages_by_product.setdefault(package[1], []).append(package)
        cur.execute(
            """
            SELECT email,trim(concat_ws(' ',first_name,last_name))
            FROM users WHERE role='sales' AND status='active'
            ORDER BY lower(last_name),lower(first_name),lower(email)
            """
        )
        salespeople = cur.fetchall()
    finally:
        conn.close()
    return render_template(
        "reservations.html",
        reservations=reservations,
        filters=filters,
        products=products,
        packages_by_product=packages_by_product,
        packages_json=packages_by_product,
        warehouses=WAREHOUSES,
        statuses=RESERVATION_STATUSES,
        salespeople=salespeople,
        can_edit_reservations=reservation_user_can_edit(),
    )


@app.route('/rezerwacje', methods=['POST'])
@login_required
def create_reservation():
    if not reservation_user_can_edit():
        return "Brak uprawnień do tworzenia rezerwacji.", 403
    approve = request.form.get("action") == "approve"
    product_ids = request.form.getlist("product_id")
    package_ids = request.form.getlist("package_id")
    quantities = request.form.getlist("qty")
    dimensions = request.form.getlist("dimension")
    locations = request.form.getlist("location")
    item_notes = request.form.getlist("item_note")
    try:
        customer = (request.form.get("customer_name") or "").strip()
        if not customer or len(customer) > 200:
            raise ValueError("Wybierz lub wpisz klienta.")
        warehouse = (request.form.get("warehouse") or "").strip()
        if warehouse not in WAREHOUSES:
            raise ValueError("Wybierz prawidłowy magazyn.")
        salesperson_email = (
            session.get("user", "")
            if current_user_role() == "sales"
            else (request.form.get("salesperson_email") or "").strip().lower()
        )
        if not salesperson_email:
            raise ValueError("Wybierz handlowca prowadzącego.")
        reservation_date = normalized_document_date(request.form.get("date"))
        prepared_items = []
        for index in range(max(len(product_ids), len(quantities))):
            pid_raw = form_value(product_ids, index).strip()
            qty_raw = form_value(quantities, index).strip()
            if not pid_raw and not qty_raw:
                continue
            if not pid_raw:
                raise ValueError(f"Pozycja {index + 1}: wybierz produkt.")
            try:
                product_id = int(pid_raw)
            except ValueError:
                raise ValueError(f"Pozycja {index + 1}: nieprawidłowy produkt.")
            package_raw = form_value(package_ids, index).strip()
            package_id = int(package_raw) if package_raw else None
            prepared_items.append((
                product_id,
                package_id,
                parse_positive_number(qty_raw, f"Pozycja {index + 1}: ilość"),
                form_value(dimensions, index).strip()[:100],
                form_value(locations, index).strip()[:200],
                form_value(item_notes, index).strip()[:500],
            ))
        if not prepared_items:
            raise ValueError("Dodaj co najmniej jedną pozycję rezerwacji.")
    except ValueError as exc:
        return str(exc), 400

    conn = db()
    cur = conn.cursor()
    try:
        number = f"REZ/{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
        cur.execute(
            "SELECT pg_advisory_xact_lock(hashtext(%s))",
            (f"reservation:{number.casefold()}",),
        )
        cur.execute(
            """
            SELECT trim(concat_ws(' ',first_name,last_name))
            FROM users WHERE lower(email)=lower(%s) LIMIT 1
            """,
            (salesperson_email,),
        )
        salesperson = cur.fetchone()
        salesperson_name = salesperson[0] if salesperson else salesperson_email
        status = "zatwierdzona" if approve else "robocza"
        cur.execute(
            """
            INSERT INTO reservations(
                reservation_number,reservation_date,customer_name,salesperson_email,
                salesperson_name,warehouse,status,notes,created_by,approved_by,approved_at
            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,
                CASE WHEN %s IS NULL THEN NULL ELSE NOW() END)
            RETURNING id
            """,
            (
                number, reservation_date, customer, salesperson_email, salesperson_name,
                warehouse, status, (request.form.get("notes") or "").strip()[:2000],
                session.get("user", "system"),
                session.get("user", "system") if approve else None,
                session.get("user", "system") if approve else None,
            ),
        )
        reservation_id = cur.fetchone()[0]
        reservation_history(cur, reservation_id, "utworzono rezerwację", number)
        for index, (product_id, package_id, qty, dimension, location, note) in enumerate(prepared_items):
            cur.execute(
                "SELECT id,name,qty,unit,warehouse,price_netto,vat FROM products WHERE id=%s FOR UPDATE",
                (product_id,),
            )
            product = cur.fetchone()
            if not product:
                raise ValueError(f"Pozycja {index + 1}: produkt nie istnieje.")
            if product[4] != warehouse:
                raise ValueError(f"Pozycja {index + 1}: produkt jest w innym magazynie.")
            package_number = None
            if package_id:
                cur.execute(
                    """
                    SELECT id,number,qty,warehouse,product_id
                    FROM packages WHERE id=%s FOR UPDATE
                    """,
                    (package_id,),
                )
                package = cur.fetchone()
                if not package or package[4] != product_id:
                    raise ValueError(f"Pozycja {index + 1}: paczka nie pasuje do produktu.")
                if package[3] != warehouse:
                    raise ValueError(f"Pozycja {index + 1}: paczka jest w innym magazynie.")
                package_available = package_available_qty(cur, package_id, package[2])
                if approve and package_available + 1e-9 < qty:
                    raise ValueError(
                        f"Pozycja {index + 1}: w paczce dostępne jest tylko {package_available:g}."
                    )
                package_number = package[1]
            product_available = product_available_qty(cur, product_id, product[2])
            if approve and product_available + 1e-9 < qty:
                raise ValueError(
                    f"Pozycja {index + 1}: dostępne jest tylko {product_available:g}."
                )
            net = float(product[5] or 0)
            vat = float(product[6] or 0)
            cur.execute(
                """
                INSERT INTO reservation_items(
                    reservation_id,product_id,package_id,product_name,package_number,
                    dimension,qty,unit,warehouse,location,price_netto,price_brutto,
                    warehouse_note
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (
                    reservation_id, product_id, package_id, product[1], package_number,
                    dimension, qty, product[3], warehouse, location, net,
                    net * (1 + vat / 100), note,
                ),
            )
        if approve:
            reservation_history(cur, reservation_id, "zatwierdzono", "Towar zablokowany do sprzedaży")
            generate_and_store_reservation_pdf(cur, reservation_id, session.get("user", "system"))
            notify_warehouse_reservation(cur, reservation_id, "Nowa rezerwacja do przygotowania")
        conn.commit()
        cache.clear()
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}")


@app.route('/rezerwacje/<int:reservation_id>')
@login_required
def reservation_detail(reservation_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            """
            SELECT id,reservation_number,reservation_date,customer_name,
                   salesperson_email,salesperson_name,warehouse,status,notes,
                   cancel_reason,created_by,approved_by,approved_at,pdf_generated_at,
                   completed_by,completed_at
            FROM reservations WHERE id=%s
            """,
            (reservation_id,),
        )
        reservation = cur.fetchone()
        if not reservation:
            return "Nie znaleziono rezerwacji.", 404
        cur.execute(
            """
            SELECT id,product_id,package_id,product_name,package_number,dimension,
                   qty,unit,warehouse,location,prepared,prepared_by,prepared_at,
                   warehouse_note
            FROM reservation_items WHERE reservation_id=%s ORDER BY id
            """,
            (reservation_id,),
        )
        items = cur.fetchall()
        cur.execute(
            """
            SELECT actor_email,action,details,created_at
            FROM reservation_history WHERE reservation_id=%s
            ORDER BY created_at DESC,id DESC
            """,
            (reservation_id,),
        )
        history = cur.fetchall()
    finally:
        conn.close()
    return render_template(
        "reservation_detail.html",
        reservation=reservation,
        items=items,
        history=history,
        statuses=RESERVATION_STATUSES,
        can_pick=reservation_user_can_pick(),
        can_edit_reservations=reservation_user_can_edit(),
    )


@app.route('/rezerwacje/<int:reservation_id>/pdf')
@login_required
def reservation_pdf(reservation_id):
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "SELECT reservation_number,pdf FROM reservations WHERE id=%s",
            (reservation_id,),
        )
        row = cur.fetchone()
        if not row:
            return "Nie znaleziono rezerwacji.", 404
        content = bytes(row[1] or b"")
        if not content:
            content = generate_and_store_reservation_pdf(cur, reservation_id, session.get("user", "system"))
            conn.commit()
    finally:
        conn.close()
    return Response(
        content,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{row[0]}.pdf"'},
    )


@app.route('/rezerwacje/<int:reservation_id>/pdf/generate', methods=['POST'])
@login_required
def regenerate_reservation_pdf(reservation_id):
    conn = db()
    cur = conn.cursor()
    try:
        generate_and_store_reservation_pdf(cur, reservation_id, session.get("user", "system"))
        conn.commit()
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}?pdf=generated")


@app.route('/rezerwacje/<int:reservation_id>/items/<int:item_id>', methods=['POST'])
@login_required
def update_reservation_item(reservation_id, item_id):
    if not reservation_user_can_pick():
        return "Brak uprawnień do kompletacji.", 403
    prepared = request.form.get("prepared") == "on"
    note = (request.form.get("warehouse_note") or "").strip()[:500]
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT status FROM reservations WHERE id=%s FOR UPDATE", (reservation_id,))
        row = cur.fetchone()
        if not row:
            return "Nie znaleziono rezerwacji.", 404
        if row[0] in {"anulowana", "wydana"}:
            return "Nie można zmieniać zakończonej rezerwacji.", 400
        cur.execute(
            """
            UPDATE reservation_items
            SET prepared=%s,
                prepared_by=CASE WHEN %s THEN %s ELSE NULL END,
                prepared_at=CASE WHEN %s THEN NOW() ELSE NULL END,
                warehouse_note=%s
            WHERE id=%s AND reservation_id=%s
            """,
            (prepared, prepared, session.get("user"), prepared, note, item_id, reservation_id),
        )
        cur.execute(
            """
            UPDATE reservations
            SET status=CASE WHEN status IN ('zatwierdzona','przekazana do magazynu')
                            THEN 'w trakcie kompletowania' ELSE status END,
                updated_at=NOW()
            WHERE id=%s
            """,
            (reservation_id,),
        )
        reservation_history(
            cur,
            reservation_id,
            "odhaczono pozycję" if prepared else "cofnięto odhaczenie pozycji",
            f"Pozycja {item_id}: {note}",
        )
        conn.commit()
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}")


@app.route('/rezerwacje/<int:reservation_id>/complete', methods=['POST'])
@login_required
def complete_reservation(reservation_id):
    if not reservation_user_can_pick():
        return "Brak uprawnień do kompletacji.", 403
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT status FROM reservations WHERE id=%s FOR UPDATE", (reservation_id,))
        row = cur.fetchone()
        if not row:
            return "Nie znaleziono rezerwacji.", 404
        cur.execute(
            "SELECT COUNT(*) FILTER (WHERE NOT prepared), COUNT(*) FROM reservation_items WHERE reservation_id=%s",
            (reservation_id,),
        )
        missing, total = cur.fetchone()
        if total == 0 or missing:
            return "Najpierw odhacz wszystkie pozycje.", 400
        cur.execute(
            """
            UPDATE reservations
            SET status='skompletowana',completed_by=%s,completed_at=NOW(),updated_at=NOW()
            WHERE id=%s
            """,
            (session.get("user"), reservation_id),
        )
        reservation_history(cur, reservation_id, "skompletowano", "Wszystkie pozycje przygotowane")
        conn.commit()
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}?completed=1")


@app.route('/rezerwacje/<int:reservation_id>/cancel', methods=['POST'])
@login_required
def cancel_reservation(reservation_id):
    if not reservation_user_can_edit():
        return "Brak uprawnień do anulowania.", 403
    reason = (request.form.get("reason") or "").strip()
    if len(reason) < 3:
        return "Podaj powód anulowania.", 400
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT status FROM reservations WHERE id=%s FOR UPDATE", (reservation_id,))
        row = cur.fetchone()
        if not row:
            return "Nie znaleziono rezerwacji.", 404
        if row[0] in {"wydana", "anulowana"}:
            return "Nie można anulować zakończonej rezerwacji.", 400
        cur.execute(
            """
            UPDATE reservations
            SET status='anulowana',cancel_reason=%s,cancelled_by=%s,
                cancelled_at=NOW(),updated_at=NOW()
            WHERE id=%s
            """,
            (reason, session.get("user"), reservation_id),
        )
        reservation_history(cur, reservation_id, "anulowano", reason)
        conn.commit()
        cache.clear()
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}?cancelled=1")


@app.route('/rezerwacje/<int:reservation_id>/issue', methods=['POST'])
@login_required
def issue_reservation(reservation_id):
    if not reservation_user_can_pick():
        return "Brak uprawnień do wydania.", 403
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute("SELECT status FROM reservations WHERE id=%s FOR UPDATE", (reservation_id,))
        row = cur.fetchone()
        if not row:
            return "Nie znaleziono rezerwacji.", 404
        if row[0] not in {"skompletowana", "w trakcie kompletowania", "zatwierdzona", "przekazana do magazynu"}:
            return "Rezerwacji w tym statusie nie można wydać.", 400
        cur.execute(
            """
            SELECT id,product_id,package_id,qty,warehouse
            FROM reservation_items WHERE reservation_id=%s ORDER BY id FOR UPDATE
            """,
            (reservation_id,),
        )
        items = cur.fetchall()
        for _item_id, product_id, package_id, qty, warehouse in items:
            cur.execute(
                """
                UPDATE products SET qty=qty-%s
                WHERE id=%s AND warehouse=%s AND qty>=%s
                """,
                (qty, product_id, warehouse, qty),
            )
            if cur.rowcount != 1:
                raise ValueError("Brak stanu magazynowego do wydania rezerwacji.")
            if package_id:
                cur.execute(
                    "UPDATE packages SET qty=qty-%s WHERE id=%s AND qty>=%s",
                    (qty, package_id, qty),
                )
                if cur.rowcount != 1:
                    raise ValueError("Brak stanu paczki do wydania rezerwacji.")
                cur.execute(
                    """
                    UPDATE packages
                    SET status=CASE WHEN qty <= 0 THEN 'issued' ELSE status END,
                        archived_at=CASE WHEN qty <= 0 THEN COALESCE(archived_at,NOW()) ELSE archived_at END
                    WHERE id=%s
                    """,
                    (package_id,),
                )
        cur.execute(
            "UPDATE reservations SET status='wydana',updated_at=NOW() WHERE id=%s",
            (reservation_id,),
        )
        reservation_history(cur, reservation_id, "wydano", "Towar zdjęty ze stanu")
        conn.commit()
        cache.clear()
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()
    return redirect(f"/rezerwacje/{reservation_id}?issued=1")


@app.route('/sklep/orders', methods=['POST'])
@login_required
def shop_create_order():
    denied = require_shop_permission("create")
    if denied: return denied
    product_ids = request.form.getlist("product_id")
    qtys = request.form.getlist("qty")
    try:
        prepared_items = []
        for index in range(max(len(product_ids), len(qtys))):
            pid_raw = form_value(product_ids, index).strip()
            qty_raw = form_value(qtys, index).strip()
            if not pid_raw and not qty_raw:
                continue
            if not pid_raw:
                raise ValueError(f"Pozycja {index + 1}: wybierz produkt.")
            if not qty_raw:
                raise ValueError(f"Pozycja {index + 1}: podaj ilość.")
            try:
                product_id = int(pid_raw)
            except ValueError:
                raise ValueError(f"Pozycja {index + 1}: nieprawidłowy produkt.")
            prepared_items.append((product_id, parse_positive_number(qty_raw)))
        if not prepared_items:
            raise ValueError("Dodaj co najmniej jeden produkt.")

        order_number = (
            request.form.get("order_number")
            or f"SK/{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
        ).strip()
        customer_name = (request.form.get("customer_name") or "").strip()
        delivery_address = (request.form.get("delivery_address") or "").strip()
        email = (request.form.get("email") or "").strip()
        phone = (request.form.get("phone") or "").strip()
        nip = (request.form.get("nip") or "").strip()
        payment_method = (request.form.get("payment_method") or "").strip()
        payment_status = (
            request.form.get("payment_status") or "Oczekuje na płatność"
        ).strip()
        if not order_number or len(order_number) > 100:
            raise ValueError("Numer zamówienia jest wymagany i może mieć maksymalnie 100 znaków.")
        if not customer_name or len(customer_name) > 200:
            raise ValueError("Nazwa klienta jest wymagana i może mieć maksymalnie 200 znaków.")
        if not delivery_address or len(delivery_address) > 500:
            raise ValueError("Adres dostawy jest wymagany i może mieć maksymalnie 500 znaków.")
        if email and ("@" not in email or len(email) > 254):
            raise ValueError("Podaj prawidłowy adres e-mail klienta.")
        if len(phone) > 50 or len(nip) > 30:
            raise ValueError("Telefon lub NIP jest zbyt długi.")
        if payment_method and payment_method not in ACCOUNTING_PAYMENT_METHODS:
            raise ValueError("Wybierz prawidłowy sposób płatności.")
        if payment_status not in {"Oczekuje na płatność", "Opłacone"}:
            raise ValueError("Wybierz prawidłowy status płatności.")
        salesperson_email = (
            session.get("user", "")
            if current_user_role() == "sales"
            else (request.form.get("salesperson_email") or "").strip().lower()
        )
        if not salesperson_email:
            raise ValueError("Wybierz handlowca prowadzącego zamówienie.")
        order_date = normalized_document_date(request.form.get("date"))
        shipping = parse_nonnegative_number(
            request.form.get("shipping_cost"), "Koszt wysyłki"
        )
    except ValueError as exc:
        return render_shop_dashboard(str(exc), 400, request.form)

    conn = db(); cur = conn.cursor()
    try:
        cur.execute(
            "SELECT pg_advisory_xact_lock(hashtext(%s))",
            (f"shop-order:{order_number.casefold()}",),
        )
        cur.execute(
            "SELECT 1 FROM shop_orders WHERE lower(order_number)=lower(%s)",
            (order_number,),
        )
        if cur.fetchone():
            raise ValueError("Zamówienie o takim numerze już istnieje.")
        cur.execute(
            """
            INSERT INTO shop_orders(
                order_number,order_date,customer_name,delivery_address,phone,email,
                shipping_cost,payment_method,payment_status,status,sales_document_number,
                tracking_number,notes,nip,created_by
            ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'Nowe zamówienie',%s,%s,%s,%s,%s)
            RETURNING id
            """,
            (
                order_number, order_date, customer_name, delivery_address, phone,
                email, shipping, payment_method, payment_status,
                request.form.get("sales_document_number", "").strip(),
                request.form.get("tracking_number", "").strip(),
                request.form.get("notes", "").strip(), nip, session.get("user"),
            ),
        )
        order_id = cur.fetchone()[0]
        shop_history(cur, order_id, "utworzono zamówienie", order_number)
        if salesperson_email:
            assign_order_salesperson(
                cur,
                order_id,
                salesperson_email,
                session.get("user", "system"),
            )
        lacking = []
        for pid, qty in prepared_items:
            cur.execute("SELECT id,name,qty,warehouse,price_netto,vat FROM products WHERE id=%s FOR UPDATE", (pid,))
            p = cur.fetchone()
            if not p: raise ValueError("Wybrany produkt nie istnieje.")
            price_netto = parse_nonnegative_number(p[4], f"Cena produktu {p[1]}")
            vat = parse_nonnegative_number(p[5], f"VAT produktu {p[1]}")
            if vat > 100:
                raise ValueError(f"Produkt {p[1]} ma nieprawidłową stawkę VAT.")
            available = product_available_qty(cur, pid, p[2])
            if available + 1e-9 < qty: lacking.append(f"{p[1]} ({available})")
            reserved = qty if available + 1e-9 >= qty else 0
            cur.execute(
                """
                INSERT INTO shop_order_items(
                    order_id,product_id,product_name,qty,price_netto,price_brutto,
                    vat,warehouse,reserved_qty
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (
                    order_id, pid, p[1], qty, price_netto,
                    price_netto * (1 + vat / 100), vat, p[3], reserved,
                ),
            )
            if reserved:
                shop_history(cur, order_id, "zarezerwowano towar", f"{p[1]} x {qty}")
        update_shop_stage(
            cur,
            order_id,
            "order_accepted",
            True,
            session.get("user", "system"),
        )
        if lacking:
            message = "Brak towaru: " + ", ".join(lacking)
            cur.execute("INSERT INTO shop_notifications(order_id,type,message) VALUES (%s,'brak towaru',%s)", (order_id, message))
            notify_order_salesperson(cur, order_id, "brak towaru", message)
        else:
            cur.execute("UPDATE shop_orders SET status='Towar zarezerwowany' WHERE id=%s", (order_id,))
            update_shop_stage(
                cur,
                order_id,
                "stock_reserved",
                True,
                session.get("user", "system"),
            )
            cur.execute("INSERT INTO shop_notifications(order_id,type,message) VALUES (%s,'oczekuje na dokument','Zamówienie oczekuje na dokument sprzedaży')", (order_id,))
        ensure_shop_accounting_row(cur, order_id)
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}?created=1")
    except ValueError as exc:
        conn.rollback()
        return render_shop_dashboard(str(exc), 400, request.form)
    except psycopg2.errors.UndefinedColumn:
        conn.rollback()
        logger.exception("Shop database schema is incomplete.")
        return render_shop_dashboard(
            "Baza sklepu nie ma wszystkich wymaganych kolumn. "
            "Uruchomiono migrację przy wdrożeniu; odśwież stronę i spróbuj ponownie.",
            503,
            request.form,
        )
    except psycopg2.IntegrityError:
        conn.rollback()
        logger.exception("Shop order integrity validation failed.")
        return render_shop_dashboard(
            "Nie można zapisać zamówienia: numer zamówienia lub dane produktu "
            "kolidują z istniejącymi danymi.",
            409,
            request.form,
        )
    except Exception:
        conn.rollback()
        logger.exception("Shop order creation failed")
        return render_shop_dashboard(
            "Nie udało się zapisać zamówienia z powodu błędu bazy danych. "
            "Żadne dane nie zostały zapisane.",
            500,
            request.form,
        )
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>')
@login_required
def shop_order_detail(order_id):
    denied = require_shop_permission("view")
    if denied: return denied
    conn=db(); cur=conn.cursor()
    cur.execute(
        """
        SELECT id,order_number,order_date,customer_name,delivery_address,phone,email,
               shipping_cost,payment_method,payment_status,status,sales_document_number,
               tracking_number,notes,nip,document_confirmed,created_by,created_at,updated_at,
               stages,salesperson_email,salesperson_name,salesperson_assigned_by,
               salesperson_assigned_at
        FROM shop_orders WHERE id=%s
        """,
        (order_id,),
    )
    order=cur.fetchone()
    if not order:
        conn.close()
        return "Nie znaleziono zamówienia.", 404
    denied = sales_order_access_error(cur, order_id)
    if denied:
        conn.close()
        return denied
    cur.execute(
        """
        SELECT id,order_id,product_id,product_name,qty,price_netto,price_brutto,vat,
               warehouse,reserved_qty,issued_qty
        FROM shop_order_items WHERE order_id=%s ORDER BY id
        """,
        (order_id,),
    )
    items=cur.fetchall()
    cur.execute(
        """
        SELECT id,document_number,editable_data,confirmed
        FROM shop_sales_documents WHERE order_id=%s AND voided_at IS NULL
        """,
        (order_id,),
    )
    document=cur.fetchone()
    ensure_shop_accounting_row(cur, order_id)
    conn.commit()
    cur.execute("SELECT * FROM shop_accounting WHERE order_id=%s", (order_id,)); accounting=cur.fetchone()
    cur.execute("SELECT user_email,action,details,created_at FROM shop_order_history WHERE order_id=%s ORDER BY created_at DESC", (order_id,)); history=cur.fetchall()
    cur.execute(
        """
        SELECT stage_key,previous_value,new_value,changed_by,changed_at,
               previous_status,new_status
        FROM shop_order_stage_history
        WHERE order_id=%s ORDER BY changed_at DESC,id DESC
        """,
        (order_id,),
    )
    stage_history = cur.fetchall()
    cur.execute(
        """
        SELECT previous_salesperson,new_salesperson,changed_by,changed_at
        FROM shop_order_salesperson_history
        WHERE order_id=%s ORDER BY changed_at DESC,id DESC
        """,
        (order_id,),
    )
    salesperson_history = cur.fetchall()
    cur.execute(
        """
        SELECT email,trim(concat_ws(' ',first_name,last_name))
        FROM users WHERE role='sales' AND status='active'
        ORDER BY lower(last_name),lower(first_name),lower(email)
        """
    )
    salespeople = cur.fetchall()
    cur.execute(
        """
        SELECT id,name,qty,unit,warehouse
        FROM products ORDER BY warehouse,lower(name),id
        """
    )
    products = cur.fetchall()
    conn.close()
    stages = shop_stage_dict(order[19])
    stage_rows = [
        {
            "key": key,
            "label": label,
            "checked": bool(stages.get(key)),
            "can_edit": shop_stage_can_edit(key),
        }
        for key, label, _permission in SHOP_ORDER_STAGES
    ]
    return render_template(
        "shop_order.html",
        order=order,
        items=items,
        document=document,
        history=history,
        stage_history=stage_history,
        stage_labels={key: value["label"] for key, value in SHOP_STAGE_BY_KEY.items()},
        stage_rows=stage_rows,
        statuses=SHOP_STATUS_FLOW,
        accounting=accounting,
        payment_methods=ACCOUNTING_PAYMENT_METHODS,
        can_ship=order_can_be_shipped(accounting),
        salesperson_history=salesperson_history,
        salespeople=salespeople,
        products=products,
        can_edit_order=(
            not order[15]
            and (
                current_user_role() in {"admin", "shop"}
                or (
                    current_user_role() == "sales"
                    and str(order[20] or "").casefold()
                    == str(session.get("user") or "").casefold()
                )
            )
        ),
    )


@app.route('/sklep/orders/<int:order_id>/edit', methods=['POST'])
@login_required
def shop_edit_order(order_id):
    if current_user_role() not in {"admin", "shop", "sales"}:
        return "Brak uprawnień do edycji zamówienia.", 403
    product_ids = request.form.getlist("product_id")
    qtys = request.form.getlist("qty")
    try:
        items = []
        for index in range(max(len(product_ids), len(qtys))):
            pid_raw = form_value(product_ids, index).strip()
            qty_raw = form_value(qtys, index).strip()
            if not pid_raw and not qty_raw:
                continue
            if not pid_raw or not qty_raw:
                raise ValueError(f"Pozycja {index + 1}: wybierz produkt i podaj ilość.")
            items.append((int(pid_raw), parse_positive_number(qty_raw)))
        if not items:
            raise ValueError("Dodaj co najmniej jeden produkt.")
        customer_name = (request.form.get("customer_name") or "").strip()
        delivery_address = (request.form.get("delivery_address") or "").strip()
        if not customer_name or len(customer_name) > 200:
            raise ValueError("Podaj prawidłową nazwę klienta.")
        if not delivery_address or len(delivery_address) > 500:
            raise ValueError("Podaj prawidłowy adres dostawy.")
    except (TypeError, ValueError) as exc:
        return str(exc), 400

    conn = db()
    cur = conn.cursor()
    try:
        denied = sales_order_access_error(cur, order_id)
        if denied:
            return denied
        cur.execute(
            "SELECT document_confirmed FROM shop_orders WHERE id=%s FOR UPDATE",
            (order_id,),
        )
        order = cur.fetchone()
        if not order:
            return "Nie znaleziono zamówienia.", 404
        if order[0]:
            return "Zatwierdzonego zamówienia nie można już edytować.", 409
        cur.execute("DELETE FROM shop_order_items WHERE order_id=%s", (order_id,))
        lacking = []
        for product_id, qty in items:
            cur.execute(
                """
                SELECT id,name,qty,warehouse,price_netto,vat
                FROM products WHERE id=%s
                """,
                (product_id,),
            )
            product = cur.fetchone()
            if not product:
                raise ValueError("Wybrany produkt nie istnieje.")
            cur.execute(
                """
                SELECT COALESCE(SUM(qty),0) FROM (
                    SELECT COALESCE(SUM(i.reserved_qty-i.issued_qty),0) AS qty
                    FROM shop_order_items i
                    JOIN shop_orders o ON o.id=i.order_id
                    WHERE i.product_id=%s AND i.order_id<>%s
                      AND o.status NOT IN ('Anulowane','Zakończone')
                    UNION ALL
                    SELECT COALESCE(SUM(ri.qty),0) AS qty
                    FROM reservation_items ri
                    JOIN reservations r ON r.id=ri.reservation_id
                    WHERE ri.product_id=%s AND r.status = ANY(%s)
                ) reserved
                """,
                (product_id, order_id, product_id, active_reservation_statuses_sql()),
            )
            available = float(product[2] or 0) - float(cur.fetchone()[0] or 0)
            reserved = qty if available + 1e-9 >= qty else 0
            if not reserved:
                lacking.append(product[1])
            net = float(product[4] or 0)
            vat = float(product[5] or 0)
            cur.execute(
                """
                INSERT INTO shop_order_items(
                    order_id,product_id,product_name,qty,price_netto,price_brutto,
                    vat,warehouse,reserved_qty
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                (
                    order_id, product_id, product[1], qty, net,
                    net * (1 + vat / 100), vat, product[3], reserved,
                ),
            )
        cur.execute(
            """
            UPDATE shop_orders
            SET customer_name=%s,delivery_address=%s,phone=%s,email=%s,nip=%s,
                notes=%s,updated_at=NOW()
            WHERE id=%s
            """,
            (
                customer_name, delivery_address,
                (request.form.get("phone") or "").strip()[:50],
                (request.form.get("email") or "").strip()[:254],
                (request.form.get("nip") or "").strip()[:30],
                (request.form.get("notes") or "").strip()[:2000],
                order_id,
            ),
        )
        shop_history(cur, order_id, "edytowano zamówienie", "Zmieniono dane lub pozycje.")
        if lacking:
            notify_order_salesperson(
                cur,
                order_id,
                "brak towaru",
                "Brak wystarczającego stanu: " + ", ".join(lacking),
            )
        ensure_shop_accounting_row(cur, order_id)
        cur.execute(
            """
            UPDATE shop_accounting a
            SET amount_due=(
                SELECT COALESCE(SUM(i.qty*i.price_brutto),0)+o.shipping_cost
                FROM shop_orders o
                LEFT JOIN shop_order_items i ON i.order_id=o.id
                WHERE o.id=%s GROUP BY o.id
            ),updated_at=NOW()
            WHERE a.order_id=%s
            """,
            (order_id, order_id),
        )
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}?edited=1")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Shop order edit failed.")
        return "Nie udało się zapisać zmian zamówienia.", 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/salesperson', methods=['POST'])
@login_required
def shop_assign_salesperson(order_id):
    if current_user_role() not in {"admin", "shop"}:
        return "Brak uprawnień do przypisywania handlowca.", 403
    conn = db()
    cur = conn.cursor()
    try:
        assign_order_salesperson(
            cur,
            order_id,
            request.form.get("salesperson_email"),
            session.get("user", "system"),
        )
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}?salesperson=saved")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Salesperson assignment failed.")
        return "Nie udało się przypisać handlowca.", 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/stage', methods=['POST'])
@login_required
def shop_update_stage(order_id):
    payload = request.get_json(silent=True) or request.form
    stage_key = str(payload.get("stage") or "").strip()
    if stage_key not in SHOP_STAGE_BY_KEY:
        return jsonify({"ok": False, "message": "Nieprawidłowy etap zamówienia."}), 400
    if not shop_stage_can_edit(stage_key):
        return jsonify({"ok": False, "message": "Brak uprawnień do tego etapu."}), 403
    raw_value = payload.get("checked")
    new_value = raw_value is True or str(raw_value).lower() in {"1", "true", "on", "yes"}
    conn = db()
    cur = conn.cursor()
    try:
        denied = sales_order_access_error(cur, order_id)
        if denied:
            return jsonify({"ok": False, "message": denied[0]}), denied[1]
        previous, current = update_shop_stage(
            cur,
            order_id,
            stage_key,
            new_value,
            session.get("user", "system"),
            str(payload.get("tracking_number") or ""),
        )
        conn.commit()
        cache.clear()
        return jsonify(
            {
                "ok": True,
                "stage": stage_key,
                "previous": previous,
                "checked": current,
                "message": "Etap zapisany.",
            }
        )
    except ValueError as exc:
        conn.rollback()
        return jsonify({"ok": False, "message": str(exc)}), 400
    except Exception:
        conn.rollback()
        logger.exception("Shop stage update failed.")
        return jsonify(
            {
                "ok": False,
                "message": "Nie udało się zapisać etapu. Spróbuj ponownie.",
            }
        ), 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/details', methods=['POST'])
@login_required
def shop_update_details(order_id):
    role = current_user_role()
    if role not in {"admin", "shop", "warehouse", "sales"}:
        return "Brak uprawnień do edycji danych operacyjnych.", 403
    conn = db()
    cur = conn.cursor()
    try:
        denied = sales_order_access_error(cur, order_id)
        if denied:
            return denied
        cur.execute(
            """
            SELECT notes,tracking_number
            FROM shop_orders WHERE id=%s FOR UPDATE
            """,
            (order_id,),
        )
        current = cur.fetchone()
        if not current:
            return "Nie znaleziono zamówienia.", 404
        notes = (
            (request.form.get("notes") or "").strip()
            if role in {"admin", "shop", "sales"}
            else (current[0] or "")
        )
        tracking = (
            (request.form.get("tracking_number") or "").strip()
            if role in {"admin", "warehouse"}
            else (current[1] or "")
        )
        if len(notes) > 2000:
            raise ValueError("Uwagi mogą mieć maksymalnie 2000 znaków.")
        if len(tracking) > 100:
            raise ValueError("Numer przesyłki może mieć maksymalnie 100 znaków.")
        cur.execute(
            """
            UPDATE shop_orders
            SET notes=%s,tracking_number=%s,updated_at=NOW()
            WHERE id=%s
            """,
            (notes, tracking, order_id),
        )
        changes = []
        if (current[0] or "") != notes:
            changes.append("zmieniono uwagi")
        if (current[1] or "") != tracking:
            changes.append(
                f"numer przesyłki: {current[1] or 'brak'} → {tracking or 'brak'}"
            )
        if changes:
            shop_history(cur, order_id, "zmieniono dane zamówienia", "; ".join(changes))
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}?details=saved")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Shop order details update failed.")
        return "Nie udało się zapisać danych zamówienia.", 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/status', methods=['POST'])
@login_required
def shop_update_status(order_id):
    status = request.form.get("status")
    if status not in SHOP_STATUS_FLOW:
        return "Nieprawidłowy status.", 400
    stage_key = next(
        (key for key, mapped_status in SHOP_STAGE_STATUS.items() if mapped_status == status),
        None,
    )
    if stage_key:
        if not shop_stage_can_edit(stage_key):
            return "Brak uprawnień do tego etapu.", 403
    else:
        action = "warehouse" if status == "Dostarczone" else "shop_edit"
        denied = require_shop_permission(action)
        if denied:
            return denied
    conn = db()
    cur = conn.cursor()
    try:
        denied = sales_order_access_error(cur, order_id)
        if denied:
            return denied
        if stage_key:
            update_shop_stage(
                cur,
                order_id,
                stage_key,
                True,
                session.get("user", "system"),
                request.form.get("tracking_number", ""),
            )
        else:
            cur.execute(
                """
                SELECT status FROM shop_orders WHERE id=%s FOR UPDATE
                """,
                (order_id,),
            )
            current = cur.fetchone()
            if not current:
                raise ValueError("Nie znaleziono zamówienia.")
            cur.execute(
                """
                UPDATE shop_orders
                SET status=%s,tracking_number=COALESCE(NULLIF(%s,''),tracking_number),
                    updated_at=NOW()
                WHERE id=%s
                """,
                (status, request.form.get("tracking_number", ""), order_id),
            )
            shop_history(
                cur,
                order_id,
                "zmieniono status",
                f"{current[0]} → {status}",
            )
            notify_order_salesperson(
                cur,
                order_id,
                "zmiana statusu",
                f"Status zamówienia zmieniono: {current[0]} → {status}.",
            )
        if status == "Dokument wystawiony":
            cur.execute(
                """
                INSERT INTO shop_notifications(order_id,type,message)
                VALUES (%s,'gotowe do pakowania','Zamówienie gotowe do pakowania')
                """,
                (order_id,),
            )
        if status == "Spakowane":
            cur.execute(
                """
                INSERT INTO shop_notifications(order_id,type,message)
                VALUES (%s,'gotowe do wysyłki','Zamówienie gotowe do wysyłki')
                """,
                (order_id,),
            )
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}")
    except ValueError as exc:
        conn.rollback(); return str(exc), 400
    except Exception:
        conn.rollback()
        logger.exception("Shop order status update failed.")
        return "Nie udało się zmienić statusu zamówienia.", 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/document', methods=['POST'])
@login_required
def shop_confirm_document(order_id):
    denied=require_shop_permission("accounting")
    if denied: return denied
    number=(request.form.get('sales_document_number') or '').strip()
    if not number: return "Podaj numer faktury lub paragonu.", 400
    conn=db(); cur=conn.cursor()
    try:
        cur.execute("UPDATE shop_orders SET sales_document_number=%s,status='Dokument wystawiony',document_confirmed=TRUE WHERE id=%s", (number,order_id))
        if cur.rowcount != 1:
            raise ValueError("Nie znaleziono zamówienia.")
        generate_shop_sales_document(cur, order_id, session.get("user"))
        cur.execute("UPDATE shop_sales_documents SET document_number=%s,confirmed=TRUE,confirmed_by=%s,confirmed_at=NOW() WHERE order_id=%s", (number,session.get('user'),order_id))
        if cur.rowcount != 1:
            raise ValueError("Nie znaleziono dokumentu sprzedaży.")
        shop_history(cur, order_id, "wystawiono dokument", number)
        conn.commit()
        return redirect(f"/sklep/orders/{order_id}")
    except ValueError as exc:
        conn.rollback()
        return str(exc), 404
    except Exception:
        conn.rollback()
        logger.exception("Shop sales document confirmation failed.")
        return "Nie udało się zatwierdzić dokumentu sprzedaży.", 500
    finally:
        conn.close()


@app.route('/sklep/orders/<int:order_id>/document/generate', methods=['POST'])
@login_required
def shop_generate_document(order_id):
    if not (can_shop("sales") or can_shop("accounting")):
        return "Brak uprawnień do generowania dokumentu sprzedaży.", 403
    conn = db()
    cur = conn.cursor()
    try:
        denied = sales_order_access_error(cur, order_id)
        if denied:
            return denied
        _document, payload = generate_shop_sales_document(
            cur, order_id, session.get("user")
        )
        shop_history(cur, order_id, "wygenerowano dokument", payload["document_number"])
        cur.execute(
            """
            INSERT INTO shop_notifications(order_id,type,message)
            VALUES (%s,'dokument wygenerowany',%s)
            """,
            (order_id, f"Dokument {payload['document_number']} jest gotowy do pobrania."),
        )
        conn.commit()
        cache.clear()
        return redirect(f"/sklep/orders/{order_id}?document=generated")
    except ValueError as exc:
        conn.rollback()
        return render_template(
            "error.html",
            title="Nie można wygenerować dokumentu",
            message=str(exc),
        ), 400
    except Exception:
        conn.rollback()
        logger.exception("Shop sales document generation failed.")
        return render_template(
            "error.html",
            title="Nie udało się wygenerować dokumentu",
            message=(
                "Dane zamówienia pozostały bez zmian. Sprawdź produkty i spróbuj ponownie."
            ),
        ), 500
    finally:
        conn.close()


@app.route('/sklep/documents/<int:doc_id>/<fmt>')
@login_required
def shop_download_document(doc_id, fmt):
    denied = require_shop_permission("view")
    if denied:
        return denied
    if fmt not in {"pdf","docx"}: return "Nieprawidłowy format.", 400
    conn=db(); cur=conn.cursor()
    if current_user_role() == "sales":
        cur.execute(
            f"""
            SELECT d.document_number,d.{fmt}
            FROM shop_sales_documents d
            JOIN shop_orders o ON o.id=d.order_id
            WHERE d.id=%s AND d.voided_at IS NULL
              AND lower(COALESCE(o.salesperson_email,''))=lower(%s)
            """,
            (doc_id, session.get("user", "")),
        )
    else:
        cur.execute(
            f"""
            SELECT document_number,{fmt} FROM shop_sales_documents
            WHERE id=%s AND voided_at IS NULL
            """,
            (doc_id,),
        )
    row=cur.fetchone(); conn.close()
    if not row or row[1] is None: return "Nie znaleziono dokumentu.", 404
    mimetype = "application/pdf" if fmt == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    filename = secure_filename(str(row[0] or "dokument")) or "dokument"
    return Response(bytes(row[1]), mimetype=mimetype, headers={"Content-Disposition": f'attachment; filename="{filename}.{fmt}"'})


@app.route('/handlowiec/raport')
@login_required
def salesperson_report():
    if current_user_role() not in {"admin", "sales"}:
        return "Brak uprawnień do raportu handlowca.", 403
    params = []
    where = ""
    if current_user_role() == "sales":
        where = "WHERE lower(COALESCE(o.salesperson_email,''))=lower(%s)"
        params.append(session.get("user", ""))
    conn = db()
    cur = conn.cursor()
    cur.execute(
        f"""
        WITH order_values AS (
            SELECT o.id,o.status,o.salesperson_email,o.salesperson_name,
                   COALESCE((
                       SELECT SUM(i.qty*i.price_brutto)
                       FROM shop_order_items i WHERE i.order_id=o.id
                   ),0)+COALESCE(o.shipping_cost,0) AS total
            FROM shop_orders o
            {where}
        )
        SELECT COALESCE(salesperson_name,salesperson_email,'Nieprzypisany'),
               COUNT(*),COALESCE(SUM(total),0),
               COUNT(*) FILTER (WHERE status='Zakończone'),
               COUNT(*) FILTER (
                   WHERE status NOT IN ('Zakończone','Anulowane')
               ),
               COUNT(*) FILTER (WHERE status='Anulowane')
        FROM order_values
        GROUP BY COALESCE(salesperson_name,salesperson_email,'Nieprzypisany')
        ORDER BY 1
        """,
        tuple(params),
    )
    rows = cur.fetchall()
    conn.close()
    return render_template("sales_report.html", rows=rows)


@app.route('/historia')
@login_required
@cache.cached(timeout=30, query_string=True)
def historia():
    conn = db()
    cur = conn.cursor()

    cur.execute("SELECT * FROM issue_docs ORDER BY id DESC")
    docs = cur.fetchall()
    conn.close()

    days = {}
    for d in docs:
        days.setdefault(d[1], []).append(d)

    return render_template("historia.html", days=days)


@app.route('/report')
@login_required
@cache.cached(timeout=30, query_string=True)
def report():
    try:
        selected_date = normalized_document_date(request.args.get("date"))
    except ValueError as exc:
        return str(exc), 400
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT d.date, COALESCE(d.kontrahent, ''), COALESCE(p.name, ''), i.qty, COALESCE(i.warehouse, p.warehouse, '')
        FROM issue_docs d
        JOIN issue_items i ON i.doc_id = d.id
        LEFT JOIN products p ON p.id = i.product_id
        WHERE d.date = %s AND d.voided_at IS NULL
        ORDER BY d.id DESC, i.id ASC
        """,
        (selected_date,)
    )
    operations = cur.fetchall()
    conn.close()
    return render_template("report.html", selected_date=selected_date, operations=operations)


@app.route('/report/pdf')
@login_required
def report_pdf():
    try:
        selected_date = normalized_document_date(request.args.get("date"))
    except ValueError as exc:
        return str(exc), 400
    conn = db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT d.date, COALESCE(d.kontrahent, ''), COALESCE(p.name, ''), i.qty, COALESCE(i.warehouse, p.warehouse, '')
        FROM issue_docs d
        JOIN issue_items i ON i.doc_id = d.id
        LEFT JOIN products p ON p.id = i.product_id
        WHERE d.date = %s AND d.voided_at IS NULL
        ORDER BY d.id DESC, i.id ASC
        """,
        (selected_date,)
    )
    operations = cur.fetchall()
    conn.close()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("Raport dzienny magazynu", styles["Title"]),
        Spacer(1, 8),
        Paragraph(f"Data: {selected_date}", styles["Normal"]),
        Spacer(1, 12),
    ]

    data = [["Data", "Kontrahent", "Produkt", "Ilość", "Magazyn"]]
    for row in operations:
        data.append([str(row[0]), str(row[1]), str(row[2]), str(row[3]), str(row[4])])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFC067")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ALIGN", (3, 1), (3, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(table)
    doc.build(elements)
    buffer.seek(0)
    return Response(
        buffer.getvalue(),
        mimetype="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=raport-{selected_date}.pdf"}
    )


if __name__ == '__main__':
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", "5000")), debug=False)
